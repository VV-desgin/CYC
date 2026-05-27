
import os
import json
import tempfile
import pickle
import io
import html
os.environ["PYTHONIOENCODING"] = "utf-8"
os.environ["PYTHONUTF8"] = "1"

import streamlit as st
import pandas as pd
import time
import re
import math
from io import BytesIO
from openai import OpenAI
from openpyxl.styles import Border, Side, Alignment, PatternFill, Font
from openpyxl.utils import get_column_letter
from docx.oxml.ns import qn

st.set_page_config(
    page_title="5G通信基建数智化交付系统",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': None,
        'Report a bug': None,
        'About': None
    }
)

# ==================== 工业科技风 CSS 主题 ====================
def inject_custom_css():
    st.markdown("""
    <style>
    :root {
        --primary: #165DFF;
        --primary-light: #E8F0FE;
        --primary-dark: #0E42D2;
        --success: #00B42A;
        --success-light: #E8F8EC;
        --warning: #FF7D00;
        --warning-light: #FFF7E8;
        --danger: #F53F3F;
        --danger-light: #FFECE8;
        --gray-50: #F5F7FA;
        --gray-100: #EEF1F5;
        --gray-200: #E1E4E8;
        --gray-300: #D1D5DB;
        --gray-400: #9CA3AF;
        --gray-500: #6B7280;
        --gray-600: #4B5563;
        --gray-700: #374151;
        --gray-800: #1F2937;
        --gray-900: #1D2129;
        --radius: 8px;
        --radius-lg: 12px;
        --shadow-card: 0 2px 8px rgba(0,0,0,0.06);
        --shadow-hover: 0 8px 24px rgba(0,0,0,0.10);
    }

    /* ===== 全局基础 ===== */
    html, body, [data-testid="stAppViewContainer"], .stMarkdown, .stDataFrame, .stButton, .stTextInput, .stSelectbox, .stTextArea, .stNumberInput, .stExpander {
        font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", sans-serif;
    }
    [data-testid="stAppViewContainer"] { background: #F5F7FA; }
    .main .block-container { padding: 24px 32px; }

    /* 隐藏 Streamlit 原生多余元素 */
    #MainMenu { visibility: hidden; }
    footer { visibility: hidden; }
    header[data-testid="stHeader"] { background: transparent; }
    [data-testid="stHeaderActionElements"] { display: none; }
    [data-testid="stDecoration"] { display: none; }

    /* ===== 侧边栏 ===== */
    section[data-testid="stSidebar"] {
        background: #FFFFFF;
        max-width: 500px !important;
        transition: width 0.3s ease !important;
        border-right: 1px solid var(--gray-200);
    }
    /* 侧边栏展开时才能达到 320px 最小宽度，折叠时还原 Streamlit 原生宽度 */
    [data-testid="stSidebar"][aria-expanded="true"] {
        min-width: 320px !important;
    }
    [data-testid="stSidebar"][aria-expanded="false"] {
        min-width: 0 !important;
        width: 0 !important;
        transform: translateX(-100%);
        opacity: 0;
        pointer-events: none;
    }
    section[data-testid="stSidebar"] p,
    section[data-testid="stSidebar"] span,
    section[data-testid="stSidebar"] label,
    section[data-testid="stSidebar"] div[data-testid="stMarkdownContainer"] {
        word-break: normal !important;
        overflow-wrap: normal !important;
        word-wrap: normal !important;
    }

    /* 侧边栏收起后主内容区 100% 撑满（直接基于 aria-expanded 属性选择器） */
    [data-testid="stSidebar"][aria-expanded="false"] ~ [data-testid="stAppViewContainer"] {
        margin-left: 0 !important;
        padding-left: 0 !important;
        width: 100vw !important;
        max-width: 100vw !important;
    }
    [data-testid="stSidebar"][aria-expanded="false"] ~ [data-testid="stAppViewContainer"] section.main {
        min-width: 100vw !important;
        max-width: 100vw !important;
        left: 0 !important;
        margin-left: 0 !important;
        padding-left: 0 !important;
    }
    [data-testid="stSidebar"][aria-expanded="false"] ~ [data-testid="stAppViewContainer"] section.main .block-container {
        margin-left: 0 !important;
        padding-left: 0 !important;
        max-width: 100vw !important;
    }

    /* 侧边栏按钮行等宽修复 */
    section[data-testid="stSidebar"] [data-testid="stHorizontalBlock"] {
        gap: 6px !important;
        flex-wrap: nowrap !important;
    }
    section[data-testid="stSidebar"] [data-testid="stHorizontalBlock"] > [data-testid="column"] {
        flex: 1 1 0% !important;
        min-width: 0 !important;
    }
    section[data-testid="stSidebar"] [data-testid="stHorizontalBlock"] button {
        width: 100% !important;
        min-width: 0 !important;
        padding: 4px 2px !important;
        font-size: 0.75rem !important;
        white-space: nowrap !important;
        overflow: hidden !important;
        text-overflow: ellipsis !important;
    }

    /* 主内容区自适应侧边栏宽度（仅侧边栏展开时生效） */
    [data-testid="stSidebar"][aria-expanded="true"] ~ [data-testid="stAppViewContainer"] section.main {
        min-width: calc(100vw - 500px) !important;
        max-width: calc(100vw - 320px) !important;
        transition: min-width 0.3s ease, max-width 0.3s ease !important;
    }
    .main-content-expanded {
        margin-left: 0 !important;
        padding-left: 0 !important;
    }

    /* ===== 主内容区标题 ===== */
    .page-title {
        font-size: 2rem !important;
        font-weight: 800 !important;
        color: var(--gray-900) !important;
        margin-bottom: 8px !important;
        padding-bottom: 12px !important;
        border-bottom: 2px solid var(--primary);
    }
    .page-subtitle {
        font-size: 0.85rem !important;
        color: var(--gray-500) !important;
        margin-bottom: 20px !important;
    }
    .section-title {
        font-size: 0.95rem !important;
        font-weight: 600 !important;
        color: var(--gray-800) !important;
        padding-bottom: 8px;
        margin-bottom: 16px;
        border-bottom: 1px solid var(--gray-200);
    }

    /* ===== 按钮 ===== */
    .stButton > button {
        border-radius: var(--radius) !important;
        font-weight: 500 !important;
        padding: 8px 20px !important;
        font-size: 0.85rem !important;
        transition: all 0.15s ease !important;
    }
    .stButton > button[kind="primary"] {
        background: var(--primary) !important;
        color: #fff !important;
        border: none !important;
    }
    .stButton > button[kind="primary"]:hover {
        background: var(--primary-dark) !important;
        transform: translateY(-1px);
    }
    .stButton > button[kind="secondary"] {
        background: transparent !important;
        color: var(--primary) !important;
        border: 1px solid var(--primary) !important;
    }
    .stButton > button[kind="secondary"]:hover {
        background: var(--primary-light) !important;
    }

    /* ===== 工程概览卡片 ===== */
    .metric-card {
        background: #FFFFFF;
        border: 1px solid var(--gray-200);
        border-left: 4px solid var(--primary);
        border-radius: var(--radius-lg);
        padding: 24px 20px;
        box-shadow: var(--shadow-card);
        text-align: left;
        display: flex;
        flex-direction: column;
        justify-content: center;
        min-height: 160px;
        height: 100%;
        width: 100%;
        box-sizing: border-box;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
        overflow: hidden;
    }
    .metric-card:hover {
        transform: translateY(-4px);
        box-shadow: var(--shadow-hover);
    }
    .metric-card .label {
        font-size: 0.75rem;
        font-weight: 500;
        color: var(--gray-500);
        margin-bottom: 6px;
    }
    .metric-card .value {
        font-size: 2.5rem;
        font-weight: 800;
        color: var(--gray-900);
        line-height: 1.2;
        margin-bottom: 2px;
    }
    .metric-card .suffix {
        font-size: 0.9rem;
        font-weight: 400;
        color: var(--gray-400);
        margin-left: 4px;
    }

    /* 卡片状态色条 */
    .metric-card.status-completed { border-left-color: var(--success); }
    .metric-card.status-in-progress { border-left-color: var(--warning); }
    .metric-card.status-pending { border-left-color: var(--gray-300); }
    .metric-card.status-info { border-left-color: var(--primary); }

    /* 工程概览列等高 */
    [data-testid="stHorizontalBlock"] {
        align-items: stretch !important;
        gap: 20px !important;
    }
    [data-testid="stHorizontalBlock"] > div { display: flex; }

    /* ===== 选择框 / 单选框 ===== */
    .stSelectbox [data-baseweb="select"] > div {
        background: #fff !important;
        border: 1px solid var(--gray-200) !important;
        border-radius: var(--radius) !important;
    }
    .stSelectbox [data-baseweb="select"] > div:focus-within {
        border-color: var(--primary) !important;
        box-shadow: 0 0 0 3px rgba(22,93,255,0.1) !important;
    }
    [role="radiogroup"],
    section[data-testid="stSidebar"] [role="radiogroup"] {
        background: transparent !important;
        border: none !important;
        border-radius: var(--radius) !important;
        padding: 0 !important;
        gap: 2px !important;
        width: 100% !important;
        box-sizing: border-box !important;
        display: flex !important;
        flex-direction: column !important;
    }
    [role="radiogroup"] label,
    section[data-testid="stSidebar"] [role="radiogroup"] label {
        border-radius: 6px !important;
        padding: 6px 14px !important;
        font-size: 0.82rem !important;
        font-weight: 500 !important;
        color: var(--gray-600) !important;
        transition: all 0.15s !important;
        width: 100% !important;
        box-sizing: border-box !important;
        flex: 1 !important;
        min-width: unset !important;
        max-width: 100% !important;
    }
    [role="radiogroup"] label[data-selected="true"] {
        background: #fff !important;
        color: var(--primary) !important;
        box-shadow: 0 1px 2px rgba(0,0,0,0.06);
    }

    /* ===== Tabs 标签页 ===== */
    .stTabs [data-baseweb="tab-list"] {
        gap: 6px;
        background: var(--gray-50);
        border-radius: var(--radius);
        padding: 4px;
        margin-bottom: 16px;
    }
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px !important;
        font-weight: 500 !important;
        font-size: 0.85rem !important;
        padding: 8px 18px !important;
        color: var(--gray-600) !important;
    }
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        background: #fff !important;
        color: var(--primary) !important;
        box-shadow: 0 1px 2px rgba(0,0,0,0.06);
    }

    /* ===== DataFrame 表格 ===== */
    [data-testid="stDataFrame"] {
        border-radius: var(--radius) !important;
        border: 1px solid var(--gray-200) !important;
        overflow: hidden;
    }



    /* ===== 状态徽章 ===== */
    .status-badge {
        display: inline-flex;
        align-items: center;
        padding: 3px 10px;
        border-radius: 20px;
        font-size: 0.72rem;
        font-weight: 600;
    }
    .status-badge.success {
        background: var(--success-light);
        color: var(--success);
    }
    .status-badge.warning {
        background: var(--warning-light);
        color: #C95D00;
    }
    .status-badge.danger {
        background: var(--danger-light);
        color: var(--danger);
    }
    .status-badge.info {
        background: var(--primary-light);
        color: var(--primary);
    }

    /* ===== 通知横幅与容器 ===== */
    .completion-banner {
        background: var(--success-light);
        border: 1px solid #6EE7B7;
        border-radius: var(--radius-lg);
        padding: 14px 20px;
        margin: 12px 0;
        font-weight: 600;
        color: var(--success);
        box-shadow: var(--shadow-card);
    }
    .download-section {
        background: var(--primary-light);
        border: 1px solid #BEDAFF;
        border-radius: var(--radius-lg);
        padding: 18px;
        margin: 12px 0;
    }
    .config-guide-box {
        background: var(--warning-light);
        border: 1px solid #FFD591;
        border-radius: var(--radius);
        padding: 14px 16px;
        margin: 10px 0;
        font-size: 0.82rem;
        line-height: 1.6;
    }
    .config-guide-box code {
        background: #fff;
        padding: 2px 6px;
        border-radius: 4px;
        font-size: 0.78rem;
        color: var(--primary);
    }
    .download-container {
        border: 1px solid var(--gray-200) !important;
        border-radius: var(--radius) !important;
        padding: 16px !important;
        margin-bottom: 16px !important;
        background: #FFFFFF !important;
    }

    /* ===== 展开面板 ===== */
    [data-testid="stExpander"] details {
        border: 1px solid var(--gray-200) !important;
        border-radius: var(--radius) !important;
    }

    /* ===== 响应式适配 ===== */
    @media (max-width: 800px) {
        section[data-testid="stSidebar"] {
            min-width: 180px !important;
            max-width: 300px !important;
        }
        .sidebar-content { font-size: 0.9rem !important; }
    }
    @media (max-width: 600px) {
        section[data-testid="stSidebar"] {
            width: 0 !important;
            min-width: 0 !important;
        }
    }
    </style>
    """, unsafe_allow_html=True)

inject_custom_css()


# ==================== 新字段模板（16项） ====================
NEW_FIELD_TEMPLATE = [
    "站点编号", "站点名称", "站点类型", "网络制式", "覆盖场景",
    "线缆类型", "路由长度(m)", "设备总台数", "接地设备数量", "室外接头数量",
    "光口配线对数", "取电方式", "起点", "终点", "BBU型号", "AAU型号"
]

REQUIRED_FIELDS = ["站点编号", "站点名称", "站点类型", "网络制式", "覆盖场景", "取电方式"]

ENUMS = {
    "站点类型": ["室外宏站", "室内分布", "杆塔站", "塔房站", "其他"],
    "网络制式": ["4G", "5G", "4G+5G", "其他"],
    "覆盖场景": ["城区", "园区", "校园", "楼宇", "隧道", "其他"],
    "线缆类型": ["光缆", "电源线", "馈线", "接地线", "网线", "其他"],
    "取电方式": ["市电直供", "直流远供", "太阳能+储能", "混合供电", "其他"],
}

BOM_FIELDS = ["编号", "项目编码", "专业类别", "设备/材料名称", "规格型号", "单位", "数量", "项目特征", "工程量计算规则", "工作内容", "安装位置", "备注"]

DEMO_DATA = {
    "站点编号": ["SZ-BS-001", "SZ-BS-002", "GZ-BS-001", "BJ-BS-001", "SH-BS-001"],
    "站点名称": ["深圳大学城1号站", "深圳科技园2号站", "广州天河3号站", "北京海淀4号站", "上海浦东5号站"],
    "站点类型": ["室外宏站", "杆塔站", "室内分布", "塔房站", "室外宏站"],
    "网络制式": ["5G", "4G+5G", "4G", "5G", "4G+5G"],
    "覆盖场景": ["城区", "园区", "楼宇", "城区", "校园"],
    "线缆类型": ["光缆", "电源线", "馈线", "光缆", "光缆"],
    "路由长度(m)": [200, 350, 180, 420, 280],
    "设备总台数": [3, 4, 2, 3, 3],
    "接地设备数量": [2, 2, 1, 2, 2],
    "室外接头数量": [2, 3, 1, 3, 2],
    "光口配线对数": [24, 12, 8, 24, 12],
    "取电方式": ["市电直供", "直流远供", "市电直供", "混合供电", "市电直供"],
    "起点": ["BBU-P1", "AAU-P1", "ODF-P1", "BBU-P2", "AAU-P2"],
    "终点": ["AAU-P1", "ODF-P2", "BBU-P1", "AAU-P1", "ODF-P1"],
    "BBU型号": ["BBU5900", "BBU5900", "BBU3910", "BBU5900", "BBU5900"],
    "AAU型号": ["AAU5613", "AAU5636", "AAU5339", "AAU5639", "AAU5613"],
}
DEMO_DF = pd.DataFrame(DEMO_DATA)

PLATFORM_PRESETS = {
    "硅基流动": {"base_url": "https://api.siliconflow.cn/v1", "model": "deepseek-ai/DeepSeek-V3"},
    "DeepSeek": {"base_url": "https://api.deepseek.com", "model": "deepseek-chat"},
    "OpenAI": {"base_url": "https://api.openai.com/v1", "model": "gpt-4o"},
    "通义千问(阿里云百炼)": {"base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1", "model": "qwen-plus"},
    "智谱AI": {"base_url": "https://open.bigmodel.cn/api/paas/v4", "model": "glm-4"},
    "百度千帆": {"base_url": "https://qianfan.baidubce.com/v2", "model": "ernie-4.0-8k"},
    "火山方舟": {"base_url": "https://ark.cn-beijing.volces.com/api/v3", "model": "doubao-pro-32k"},
    "腾讯混元": {"base_url": "https://api.hunyuan.cloud.tencent.com/v1", "model": "hunyuan-pro"},
    "Groq": {"base_url": "https://api.groq.com/openai/v1", "model": "llama-3.3-70b-versatile"},
    "自定义/本地": {"base_url": "http://localhost:11434/v1", "model": "qwen2.5:7b"},
}

PLATFORM_KEY_RULES = {
    "硅基流动":     {"prefix": "sk-",  "min_len": 40, "max_len": 64,  "pattern": r"^sk-[a-zA-Z0-9]{35,}$"},
    "DeepSeek":     {"prefix": "sk-",  "min_len": 30, "max_len": 64,  "pattern": r"^sk-[a-zA-Z0-9]{28,}$"},
    "OpenAI":       {"prefix": "sk-",  "min_len": 40, "max_len": 200, "pattern": r"^sk-(?:proj-[a-zA-Z0-9_-]+-)?[a-zA-Z0-9]{32,}$"},
    "通义千问(阿里云百炼)": {"prefix": "sk-",  "min_len": 25, "max_len": 64,  "pattern": r"^sk-[a-zA-Z0-9]{20,}$"},
    "智谱AI":       {"prefix": "",     "min_len": 25, "max_len": 64,  "pattern": r"^[a-zA-Z0-9]{24,}(?:\.[a-zA-Z0-9]+){2}$"},
    "百度千帆":     {"prefix": "",     "min_len": 30, "max_len": 128, "pattern": r"^[a-zA-Z0-9._-]{30,}$"},
    "火山方舟":     {"prefix": "",     "min_len": 25, "max_len": 64,  "pattern": r"^[a-zA-Z0-9_-]{24,}$"},
    "腾讯混元":     {"prefix": "sk-",  "min_len": 30, "max_len": 64,  "pattern": r"^sk-[a-zA-Z0-9]{28,}$"},
    "Groq":         {"prefix": "gsk_", "min_len": 40, "max_len": 64,  "pattern": r"^gsk_[a-zA-Z0-9]{35,}$"},
    "自定义/本地": {"prefix": "",     "min_len": 0,  "max_len": 512, "pattern": r"^.*$"},
}


def validate_api_key(key, platform):
    """强校验 API Key 格式、前缀、长度。

    返回 (is_valid, error_message)。
    规则来自 PLATFORM_KEY_RULES，对"自定义/本地"仅做非空检查。
    """
    if not key or not key.strip():
        return False, "API Key 不能为空"

    key = key.strip()
    rule = PLATFORM_KEY_RULES.get(platform)

    if rule is None:
        return True, ""

    # 长度校验
    if len(key) < rule["min_len"]:
        return False, f"Key 过短（{len(key)}字符，至少{rule['min_len']}字符）"
    if len(key) > rule["max_len"]:
        return False, f"Key 过长（{len(key)}字符，上限{rule['max_len']}字符）"

    # 前缀校验
    if rule["prefix"] and not key.startswith(rule["prefix"]):
        return False, f"Key 前缀错误，应为 {rule['prefix']}... 开头"

    # 正则格式校验
    if rule.get("pattern"):
        if not re.match(rule["pattern"], key):
            return False, f"Key 格式不符合 {platform} 规范"

    # 除自定义/本地的常规平台，必须 >= min_len（再次确保）
    if platform != "自定义/本地" and len(key) < rule["min_len"]:
        return False, f"Key 长度不足（{len(key)}字符，至少{rule['min_len']}字符）"

    return True, ""

AI_STEPS = [
    {"key": "bom", "label": "施工BOM",
     "system": "你是5G基站施工BOM专家。按12项BOM字段生成物料清单：编号、项目编码、专业类别、设备/材料名称、规格型号、单位、数量、项目特征、工程量计算规则、工作内容、安装位置、备注。必须以标准Markdown表格格式输出，第一行为列名，第二行为分隔线，之后每行一条数据。每行必须以 | 开头和结尾。"},
    {"key": "bor", "label": "资源需求清单",
     "system": "你是5G基站施工管理专家。汇总全部站点生成资源需求清单。必须以标准Markdown表格格式输出，列名：资源类别、资源名称、规格/型号、数量、单位、备注。第一行为列名，第二行为分隔线，之后每行一条数据。每行必须以 | 开头和结尾。"},
    {"key": "bop", "label": "工艺指导书",
     "system": "你是5G基站施工工艺专家。按6章结构生成工艺指导书：一、站点勘测与规划 二、设备安装要求 三、线缆敷设要求 四、取电与接地要求 五、质量检查标准 六、施工注意事项。直接输出。"},
    {"key": "fiber", "label": "纤芯分配表",
     "system": "你是5G光缆工程专家。生成纤芯分配表。必须以标准Markdown表格格式输出，列名：站点编号、光缆编号、纤芯序号、纤芯颜色、起始端子、终止端子、业务类型。第一行为列名，第二行为分隔线，之后每行一条数据。每行必须以 | 开头和结尾。"},
    {"key": "risk", "label": "风险提示",
     "system": "你是5G通信基建审查专家。依据YD/T 5264-2021分析全部站点数据，输出施工风险提示、注意事项、合规建议。"},
]

# ==================== Session State 初始化 ====================
def init_session_state():
    defaults = {
        "uploaded_files": [{"name": "样例数据", "df": DEMO_DF, "valid": True, "errors": []}],
        "current_idx": 0,
        "result_df": None,
        "pending_upload": None,
        "ai_data": {},
        "ai_running": False,
        "ai_generation_done": False,
        "ai_step_index": 0,
        "ai_start_time": 0,
        "ai_step_times": {},
        "ai_timeout": False,
        "offline_mode": False,
        "review_results": None,
        "ai_dataframes": {},              # v3: AI生成的DataFrame存储
        "use_builtin_ai": True,          # v3: 默认内置模式
        "ai_platform": "硅基流动",
        "ai_base_url": "https://api.siliconflow.cn/v1",
        "ai_model": "deepseek-ai/DeepSeek-V3",
        "ai_api_key": "",
        "ai_config_expanded": False,     # v3: 默认折叠，简化界面
        "api_connection_verified": False,  # 测试连接是否验证通过
        "show_api_guide": False,          # v3: API获取指引
        "show_file_detail": False,
        "expand_raw_data": False,       # 上传文件后自动展开原始数据
        "excel_bytes": None,
        "word_report_bytes": None,
        "word_bop_bytes": None,
        "review_failures_bytes": None,    # v3: 合规审查不通过项Excel
        "compliance_full_excel": None,   # v3: 合规审查完整报告（3sheet）
        "review_failed": False,           # v3: 合规审查是否失败
        "task_history": [],               # v5: 任务历史记录
        "sidebar_collapsed": False,
        "sidebar_width": 280,
        "validation_errors": {},          # v3: 文件验证错误记录
        "ai_builtin_running": False,       # v4: 内置模式分步执行中
        "ai_builtin_step": 0,              # v4: 0=空闲, 1=合规审查, 2=BOM, 3=BOR, 4=BOP, 5=纤芯, 6=风险, 7=导出, 8=完成
        "ai_builtin_start_time": 0,        # v4: 内置模式开始时间
        "ai_builtin_sites": None,           # v4: 处理中的站点数据副本
        "cr_chunk_start": 0,               # v4: 合规审查分块起始行
        "cr_chunk_results": None,           # v4: 合规审查累积结果
        "cr_chunk_cp": None,               # v4: 合规审查累积连接对
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_session_state()

# ==================== 侧边栏折叠控制 ====================
# 折叠时注入 CSS 隐藏侧边栏并调整主页面
if st.session_state.sidebar_collapsed:
    st.markdown("""
    <style>
    section[data-testid="stSidebar"] { 
        display: none !important; 
        width: 0 !important;
        min-width: 0 !important;
        max-width: 0 !important;
        flex: 0 0 0 !important;
    }
    section.main, section[data-testid="stMain"], .stMain {
        width: 100vw !important; 
        margin-left: 0 !important;
        padding-left: 0 !important;
        min-width: 100vw !important;
        max-width: 100vw !important;
    }
    .stAppViewContainer, [data-testid="stAppViewContainer"] {
        margin-left: 0 !important;
        padding-left: 0 !important;
        width: 100vw !important;
    }
    .stApp .block-container, section.main .block-container {
        max-width: 100vw !important;
        padding-left: 1rem !important;
        padding-right: 1rem !important;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # 展开按钮
    col_expand_btn = st.columns([0.03, 0.97])
    with col_expand_btn[0]:
        if st.button("▶", key="sidebar_expand_btn", help="展开侧边栏"):
            st.session_state.sidebar_collapsed = False
            st.rerun()

# ==================== v3 新增：文件持久化 ====================
def save_uploaded_files():
    """保存上传文件到临时目录"""
    try:
        data_to_save = []
        for item in st.session_state.uploaded_files:
            # 将 DataFrame 转换为 CSV 字符串保存
            if "df" in item:
                csv_str = item["df"].to_csv(index=False)
                data_item = {
                    "name": item["name"],
                    "csv_data": csv_str,
                    "valid": item.get("valid", True),
                    "errors": item.get("errors", []),
                    "warnings": item.get("warnings", [])
                }
                data_to_save.append(data_item)
        
        # 保存到临时文件
        temp_dir = tempfile.gettempdir()
        save_path = os.path.join(temp_dir, "5g_delivery_system_uploaded_files.json")
        with open(save_path, "w", encoding="utf-8") as f:
            json.dump(data_to_save, f, ensure_ascii=False)
    except Exception as e:
        st.error(f"保存文件失败: {e}")

def load_uploaded_files():
    """从临时目录加载上传文件"""
    try:
        temp_dir = tempfile.gettempdir()
        save_path = os.path.join(temp_dir, "5g_delivery_system_uploaded_files.json")
        if os.path.exists(save_path):
            with open(save_path, "r", encoding="utf-8") as f:
                loaded_data = json.load(f)
            
            loaded_files = []
            for item in loaded_data:
                # 从 CSV 字符串恢复 DataFrame
                df = pd.read_csv(io.StringIO(item["csv_data"]))
                loaded_item = {
                    "name": item["name"],
                    "df": df,
                    "valid": item.get("valid", True),
                    "errors": item.get("errors", []),
                    "warnings": item.get("warnings", [])
                }
                loaded_files.append(loaded_item)
            
            # 如果当前 session 没有文件，则加载
            if not st.session_state.uploaded_files or len(st.session_state.uploaded_files) == 1:
                st.session_state.uploaded_files = loaded_files
            elif loaded_files:
                # 合并已加载文件，避免重复
                existing_names = [f["name"] for f in st.session_state.uploaded_files]
                for f in loaded_files:
                    if f["name"] not in existing_names:
                        st.session_state.uploaded_files.append(f)
    except Exception as e:
        # 加载失败不影响正常使用
        pass

# 页面初始化时加载
load_uploaded_files()

# ==================== v3 新增：文件上传验证 ====================
def validate_uploaded_file(df, filename):
    """验证上传文件的字段完整性和数据有效性"""
    errors = []
    warnings = []

    # 检查必备字段是否缺失
    missing_fields = [c for c in REQUIRED_FIELDS if c not in df.columns]
    if missing_fields:
        errors.append(f"缺少必备字段: {', '.join(missing_fields)}")

    # 检查全部16项模板字段
    missing_new = [c for c in NEW_FIELD_TEMPLATE if c not in df.columns]
    if missing_new:
        warnings.append(f"缺少推荐字段(不影响使用): {', '.join(missing_new)}")

    # 检查空行
    if "站点编号" in df.columns:
        empty_ids = df["站点编号"].isna().sum() + (df["站点编号"].astype(str).str.strip() == "").sum()
        if empty_ids > 0:
            errors.append(f"存在{empty_ids}行站点编号为空，将被跳过")

    # 检查路由长度合法性
    if "路由长度(m)" in df.columns:
        try:
            route_len = df["路由长度(m)"].apply(safe_float_route)
            invalid_routes = (route_len == 0).sum()
            if invalid_routes > 0:
                errors.append(f"存在{invalid_routes}行路由长度为0或无法解析")
            zero_routes = (route_len < 0).sum()
            if zero_routes > 0:
                warnings.append(f"存在{zero_routes}行路由长度<0，将触发合规审查告警")
        except Exception:
            errors.append("路由长度列格式异常，无法解析")

    # 检查枚举值有效性
    for field, valid_values in ENUMS.items():
        if field in df.columns:
            actual_values = df[field].dropna().astype(str).unique()
            invalid = [v for v in actual_values if v not in valid_values and v != "nan"]
            if invalid:
                warnings.append(f"字段'{field}'包含非标准值: {', '.join(invalid[:5])}")

    is_valid = len(errors) == 0
    return is_valid, errors, warnings, missing_new


def read_file(f):
    """读取上传文件，支持 CSV 和 Excel"""
    try:
        if f.name.endswith(".csv"):
            return pd.read_csv(f)
        else:
            return pd.read_excel(f)
    except Exception as e:
        return None


def get_client(url, key, platform=None):
    """创建 OpenAI 客户端，含 Key 格式预校验。

    若 platform 非空，先调用 validate_api_key 做格式/长度/前缀校验，
    不通过则直接抛出 ValueError，避免无效 Key 进入网络请求。
    """
    if platform and platform != "自定义/本地":
        is_valid, err_msg = validate_api_key(key, platform)
        if not is_valid:
            raise ValueError(f"API Key 校验失败: {err_msg}")

    keys = [k.strip() for k in key.split(",") if k.strip()]
    return OpenAI(base_url=normalize_base_url(url, platform), api_key=keys[0] if keys else key)


def normalize_base_url(url, platform=None):
    """规范化 API 地址：自定义/本地模式下自动补全 /v1 路径。

    OpenAI SDK 约定 base_url 应指向 /v1 根路径，拼装请求时为 {base_url}/chat/completions。
    用户可能只输入 http://host:port，需自动补全为标准格式。
    """
    import re as _re
    url = url.strip().rstrip("/")
    if platform == "自定义/本地":
        # 已含 /v1 后缀的保持不动
        if not _re.search(r'/v\d+$', url):
            url += "/v1"
    return url


def parse_markdown_table(text):
    """从 Markdown 文本中提取表格为 DataFrame。多策略解析。"""
    if not text:
        return None

    lines = text.strip().split('\n')

    # ---- 策略 A：标准 Markdown 表格 ----
    table_lines = []
    in_table = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if in_table and table_lines:
                break
            continue
        is_pipe_line = stripped.startswith('|') or stripped.endswith('|') or '|' in stripped
        if is_pipe_line:
            is_sep = bool(re.match(r'^[\s\-:|]+$', stripped))
            if is_sep:
                in_table = True
                continue
            table_lines.append(stripped)
            in_table = True
        else:
            if in_table and table_lines:
                break
            in_table = False

    if len(table_lines) >= 2:
        df = _lines_to_df(table_lines)
        if df is not None and not df.empty:
            return df

    # ---- 策略 B：按 | 分割所有行 ----
    pipe_lines = [l.strip() for l in lines if l.strip() and '|' in l.strip()]
    pipe_lines = [l for l in pipe_lines if not re.match(r'^[\s\-:|]+$', l)]
    if len(pipe_lines) >= 2:
        df = _lines_to_df(pipe_lines)
        if df is not None and not df.empty:
            return df

    # ---- 策略 C：pd.read_csv 自动检测 ----
    try:
        df = pd.read_csv(io.StringIO(text), sep=None, engine='python', on_bad_lines='skip')
        df = df.dropna(axis=1, how='all').dropna(axis=0, how='all')
        if not df.empty and len(df.columns) >= 2:
            # 清理 Unnamed 列名
            df.columns = [c if not str(c).startswith('Unnamed') else '' for c in df.columns]
            return df
    except Exception:
        pass

    return None


def _lines_to_df(table_lines):
    """将 | 分隔的行列表转为 DataFrame，自动检测表头。"""
    # 归一化：去掉首尾 |
    normalized = []
    for tl in table_lines:
        t = tl.strip()
        if t.startswith('|'):
            t = t[1:]
        if t.endswith('|'):
            t = t[:-1]
        normalized.append(t)

    # 解析单元格
    parsed_rows, max_cols = [], 0
    for line in normalized:
        cells = [cell.strip() for cell in line.split('|')]
        max_cols = max(max_cols, len(cells))
        parsed_rows.append(cells)

    # 补齐列数
    for row in parsed_rows:
        row += [''] * (max_cols - len(row))

    if len(parsed_rows) < 2:
        return None

    # 检测第一行是否为表头（表头通常全文本，数据行常含数字）
    first_row = parsed_rows[0]
    has_numeric_in_first = any(re.search(r'\d', c) for c in first_row if c)
    second_row = parsed_rows[1]
    has_numeric_in_second = any(re.search(r'\d', c) for c in second_row if c)

    if has_numeric_in_first and not has_numeric_in_second:
        # 第一行更像是数据，第二行更像表头 → 交换
        header = second_row
        data = [first_row] + parsed_rows[2:]
    else:
        header = first_row
        data = parsed_rows[1:]

    # 确保 header 无空名且唯一
    unique_header = []
    for j, h in enumerate(header):
        h = h.strip()
        if not h:
            h = f'列{j}'
        # 确保唯一性
        base = h
        counter = 1
        while h in unique_header:
            h = f'{base}_{counter}'
            counter += 1
        unique_header.append(h)
    
    df = pd.DataFrame(data, columns=unique_header)

    # 去全空列
    empty_cols = [c for c in df.columns if df[c].astype(str).str.strip().eq('').all()]
    if empty_cols:
        df = df.drop(columns=empty_cols)

    return df


# 设备厂家映射表（型号 → 厂家）
BBU_VENDOR = {
    "BBU5900": "华为", "BBU5901": "华为", "BBU5905": "华为",
    "BBU3910": "华为", "BBU3900": "华为", "BBU5910": "华为",
    "BBU V9200": "中兴", "BBU V9600": "中兴", "BBU V9800": "中兴",
    "BBU 6648": "爱立信", "BBU 6651": "爱立信",
}
AAU_VENDOR = {
    "AAU5336": "华为", "AAU5339": "华为", "AAU5613": "华为",
    "AAU5619": "华为", "AAU5636": "华为", "AAU5639": "华为",
    "AAU5726": "华为", "AAU5910": "华为",
    "AAU 7221": "中兴", "AAU 7528": "中兴", "AAU 7721": "中兴",
    "AAU AIR 6488": "爱立信", "AAU AIR 6419": "爱立信",
}

def get_device_vendor(bbu, aau):
    """根据设备型号查厂家，返回 (BBU厂家, AAU厂家)，未知则返回空"""
    bvu = BBU_VENDOR.get(bbu.strip() if bbu else "", "")
    avu = AAU_VENDOR.get(aau.strip() if aau else "", "")
    return bvu, avu


# ==================== 中文数字转阿拉伯数字 ====================
_CN_NUM_MAP = {
    "零": 0, "一": 1, "二": 2, "两": 2, "三": 3, "四": 4,
    "五": 5, "六": 6, "七": 7, "八": 8, "九": 9,
    "十": 10, "百": 100, "千": 1000, "万": 10000
}


def chinese_to_arabic(cn_str):
    """将中文数字字符串转为阿拉伯整数。

    支持范围 0~9999，含"四百二十"、"十五"、"一百零五"等格式。
    无法识别时返回 None。
    """
    if not cn_str or not isinstance(cn_str, str):
        return None
    cn_str = cn_str.strip()
    if not cn_str:
        return None

    # 已是纯数字则直接返回
    if cn_str.isdigit():
        return int(cn_str)

    # 过滤非中文数字字符（如"米"、"m"等后缀）
    cleaned = "".join(ch for ch in cn_str if ch in _CN_NUM_MAP)
    if not cleaned:
        return None

    # 处理零
    if cleaned == "零":
        return 0

    total = 0
    num = 0

    for ch in cleaned:
        val = _CN_NUM_MAP[ch]

        if val >= 10:
            # 单位：十/百/千/万
            if num == 0:
                num = 1
            total += num * val
            num = 0
        else:
            # 数字：0-9
            num = val

    total += num
    return total if total > 0 else 0


def safe_float_route(value):
    """安全转换路由长度(m)字段为 float。

    按顺序尝试：直接 float → 中文数字转换 → str.isdigit → 正则提取。
    全部失败返回 0.0，绝不抛异常。
    """
    if value is None:
        return 0.0
    if isinstance(value, (int, float)):
        v = float(value)
        if math.isnan(v) or math.isinf(v):
            return 0.0
        return max(v, 0.0)

    s = str(value).strip()
    if not s or s.lower() == "nan":
        return 0.0

    # 1) 直接 float 转换
    try:
        v = float(s)
        return max(v, 0.0)
    except (ValueError, TypeError):
        pass

    # 2) 中文数字
    arabic = chinese_to_arabic(s)
    if arabic is not None:
        return float(arabic)

    # 3) 纯数字字符串（含小数点）
    cleaned = "".join(ch for ch in s if ch.isdigit() or ch == ".")
    if cleaned and cleaned != ".":
        try:
            v = float(cleaned)
            return max(v, 0.0)
        except ValueError:
            pass

    # 4) 兜底
    return 0.0


def run_compliance_review(df, start_row=0, chunk_size=None, prev_results=None, prev_connection_pairs=None):
    """合规审查。支持分块处理：传入chunk_size时仅处理[start_row, start_row+chunk_size)行，
    返回 (done, results, connection_pairs, next_start)，其中done表示是否处理完所有行。"""
    results = prev_results if prev_results is not None else []
    connection_pairs = prev_connection_pairs if prev_connection_pairs is not None else {}
    is_chunked = chunk_size is not None and chunk_size > 0
    end_row = min(start_row + chunk_size, len(df)) if is_chunked else len(df)
    
    for idx in range(start_row, end_row):
        row = df.iloc[idx]
        site_id = str(row.get("站点编号", ""))
        site_type = str(row.get("站点类型", ""))
        network = str(row.get("网络制式", ""))
        route_len = safe_float_route(row.get("路由长度(m)", 0))
        cable_type = str(row.get("线缆类型", ""))
        grounding = int(safe_float_route(row.get("接地设备数量", 0)) or 0)
        fiber_pairs = int(safe_float_route(row.get("光口配线对数", 0)) or 0)
        power_mode = str(row.get("取电方式", ""))
        start = str(row.get("起点", ""))
        end = str(row.get("终点", ""))
        bbu = str(row.get("BBU型号", ""))
        aau = str(row.get("AAU型号", ""))
        outdoor_types = ["室外宏站", "杆塔站", "塔房站"]

        if re.match(r'^[A-Z0-9-]{6,25}$', site_id):
            results.append({"站点": site_id, "规则": "RK-001", "结果": "通过", "风险": "建议", "提示": "站点编号格式正确"})
        else:
            results.append({"站点": site_id, "规则": "RK-001", "结果": "失败", "风险": "高风险", "提示": f"站点编号'{site_id}'格式不符，需满足大写字母+数字+连字符，6-25位"})

        if site_type in ENUMS["站点类型"]:
            results.append({"站点": site_id, "规则": "RK-002", "结果": "通过", "风险": "建议", "提示": "站点类型有效"})
        else:
            results.append({"站点": site_id, "规则": "RK-002", "结果": "失败", "风险": "高风险", "提示": f"站点类型'{site_type}'不在枚举范围内"})

        if network in ENUMS["网络制式"]:
            results.append({"站点": site_id, "规则": "RK-003", "结果": "通过", "风险": "建议", "提示": "网络制式有效"})
        else:
            results.append({"站点": site_id, "规则": "RK-003", "结果": "失败", "风险": "高风险", "提示": f"网络制式'{network}'不在枚举范围内"})

        if route_len > 0 and not cable_type:
            results.append({"站点": site_id, "规则": "RK-004", "结果": "失败", "风险": "高风险", "提示": "路由长度>0时必须填写线缆类型"})
        else:
            results.append({"站点": site_id, "规则": "RK-004", "结果": "通过", "风险": "建议", "提示": "线缆类型已填写"})

        if site_type in outdoor_types and grounding < 1:
            results.append({"站点": site_id, "规则": "RK-005", "结果": "失败", "风险": "高风险", "提示": f"室外站点必须填写接地设备数量且>=1"})
        else:
            results.append({"站点": site_id, "规则": "RK-005", "结果": "通过", "风险": "建议", "提示": "接地设备数量满足要求"})

        if network in ("5G", "4G+5G") and fiber_pairs < 12:
            results.append({"站点": site_id, "规则": "RK-006", "结果": "失败", "风险": "高风险", "提示": f"5G站点光口配线对数必须>=12，当前为{fiber_pairs}"})
        else:
            results.append({"站点": site_id, "规则": "RK-006", "结果": "通过", "风险": "建议", "提示": "光口配线对数满足要求"})

        if start and end and start == end:
            results.append({"站点": site_id, "规则": "RK-007", "结果": "失败", "风险": "中风险", "提示": "起点与终点不能相同"})
        else:
            results.append({"站点": site_id, "规则": "RK-007", "结果": "通过", "风险": "建议", "提示": "起止点有效"})

        if "直流远供" in power_mode and route_len > 150:
            results.append({"站点": site_id, "规则": "RK-008", "结果": "警告", "风险": "高风险", "提示": f"直流远供线路长度{route_len}米>150米，高风险"})
        else:
            results.append({"站点": site_id, "规则": "RK-008", "结果": "通过", "风险": "建议", "提示": "直流远供长度合规"})

        # RK-009: 端口重复连接冲突 — 检测是否存在重复的(起点, 终点)连接对
        line_start = str(row.get("起点", "")).strip()
        line_end = str(row.get("终点", "")).strip()
        pair_key = f"{line_start} → {line_end}"
        if line_start and line_end and line_start != "nan" and line_end != "nan":
            if pair_key not in connection_pairs:
                connection_pairs[pair_key] = []
            connection_pairs[pair_key].append(site_id)

        # RK-010: BBU与AAU厂家一致性（基于型号映射，非字符串前缀）
        if bbu and aau and bbu != "nan" and aau != "nan":
            bbu_vendor, aau_vendor = get_device_vendor(bbu, aau)
            if bbu_vendor and aau_vendor:
                if bbu_vendor != aau_vendor:
                    results.append({"站点": site_id, "规则": "RK-010", "结果": "警告", "风险": "建议", 
                                   "提示": f"BBU({bbu})厂家[{bbu_vendor}]与AAU({aau})厂家[{aau_vendor}]不一致，建议统一"})
                else:
                    results.append({"站点": site_id, "规则": "RK-010", "结果": "通过", "风险": "建议", 
                                   "提示": f"设备厂家一致[{bbu_vendor}]"})
            else:
                # 未知型号，无法判断
                results.append({"站点": site_id, "规则": "RK-010", "结果": "通过", "风险": "建议", 
                               "提示": f"设备型号未知，无法判断厂家"})
        else:
            # 缺少设备信息
            results.append({"站点": site_id, "规则": "RK-010", "结果": "通过", "风险": "建议", 
                           "提示": "缺少设备信息，跳过检查"})

    # 汇总重复连接（仅在全量完成时执行）
    if not is_chunked or end_row >= len(df):
        all_sites = set(df["站点编号"].astype(str))
        conflict_sites = set()
        for pair, sites in connection_pairs.items():
            if len(sites) > 1:
                for s in sites:
                    conflict_sites.add(s)
                    results.append({"站点": s, "规则": "RK-009", "结果": "警告", "风险": "高风险",
                                    "提示": f"端口连接重复冲突: {pair} 被 {len(sites)} 个站点共用（{', '.join(sites)}）"})
        for s in all_sites - conflict_sites:
            results.append({"站点": s, "规则": "RK-009", "结果": "通过", "风险": "建议", "提示": "端口连接无重复冲突"})
    
    if is_chunked and end_row < len(df):
        # 分块模式未完成，返回中间状态
        return (False, results, connection_pairs, end_row)
    else:
        return pd.DataFrame(results)


# ==================== 内置生成函数（保持原逻辑） ====================
def generate_bom_data(sites_df):
    total_sites = len(sites_df)
    L = sum(safe_float_route(row.get("路由长度(m)", 0)) for _, row in sites_df.iterrows())
    J_val = sum(int(safe_float_route(row.get("室外接头数量", 0)) or 0) for _, row in sites_df.iterrows())
    K = sum(int(safe_float_route(row.get("光口配线对数", 0)) or 0) for _, row in sites_df.iterrows())
    N_d = sum(int(safe_float_route(row.get("接地设备数量", 0)) or 0) for _, row in sites_df.iterrows())

    bom_items = [
        {"编号": 1, "项目编码": "TX-GC-001", "专业类别": "通信线路工程", "设备/材料名称": "光缆", "规格型号": "GYTA-按芯数", "单位": "米", "数量": round(L * 1.06, 1), "项目特征": "单模铠装", "工程量计算规则": "路由长度x1.06", "工作内容": "敷设接续", "安装位置": "站点路由", "备注": "预留6%"},
        {"编号": 2, "项目编码": "TX-GC-002", "专业类别": "通信线路工程", "设备/材料名称": "防水套件", "规格型号": "热缩式防水胶带+胶泥套装", "单位": "套", "数量": J_val, "项目特征": "户外防水", "工程量计算规则": "室外接头数量", "工作内容": "接头防水处理", "安装位置": "室外接头处", "备注": ""},
        {"编号": 3, "项目编码": "TX-GC-003", "专业类别": "无线通信设备安装工程", "设备/材料名称": "光纤跳线", "规格型号": "LC-LC单模双芯", "单位": "条", "数量": math.ceil(K * 1.15), "项目特征": "双芯跳线", "工程量计算规则": "光口配线对数x1.15", "工作内容": "端口连接", "安装位置": "ODF-AAU/BBU", "备注": "含15%冗余"},
        {"编号": 4, "项目编码": "TX-GC-004", "专业类别": "通信线路工程", "设备/材料名称": "接地线", "规格型号": "BVR 16mm2", "单位": "条", "数量": N_d, "项目特征": "黄绿双色", "工程量计算规则": "接地设备数量", "工作内容": "接地连接", "安装位置": "设备接地排", "备注": ""},
    ]
    return pd.DataFrame(bom_items, columns=BOM_FIELDS)


def generate_bor_data(sites_df):
    total_sites = len(sites_df)
    bor_data = [
        {"资源类别": "工具", "资源名称": "光纤熔接机", "规格/型号": "单芯熔接", "数量": 1, "单位": "台", "备注": "光缆接续"},
        {"资源类别": "工具", "资源名称": "OTDR", "规格/型号": "光时域反射仪", "数量": 1, "单位": "台", "备注": "光缆测试"},
        {"资源类别": "人员", "资源名称": "光纤技工", "规格/型号": "持证", "数量": 2, "单位": "人", "备注": "光缆接续"},
        {"资源类别": "人员", "资源名称": "电工", "规格/型号": "低压电工证", "数量": 1, "单位": "人", "备注": "取电接入"},
        {"资源类别": "安全", "资源名称": "安全帽", "规格/型号": "-", "数量": total_sites * 2, "单位": "个", "备注": "全员"},
    ]
    return pd.DataFrame(bor_data)


def generate_fiber_data(sites_df):
    fiber_data = []
    color_map = {1: "蓝", 2: "橙", 3: "绿", 4: "棕", 5: "灰", 6: "白", 7: "红", 8: "黑", 9: "黄", 10: "紫", 11: "粉", 12: "青"}

    for _, row in sites_df.iterrows():
        sid = row["站点编号"]
        cores = int(safe_float_route(row.get("光口配线对数", 12)) or 0) // 2
        start = str(row.get("起点", "起点") or "起点")
        end = str(row.get("终点", "终点") or "终点")
        for c in range(1, min(max(cores, 4), 12) + 1):
            fiber_data.append({
                "工程编号": f"SZ-{sid}",
                "光缆编号": f"FC-{sid}-{c:02d}",
                "站点编号": sid,
                "纤芯序号": c,
                "纤芯颜色": color_map.get(c, '备用'),
                "起始端子": f"{start}-P{c}",
                "终止端子": f"{end}-P{c}",
                "BBU端口": f"BBU-P1-P{c}",
                "AAU端口": f"AAU-P1-P{c}",
                "业务类型": '数据' if c <= cores//2 else '备用'
            })
    return pd.DataFrame(fiber_data)


def generate_bop_content(sites_df):
    total_sites = len(sites_df)
    site_types = sites_df["站点类型"].value_counts().to_dict() if "站点类型" in sites_df.columns else {}
    type_summary = "、".join([f"{k}{v}个" for k, v in site_types.items()])
    power_modes = sites_df["取电方式"].value_counts().to_dict() if "取电方式" in sites_df.columns else {}
    power_summary = "、".join([f"{k}{v}个" for k, v in power_modes.items()])
    total_len = sum(safe_float_route(row.get("路由长度(m)", 0)) for _, row in sites_df.iterrows())

    content = f"""# 5G基站施工工艺指导书

## 一、站点勘测与规划

**勘测目标**：确认{total_sites}个站点的实际位置、环境条件、取电方式及线缆路由走向。

**站点分布**：{type_summary}，分布在{total_sites}个点位，总路由长度约{total_len:.0f}米。

**勘测内容**：
1. 确认站点经纬度坐标，核对设计图纸与实际场地的一致性
2. 检查现场供电条件，确认取电方式的可行性（{power_summary}）
3. 勘察线缆路由，测量实际距离，标记高风险路段
4. 记录现场障碍物、施工限制条件

**风险点记录**：对长距离线缆站点标记为高风险段，制定专项施工方案。

## 二、设备安装要求

**安装位置要求**：
- AAU安装于抱杆顶端，确保天线高度符合设计要求
- BBU安装于机柜内，保持通风良好，便于维护
- 设备安装应水平、牢固，水平误差<=3mm

**固定要求**：
- 宏站/杆塔站：AAU采用抱杆安装，紧固螺栓力矩>=40N-m
- 室内分布：设备壁挂或机架安装，确保承重安全
- 室外设备必须做好防水密封，防水等级>=IP65

**接地要求**：
- 所有设备必须可靠接地
- 接地电阻<=10欧（室外宏站<=4欧）
- 接地线采用黄绿双色BVR 16mm2线缆
- 防雷接地与工作接地分设，间距>=5m

## 三、线缆敷设要求

**光纤敷设规范**：
- 选用G.652D单模光缆，按芯数需求配置
- 弯曲半径>=20倍光缆外径（施工中>=30倍）
- 敷设张力不超过光缆短期允许张力的80%
- 光缆接续采用熔接方式，单芯熔接损耗<=0.05dB
- 光缆余长盘留整齐，标签清晰

**电源线敷设规范**：
- 电源线与其他线缆保持安全距离（>=150mm）
- 直流远供线路长度不超过150米（超过需增加中继）
- 电源接头做好防水、绝缘处理
- 线径根据传输距离和功率核算，确保末端压降<=5%

## 四、取电与接地要求

**取电方式**（本工程涉及：{power_summary}）：
- 市电直供：确认供电容量满足设备需求，安装自动切换装置
- 直流远供：检查远供设备输出电压、电流，配置隔离变压器和漏电保护
- 太阳能+储能：确认储能容量满足备用时间要求，定期检查电池状态

**安全规范**：
- 电工必须持低压电工证上岗
- 操作前必须确认断电，使用验电器验证
- 做好绝缘防护，使用绝缘工具
- 配电箱内空开、浪涌保护器、电表安装规范

**接地检查**：
- 施工前测量原有机房/铁塔接地电阻
- 接地体埋深>=0.6m（冻土层以下）
- 接地引入线截面积>=16mm2
- 所有接地点用防锈螺栓紧固

## 五、质量检查标准

| 检查项 | 验收指标 | 异常处理 |
|--------|----------|----------|
| 设备安装水平度 | <=1.5mm/m | 重新调整并复测 |
| 光缆熔接损耗 | <=0.05dB/点 | 重新熔接 |
| 接地电阻 | <=10欧（室外<=4欧） | 检查接地连接，增加接地体 |
| 电源电压 | 额定值+-10% | 检查供电设备，调整分接头 |
| 信号覆盖RSRP | >=-105dBm（边缘） | 调整天线方向角、下倾角 |
| 驻波比 | <=1.5 | 检查馈线接头、天线匹配 |
| 光纤接口清洁度 | 符合IEC61300-3-35 | 重新清洁并测试 |

## 六、施工注意事项

**安全风险**：
1. 高空作业必须持证上岗，佩戴安全带和安全帽
2. 带电作业必须两人以上，一人操作一人监护
3. 施工现场设置安全警示标志，禁止无关人员进入
4. 雷雨天气严禁室外作业

**现场协调**：
1. 与物业、业主提前沟通，取得施工许可
2. 施工时间避开高峰期，减少对周边影响
3. 做好现场清洁，施工结束后恢复原状
4. 与其他施工单位协调交叉作业，避免冲突

**施工记录**：
1. 填写施工日志，记录每日工作内容和人员
2. 关键工序拍照留存（设备安装、光缆接续、接地连接）
3. 整理竣工资料，包括测试报告、竣工图纸、隐蔽工程记录
4. 及时归档，确保可追溯
"""
    return content


def generate_risk_content(sites_df):
    total_sites = len(sites_df)
    outdoor_sites = 0
    long_route_sites = []
    for _, row in sites_df.iterrows():
        stype = str(row.get("站点类型", ""))
        if stype in ["室外宏站", "杆塔站", "塔房站"]:
            outdoor_sites += 1
        route_len = safe_float_route(row.get("路由长度(m)", 0))
        if route_len > 300:
            long_route_sites.append((str(row.get("站点编号", "")), route_len))

    long_route_text = ""
    if long_route_sites:
        long_route_text = "\n### 长距离站点\n\n"
        for sid, rlen in long_route_sites:
            long_route_text += f"- {sid}：{rlen:.0f}米，需专项施工方案\n"

    return f"""**依据YD/T 5264-2021《5G数字蜂窝移动通信网工程施工监理规范》分析**

本工程共{total_sites}个站点，其中室外站点{outdoor_sites}个。{long_route_text}

**施工风险提示：**

1. **线缆敷设风险（中）** — 长距离站点需注意牵引力控制，敷设张力不超过光缆短期允许张力的80%，避免光缆损伤。管沟或架空敷设时注意弯曲半径>=30倍光缆外径。

2. **取电安全风险（高）** — 市电直供站点需确认断电后再操作，防止触电事故。直流远供线路需配置隔离变压器和漏电保护，电压波动范围+-10%。

3. **高空作业风险（高）** — 室外宏站/杆塔站共{outdoor_sites}个，需登高作业。必须持证上岗，佩戴安全带、安全帽，雷雨天气严禁室外作业。

4. **光缆熔接质量风险（中）** — 熔接损耗超标需重新接续。单芯熔接损耗<=0.05dB，OTDR测试曲线无异常反射。光纤接口清洁度符合IEC61300-3-35标准。

5. **接地系统风险（高）** — 室外站点接地电阻<=4欧，室内站点<=10欧。接地电阻不达标影响设备安全，需定期检查。接地体埋深>=0.6m，引入线截面>=16mm2。

**合规建议：**
- 严格按照YD/T 5264-2021《5G数字蜂窝移动通信网工程施工监理规范》执行
- 施工人员必须持证上岗（电工证、登高证），做好安全培训和交底
- 关键工序必须拍照留存（设备安装、光缆接续、接地连接），确保可追溯
- 竣工后进行全面测试（RSRP>=-105dBm、SINR>=15dB、业务成功率100%）
- 填写施工日志，每日记录工作内容、人员和异常情况
"""


def style_excel_worksheet(ws, df, title=None):
    """为Excel工作表添加真实单元格样式，可选标题行"""
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    header_fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    data_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
    alt_fill = PatternFill(start_color="F4F6FA", end_color="F4F6FA", fill_type="solid")

    # 如果有标题，在第1行插入合并标题行，数据从第3行开始
    data_start_row = 1
    if title:
        ws.insert_rows(1)
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(df.columns))
        title_cell = ws.cell(row=1, column=1, value=title)
        title_cell.font = Font(bold=True, size=14, color="1F2937")
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        title_cell.fill = PatternFill(start_color="F0F4FF", end_color="F0F4FF", fill_type="solid")
        ws.row_dimensions[1].height = 30
        data_start_row = 2
        # 给标题行加边框
        for c in range(1, len(df.columns) + 1):
            ws.cell(row=1, column=c).border = thin_border

    for row_idx, row in enumerate(ws.iter_rows(min_row=data_start_row, max_row=ws.max_row, max_col=ws.max_column), data_start_row):
        for cell in row:
            cell.border = thin_border
            cell.alignment = Alignment(wrap_text=True, vertical='center', horizontal='center')
            if row_idx == data_start_row:
                cell.fill = header_fill
                cell.font = header_font
            else:
                cell.fill = alt_fill if (row_idx - data_start_row) % 2 == 1 else data_fill

    for col_idx, column in enumerate(df.columns, 1):
        col_letter = get_column_letter(col_idx)
        max_length = 0
        for cell in ws[col_letter]:
            try:
                if cell.value:
                    cell_str = str(cell.value)
                    length = sum(2 if ord(char) > 127 else 1 for char in cell_str)
                    max_length = max(max_length, length)
            except:
                pass
        adjusted_width = min(max(max_length + 2, 8), 50)
        ws.column_dimensions[col_letter].width = adjusted_width

    freeze_row = data_start_row + 1
    ws.freeze_panes = f"A{freeze_row}"


def create_word_document(content, title="5G基站施工工艺指导书"):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    try:
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        if style._element.rPr is not None and style._element.rPr.rFonts is not None:
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    except:
        pass

    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.name = '黑体'
        run.font.size = Pt(18)
        run.font.bold = True
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        except:
            pass

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith('# ') and not line.startswith('## '):
            heading = doc.add_heading(line[2:], level=1)
            for run in heading.runs:
                run.font.name = '黑体'
                run.font.size = Pt(16)
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                except:
                    pass
            i += 1
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            for run in heading.runs:
                run.font.name = '黑体'
                run.font.size = Pt(14)
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                except:
                    pass
            i += 1
        elif line.startswith('|') and line.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            data_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl[1:-1].split('|')]
                if all(c.replace('-','').replace(':','').replace(' ','') == '' for c in cells):
                    continue
                data_rows.append(cells)
            if data_rows:
                max_cols = max(len(r) for r in data_rows)
                table = doc.add_table(rows=len(data_rows), cols=max_cols, style='Table Grid')
                for r_idx, row in enumerate(data_rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < max_cols:
                            cell = table.cell(r_idx, c_idx)
                            cell.text = cell_text
                            if r_idx == 0:
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.bold = True
                                        run.font.name = '宋体'
                                        try:
                                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                        except:
                                            pass
                doc.add_paragraph()
        else:
            text = line.replace('**', '')
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = '宋体'
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                except:
                    pass
            i += 1
    return doc


def build_excel_bytes(current_df):
    """预构建 Excel 交付结果（一键下载优化）"""
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as w:
        sheet_titles = {"bom": "施工BOM清单", "bor": "资源需求清单", "fiber": "纤芯分配表", "review": "合规审查结果"}

        if st.session_state.review_results is not None:
            review_df = st.session_state.review_results.copy()
            review_df.to_excel(w, index=False, sheet_name="合规审查结果")
            ws = w.sheets["合规审查结果"]
            style_excel_worksheet(ws, review_df, title="合规审查结果")
            # 失败 → 红色背景；警告 → 橙色背景
            red_fill = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")
            orange_fill = PatternFill(start_color="FFE5CC", end_color="FFE5CC", fill_type="solid")
            for row in ws.iter_rows(min_row=3, max_row=ws.max_row):
                for cell in row:
                    if cell.value and "失败" in str(cell.value):
                        cell.fill = red_fill
                    elif cell.value and "警告" in str(cell.value):
                        cell.fill = orange_fill

        sheet_names = {"bom": "施工BOM", "bor": "资源需求清单", "fiber": "纤芯分配表"}
        if st.session_state.offline_mode and hasattr(st.session_state, 'ai_dataframes'):
            for key in ["bom", "bor", "fiber"]:
                if key in st.session_state.ai_dataframes:
                    df = st.session_state.ai_dataframes[key]
                    sheet_name = sheet_names.get(key, key)[:31]
                    df.to_excel(w, index=False, sheet_name=sheet_name)
                    ws = w.sheets[sheet_name]
                    style_excel_worksheet(ws, df, title=sheet_titles.get(key))
        else:
            for s in AI_STEPS:
                key = s["key"]
                if key in ("bop", "risk"):
                    continue
                content = st.session_state.ai_data.get(key, "")
                sheet_name = sheet_names.get(key, key)[:31]
                df_parsed = parse_markdown_table(content)
                if df_parsed is not None and not df_parsed.empty:
                    df_parsed.columns = [str(c).replace("*","").replace("_","").strip() for c in df_parsed.columns]
                    df_parsed.to_excel(w, index=False, sheet_name=sheet_name)
                    ws = w.sheets[sheet_name]
                    style_excel_worksheet(ws, df_parsed, title=sheet_titles.get(key))
                else:
                    pd.DataFrame({"内容": [content]}).to_excel(w, index=False, sheet_name=sheet_name)
    return output.getvalue()


def build_word_report_bytes(current_df):
    """预构建综合交付报告 Word — 仅含汇总统计，无逐条错误明细"""
    total_sites = len(current_df)

    # 基础工程信息
    report_content = f"""# 5G基站工程交付报告

## 一、项目基础信息

| 项目 | 数值 |
|------|------|
| 总审查站点数量 | {total_sites}个 |
"""
    if "站点类型" in current_df.columns:
        types_count = current_df["站点类型"].value_counts().to_dict()
        for t, c in types_count.items():
            report_content += f"| 站点类型-{t} | {c}个 |\n"
    if "网络制式" in current_df.columns:
        network_count = current_df["网络制式"].value_counts().to_dict()
        for n, c in network_count.items():
            report_content += f"| 网络制式-{n} | {c}个 |\n"
    if "路由长度(m)" in current_df.columns:
        total_len = current_df["路由长度(m)"].apply(safe_float_route).sum()
        report_content += f"| 总路由长度 | {total_len:.1f}米 |\n"
    report_content += "\n"

    # 合规审查汇总统计
    if st.session_state.review_results is not None:
        review_df = st.session_state.review_results
        total_rules = len(review_df)
        fail_count = len(review_df[review_df["结果"].str.contains("失败", na=False)])
        warn_count = len(review_df[review_df["结果"].str.contains("警告", na=False)])
        pass_count = len(review_df[review_df["结果"].str.contains("通过", na=False)])
        high_risk_count = len(review_df[review_df["风险"].str.contains("高风险", na=False)])
        mid_risk_count = len(review_df[review_df["风险"].str.contains("中风险", na=False)])
        low_risk_count = len(review_df[review_df["风险"].str.contains("建议", na=False)])
        pass_rate = round(pass_count / total_rules * 100, 1) if total_rules > 0 else 100.0

        # Top 5 高频不通过问题
        fail_df = review_df[review_df["结果"].str.contains("失败", na=False)]
        top5_fails = []
        if not fail_df.empty:
            rule_counts = fail_df["规则"].value_counts().head(5)
            for rule, cnt in rule_counts.items():
                # 获取规则对应的提示摘要
                sample_tip = fail_df[fail_df["规则"] == rule]["提示"].iloc[0][:60]
                top5_fails.append((rule, cnt, sample_tip))

        report_content += f"""## 二、合规审查汇总统计

| 统计项 | 数值 |
|------|------|
| 总审查规则项数量 | {total_rules}项 |
| 通过项总数 | {pass_count}项 |
| 不通过项总数 | {fail_count}项 |
| 警告项总数 | {warn_count}项 |
| 整体合规通过率 | {pass_rate}% |
| 高风险问题数量 | {high_risk_count}项 |
| 一般风险问题数量 | {mid_risk_count}项 |
| 轻微风险问题数量 | {low_risk_count}项 |

### 高频不通过问题 TOP5

| 排名 | 规则编号 | 出现次数 | 问题说明 |
|------|----------|----------|----------|
"""
        if top5_fails:
            for rank, (rule, cnt, tip) in enumerate(top5_fails, 1):
                report_content += f"| {rank} | {rule} | {cnt}次 | {tip} |\n"
        else:
            report_content += "| — | — | 0 | 无不通过项 |\n"

        report_content += f"""

---

> **全量不通过项明细、完整错误数据详见附件《合规审查不通过项.xlsx》**
"""

    # 施工工艺指导书
    report_content += f"""

## 三、施工工艺指导书

{st.session_state.ai_data.get('bop', '暂无数据')}
"""

    # 风险提示
    report_content += f"""

## 四、风险提示

{st.session_state.ai_data.get('risk', '暂无数据')}
"""

    doc = create_word_document(report_content, "5G基站工程交付报告")
    word_output = BytesIO()
    doc.save(word_output)
    return word_output.getvalue()


def build_word_bop_bytes():
    """预构建工艺指导书 Word"""
    bop_content = st.session_state.ai_data.get("bop", "暂无数据")
    doc = create_word_document(bop_content, "5G基站施工工艺指导书")
    word_output = BytesIO()
    doc.save(word_output)
    return word_output.getvalue()


def build_review_failures_excel_bytes():
    """生成合规审查不通过项 Excel — 全量逐条错误数据 + 条件格式"""
    if st.session_state.review_results is None:
        return None

    review_df = st.session_state.review_results.copy()
    # 提取所有失败+警告项作为不通过项
    fail_df = review_df[review_df["结果"].str.contains("失败|警告", na=False)].copy()

    # 计算汇总统计
    total_rules = len(review_df)
    fail_count = len(review_df[review_df["结果"].str.contains("失败", na=False)])
    warn_count = len(review_df[review_df["结果"].str.contains("警告", na=False)])
    pass_count = len(review_df[review_df["结果"].str.contains("通过", na=False)])
    pass_rate = round(pass_count / total_rules * 100, 1) if total_rules > 0 else 100.0

    # 风险等级重新归类（用于Excel展示）
    def remap_risk(row):
        result = str(row.get("结果", ""))
        rule = str(row.get("规则", ""))
        # 必填字段缺失相关的规则 → 高
        high_risk_rules = ["RK-001", "RK-004", "RK-005"]
        # 枚举值不符 → 一般
        enum_rules = ["RK-002", "RK-003"]
        if rule in high_risk_rules:
            return "高"
        elif rule in enum_rules:
            return "一般"
        elif rule in ["RK-006", "RK-008"]:
            return "高" if "失败" in result else "一般"
        elif rule in ["RK-007"]:
            return "一般" if "失败" in result else "轻微"
        elif rule in ["RK-009"]:
            return "高"
        elif rule in ["RK-010"]:
            return "轻微"
        return "轻微"

    def get_suggestion(rule):
        suggestions = {
            "RK-001": "修正站点编号格式，确保符合大写字母+数字+连字符，6-25位规范",
            "RK-002": "修正站点类型为枚举值之一，参考系统预设站点类型列表",
            "RK-003": "修正网络制式为枚举值之一，参考系统预设网络制式列表",
            "RK-004": "补充线缆类型字段，路由长度>0时线缆类型为必填项",
            "RK-005": "补充接地设备数量，室外站点必须填写且>=1",
            "RK-006": "增加光口配线对数至>=12，满足5G站点最低配线要求",
            "RK-007": "修正起止点，确保每个站点的起点与终点不重复",
            "RK-008": "直流远供线路长度超过150米，需增加中继设备或改用其他取电方式",
            "RK-009": "检查端口连接规划，消除重复连接冲突",
            "RK-010": "统一BBU与AAU设备厂家，降低兼容性风险",
        }
        return suggestions.get(rule, "请核实数据后修正")

    def get_department(rule):
        dept_map = {
            "RK-001": "数据管理组",
            "RK-002": "工程设计组",
            "RK-003": "工程设计组",
            "RK-004": "工程设计组",
            "RK-005": "施工管理组",
            "RK-006": "工程设计组",
            "RK-007": "工程设计组",
            "RK-008": "施工管理组",
            "RK-009": "工程设计组",
            "RK-010": "设备采购组",
        }
        return dept_map.get(rule, "项目管理组")

    def get_deadline(risk_level):
        dl_map = {"高": "3个工作日", "一般": "5个工作日", "轻微": "7个工作日"}
        return dl_map.get(risk_level, "7个工作日")

    # 始终定义 output_cols，避免空不通过项时 UnboundLocalError
    output_cols = ["站点编号", "站点名称", "问题规则", "不通过原因", "原值", "期望值/规则",
                   "风险等级", "整改建议", "责任部门", "整改期限"]

    if not fail_df.empty:
        # 构建输出DataFrame
        output_data = []
        for _, row in fail_df.iterrows():
            risk_level = remap_risk(row)
            rule = str(row.get("规则", ""))
            output_data.append({
                "站点编号": row.get("站点", ""),
                "站点名称": row.get("站点名称", ""),
                "问题规则": rule,
                "不通过原因": row.get("提示", ""),
                "原值": row.get("原值", ""),
                "期望值/规则": row.get("期望值/规则", ""),
                "风险等级": risk_level,
                "整改建议": get_suggestion(rule),
                "责任部门": get_department(rule),
                "整改期限": get_deadline(risk_level),
            })
        output_df = pd.DataFrame(output_data, columns=output_cols)
    else:
        # 全部通过：生成占位信息
        output_df = pd.DataFrame([{
            "站点编号": "—", "站点名称": "全部通过，无不通过项",
            "问题规则": "—", "不通过原因": "—", "原值": "—", "期望值/规则": "—",
            "风险等级": "—", "整改建议": "—", "责任部门": "—", "整改期限": "—"
        }])

    # 统计计数（用于汇总行）
    if not fail_df.empty:
        high_count = len([d for d in output_data if d["风险等级"] == "高"])
        mid_count = len([d for d in output_data if d["风险等级"] == "一般"])
        low_count = len([d for d in output_data if d["风险等级"] == "轻微"])
    else:
        high_count = mid_count = low_count = 0

    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as w:
        output_df.to_excel(w, index=False, sheet_name="合规审查不通过项", startrow=2)
        ws = w.sheets["合规审查不通过项"]

        # Row 1: 汇总统计行（合并单元格）
        summary_text = f"总不通过项数: {len(output_df)} | 高风险: {high_count} | 一般风险: {mid_count} | 轻微风险: {low_count} | 整体通过率: {pass_rate}%"
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(output_cols))
        summary_cell = ws.cell(row=1, column=1, value=summary_text)
        summary_cell.font = Font(bold=True, size=12, color="1F2937")
        summary_cell.fill = PatternFill(start_color="F0F4FF", end_color="F0F4FF", fill_type="solid")
        summary_cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 28

        # Row 2: 表头样式
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        header_fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)

        for col_idx in range(1, len(output_cols) + 1):
            cell = ws.cell(row=2, column=col_idx)
            cell.border = thin_border
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")

        # Row 3+: 数据行 + 条件格式
        red_fill = PatternFill(start_color="FFB3B3", end_color="FFB3B3", fill_type="solid")
        yellow_fill = PatternFill(start_color="FFF3B3", end_color="FFF3B3", fill_type="solid")
        gray_fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
        data_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
        alt_fill = PatternFill(start_color="F4F6FA", end_color="F4F6FA", fill_type="solid")

        risk_col_idx = output_cols.index("风险等级") + 1  # 1-based

        for row_idx in range(3, ws.max_row + 1):
            risk_val = ws.cell(row=row_idx, column=risk_col_idx).value
            if risk_val == "高":
                row_fill = red_fill
            elif risk_val == "一般":
                row_fill = yellow_fill
            elif risk_val == "轻微":
                row_fill = gray_fill
            else:
                row_fill = alt_fill if (row_idx - 3) % 2 == 1 else data_fill

            for col_idx in range(1, len(output_cols) + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.border = thin_border
                cell.alignment = Alignment(wrap_text=True, vertical="center")
                cell.fill = row_fill

        # 列宽自适应
        for col_idx, col_name in enumerate(output_cols, 1):
            col_letter = get_column_letter(col_idx)
            max_len = len(col_name) * 2  # 中文字符约2个单位
            for row_idx in range(3, min(ws.max_row + 1, 50)):
                val = ws.cell(row=row_idx, column=col_idx).value
                if val:
                    max_len = max(max_len, sum(2 if ord(c) > 127 else 1 for c in str(val)))
            ws.column_dimensions[col_letter].width = min(max(max_len + 4, 10), 50)

        # 冻结窗格 + 自动筛选
        ws.freeze_panes = "A3"
        ws.auto_filter.ref = f"A2:{get_column_letter(len(output_cols))}{ws.max_row}"

        # 给汇总行加边框
        for col_idx in range(1, len(output_cols) + 1):
            ws.cell(row=1, column=col_idx).border = thin_border

    return output.getvalue()


def _compact_site_range(site_list, max_show=4):
    """将站点列表压缩为区间表示，如 A01~D03 (共15个)"""
    if not site_list:
        return "—"
    sites = sorted(set(str(s) for s in site_list))
    n = len(sites)
    if n == 1:
        return sites[0]
    if n <= max_show:
        return "、".join(sites) + f" (共{n}个)"
    # 提取首尾站点作为区间
    first, last = sites[0], sites[-1]
    return f"{first}~{last} (共{n}个)"


def _remap_risk(result_str, rule_str):
    """风险等级映射"""
    if rule_str in ("RK-001", "RK-004", "RK-005", "RK-009"):
        return "高"
    elif rule_str in ("RK-002", "RK-003"):
        return "一般"
    elif rule_str in ("RK-006", "RK-008"):
        return "高" if "失败" in result_str else "一般"
    elif rule_str == "RK-007":
        return "一般" if "失败" in result_str else "轻微"
    elif rule_str == "RK-010":
        return "轻微"
    return "轻微"


def _get_suggestion(rule):
    suggestions = {
        "RK-001": "修正站点编号格式，确保符合大写字母+数字+连字符，6-25位规范",
        "RK-002": "修正站点类型为枚举值之一，参考系统预设站点类型列表",
        "RK-003": "修正网络制式为枚举值之一，参考系统预设网络制式列表",
        "RK-004": "补充线缆类型字段，路由长度>0时线缆类型为必填项",
        "RK-005": "补充接地设备数量，室外站点必须填写且>=1",
        "RK-006": "增加光口配线对数至>=12，满足5G站点最低配线要求",
        "RK-007": "修正起止点，确保每个站点的起点与终点不重复",
        "RK-008": "线路长度超过150米，需增加中继设备或改用其他取电方式",
        "RK-009": "调整纤芯分配，避免多站点共用同一端口",
        "RK-010": "统一BBU与AAU设备厂家，降低兼容性风险",
    }
    return suggestions.get(rule, "请核实数据后修正")


def _get_department(rule):
    dept_map = {
        "RK-001": "数据管理组",
        "RK-002": "工程设计组",
        "RK-003": "工程设计组",
        "RK-004": "工程设计组",
        "RK-005": "施工管理组",
        "RK-006": "工程设计组",
        "RK-007": "工程设计组",
        "RK-008": "施工管理组",
        "RK-009": "工程设计组",
        "RK-010": "设备采购组",
    }
    return dept_map.get(rule, "项目管理组")


def _get_deadline(risk_level):
    return {"高": "3个工作日", "一般": "5个工作日", "轻微": "7个工作日"}.get(risk_level, "7个工作日")


def build_compliance_review_full_excel_bytes(current_df):
    """生成合规审查完整报告 Excel — 2个工作表：总览表 + 问题整改台账
    返回 (bytes, filename) 或 None"""
    if st.session_state.review_results is None:
        return None

    review_df = st.session_state.review_results.copy()
    fail_df = review_df[review_df["结果"].str.contains("失败", na=False)].copy()
    warn_df = review_df[review_df["结果"].str.contains("警告", na=False)].copy()
    pass_df = review_df[review_df["结果"].str.contains("通过", na=False)].copy()

    total_rules = len(review_df)
    fail_count = len(fail_df)
    warn_count = len(warn_df)
    pass_count = len(pass_df)
    pass_rate = round(pass_count / total_rules * 100, 1) if total_rules > 0 else 100.0

    # 确定批次名
    batch_name = "站点批次"
    if current_df is not None and "站点编号" in current_df.columns:
        sites = current_df["站点编号"].dropna().unique()
        if len(sites) == 1:
            batch_name = str(sites[0])
        elif len(sites) > 1:
            batch_name = str(sites[0])[:8] if len(str(sites[0])) >= 8 else str(sites[0])

    from datetime import datetime
    date_str = datetime.now().strftime("%Y%m%d")
    filename = f"{batch_name}_合规审查报告_{date_str}.xlsx"

    # 合并失败+警告项，构建整改台账
    issue_df = pd.concat([fail_df, warn_df], ignore_index=True) if not fail_df.empty or not warn_df.empty else pd.DataFrame()
    if not issue_df.empty:
        # 按规则去重合并：同一规则 → 1条核心问题
        merged_rows = []
        risk_order = {"高": 0, "一般": 1, "轻微": 2}
        for rule, grp in issue_df.groupby("规则"):
            result_samples = grp["结果"].unique()
            worst_result = "失败" if "失败" in str(result_samples) else "警告"
            risk_levels = [_remap_risk(str(r["结果"]), str(r["规则"])) for _, r in grp.iterrows()]
            worst_risk = min(risk_levels, key=lambda x: risk_order.get(x, 99))
            site_list = grp["站点"].astype(str).tolist()
            site_range = _compact_site_range(site_list)
            merged_rows.append({
                "序号": len(merged_rows) + 1,
                "规则编号": rule,
                "核心问题": grp["提示"].iloc[0] if "提示" in grp.columns else rule,
                "受影响站点": site_range,
                "影响站点数": len(site_list),
                "风险等级": worst_risk,
                "整改建议": _get_suggestion(rule),
                "责任部门": _get_department(rule),
                "整改期限": _get_deadline(worst_risk),
            })
        ledger_df = pd.DataFrame(merged_rows)
        ledger_df = ledger_df.sort_values("风险等级", key=lambda s: s.map(risk_order))
        ledger_df["序号"] = range(1, len(ledger_df) + 1)
    else:
        ledger_df = pd.DataFrame()

    # ======================
    # 风险等级分布统计
    risk_counts = {"高": 0, "一般": 0, "轻微": 0}
    for _, row in issue_df.iterrows() if not issue_df.empty else []:
        rl = _remap_risk(str(row["结果"]), str(row["规则"]))
        risk_counts[rl] = risk_counts.get(rl, 0) + 1

    # 风险等级文本
    risk_dist_text = " · ".join(f"{k}风险 {v} 条" for k, v in risk_counts.items() if v > 0)
    if not risk_dist_text:
        risk_dist_text = "全部通过，无风险项"

    # ======================
    # 样式
    thin_border = Border(
        left=Side(style='thin', color='000000'),
        right=Side(style='thin', color='000000'),
        top=Side(style='thin', color='000000'),
        bottom=Side(style='thin', color='000000')
    )
    header_fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    red_fill = PatternFill(start_color="FFB3B3", end_color="FFB3B3", fill_type="solid")
    yellow_fill = PatternFill(start_color="FFF3B3", end_color="FFF3B3", fill_type="solid")
    gray_fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")

    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as w:
        # ================================================================
        # Sheet 1: 合规审查总览表
        # ================================================================
        ws1 = w.book.create_sheet("合规审查总览表", 0)

        # --- 表头说明 (Row 1-2) ---
        ws1.merge_cells(start_row=1, start_column=1, end_row=1, end_column=6)
        guide1 = ws1.cell(row=1, column=1, value="合规审查总览表 — 本次审查的整体评估与核心问题概览")
        guide1.font = Font(bold=True, size=14, color="1F2937")
        guide1.fill = PatternFill(start_color="F0F4FF", end_color="F0F4FF", fill_type="solid")
        guide1.alignment = Alignment(horizontal="left", vertical="center")
        ws1.row_dimensions[1].height = 36

        ws1.merge_cells(start_row=2, start_column=1, end_row=2, end_column=6)
        guide2 = ws1.cell(row=2, column=1, value=f"使用指引：本表展示审查结论与核心问题清单。具体整改事项详见「问题整改台账」。审查时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        guide2.font = Font(size=9, color="6B7280", italic=True)
        guide2.alignment = Alignment(horizontal="left", vertical="center")
        ws1.row_dimensions[2].height = 24

        # --- 审查结论区 (Row 4-9) ---
        start_r = 4
        ws1.merge_cells(start_row=start_r, start_column=1, end_row=start_r, end_column=6)
        section_title = ws1.cell(row=start_r, column=1, value="审查结论")
        section_title.font = Font(bold=True, size=12, color="4F46E5")
        section_title.fill = PatternFill(start_color="EEF2FF", end_color="EEF2FF", fill_type="solid")
        section_title.alignment = Alignment(horizontal="left", vertical="center")
        ws1.cell(row=start_r, column=1).border = thin_border
        for ci in range(1, 7):
            ws1.cell(row=start_r, column=ci).fill = PatternFill(start_color="EEF2FF", end_color="EEF2FF", fill_type="solid")
        for ci in range(2, 7):
            ws1.cell(row=start_r, column=ci).border = Border()
        ws1.row_dimensions[start_r].height = 28

        overview_fields = [
            ("审查规则总数", f"{total_rules} 条"),
            ("合规通过", f"{pass_count} 条 · 通过率 {pass_rate}%"),
            ("不通过项", f"{fail_count} 条"),
            ("警告项", f"{warn_count} 条"),
            ("风险分布", risk_dist_text if fail_count + warn_count > 0 else "无风险项"),
        ]
        for i, (label, val) in enumerate(overview_fields):
            r = start_r + 1 + i
            c_lbl = ws1.cell(row=r, column=1, value=label)
            c_lbl.font = Font(bold=True, size=10, color="374151")
            c_lbl.fill = PatternFill(start_color="F9FAFB", end_color="F9FAFB", fill_type="solid")
            c_lbl.border = thin_border
            c_lbl.alignment = Alignment(vertical="center")

            ws1.merge_cells(start_row=r, start_column=2, end_row=r, end_column=6)
            c_val = ws1.cell(row=r, column=2, value=val)
            c_val.font = Font(bold=True, size=10, color="DC2626" if ("不通过项" in label and fail_count > 0) else "1F2937")
            c_val.alignment = Alignment(vertical="center")
            # 合并区域左上角单元格设完整外边框，内部单元格清除边框
            ws1.cell(row=r, column=2).border = thin_border
            for ci_b in range(3, 7):
                ws1.cell(row=r, column=ci_b).border = Border()
            ws1.row_dimensions[r].height = 24

        # --- 核心问题区 (Row 12+) ---
        issues_start = start_r + len(overview_fields) + 3
        ws1.merge_cells(start_row=issues_start, start_column=1, end_row=issues_start, end_column=6)
        sec_title2 = ws1.cell(row=issues_start, column=1, value="核心问题清单" if not ledger_df.empty else "核心问题清单 — 全部通过")
        sec_title2.font = Font(bold=True, size=12, color="4F46E5")
        sec_title2.fill = PatternFill(start_color="EEF2FF", end_color="EEF2FF", fill_type="solid")
        sec_title2.alignment = Alignment(horizontal="left", vertical="center")
        ws1.cell(row=issues_start, column=1).border = thin_border
        for ci in range(1, 7):
            ws1.cell(row=issues_start, column=ci).fill = PatternFill(start_color="EEF2FF", end_color="EEF2FF", fill_type="solid")
        for ci in range(2, 7):
            ws1.cell(row=issues_start, column=ci).border = Border()
        ws1.row_dimensions[issues_start].height = 28

        if ledger_df.empty:
            # 空状态友好提示
            ws1.merge_cells(start_row=issues_start + 1, start_column=1, end_row=issues_start + 1, end_column=6)
            empty_cell = ws1.cell(row=issues_start + 1, column=1,
                                  value="无高风险不通过项，所有站点均已通过合规审查")
            empty_cell.font = Font(size=10, color="059669")
            empty_cell.fill = PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid")
            empty_cell.alignment = Alignment(horizontal="center", vertical="center")
            empty_cell.border = thin_border
            ws1.row_dimensions[issues_start + 1].height = 36
        else:
            # 核心问题表头
            issue_cols = ["序号", "规则编号", "核心问题", "受影响站点", "风险等级", "整改建议"]
            for ci, col_name in enumerate(issue_cols, 1):
                cell = ws1.cell(row=issues_start + 1, column=ci, value=col_name)
                cell.fill = PatternFill(start_color="6B7280", end_color="6B7280", fill_type="solid")
                cell.font = Font(bold=True, color="FFFFFF", size=10)
                cell.border = thin_border
                cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")
            ws1.row_dimensions[issues_start + 1].height = 26

            for di, (_, row) in enumerate(ledger_df.iterrows()):
                r = issues_start + 2 + di
                risk = row["风险等级"]
                row_fill = red_fill if risk == "高" else (yellow_fill if risk == "一般" else None)
                values = [row.get(c, "") for c in ["序号", "规则编号", "核心问题", "受影响站点", "风险等级", "整改建议"]]
                for ci, val in enumerate(values, 1):
                    cell = ws1.cell(row=r, column=ci, value=val)
                    cell.border = thin_border
                    cell.alignment = Alignment(wrap_text=True, vertical="center")
                    if row_fill:
                        cell.fill = row_fill
                    if ci == 5 and risk == "高":
                        cell.font = Font(bold=True, color="DC2626")
                ws1.row_dimensions[r].height = 28

        # 列宽
        col_widths_s1 = {1: 8, 2: 14, 3: 42, 4: 28, 5: 12, 6: 46}
        for ci, col_w in col_widths_s1.items():
            ws1.column_dimensions[get_column_letter(ci)].width = col_w

        # ================================================================
        # Sheet 2: 问题整改台账
        # ================================================================
        ws2 = w.book.create_sheet("问题整改台账", 1)

        # 表头说明
        ws2.merge_cells(start_row=1, start_column=1, end_row=1, end_column=9)
        g2a = ws2.cell(row=1, column=1,
                        value="问题整改台账 — 仅含不通过项与警告项，按风险等级降序，用于整改追踪")
        g2a.font = Font(bold=True, size=14, color="1F2937")
        g2a.fill = PatternFill(start_color="F0F4FF", end_color="F0F4FF", fill_type="solid")
        g2a.alignment = Alignment(horizontal="left", vertical="center")
        ws2.row_dimensions[1].height = 36

        ws2.merge_cells(start_row=2, start_column=1, end_row=2, end_column=9)
        g2b = ws2.cell(row=2, column=1,
                        value=f"使用指引：同规则下重复项已去重合并为一条核心问题。受影响的站点使用区间+数量表示。高风险项红色高亮，一般风险黄色高亮。{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        g2b.font = Font(size=9, color="6B7280", italic=True)
        g2b.alignment = Alignment(horizontal="left", vertical="center")
        ws2.row_dimensions[2].height = 24

        if ledger_df.empty:
            ws2.merge_cells(start_row=4, start_column=1, end_row=4, end_column=9)
            ec = ws2.cell(row=4, column=1,
                          value="全部站点均已通过合规审查，无整改事项")
            ec.font = Font(size=11, color="059669", bold=True)
            ec.fill = PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid")
            ec.alignment = Alignment(horizontal="center", vertical="center")
            ec.border = thin_border
            ws2.row_dimensions[4].height = 36
        else:
            # 台账表头 (Row 4)
            ledger_cols = ["序号", "规则编号", "核心问题", "受影响站点", "影响站点数",
                           "风险等级", "整改建议", "责任部门", "整改期限"]
            for ci, col_name in enumerate(ledger_cols, 1):
                cell = ws2.cell(row=4, column=ci, value=col_name)
                cell.fill = header_fill
                cell.font = header_font
                cell.border = thin_border
                cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")
            ws2.row_dimensions[4].height = 28

            for di, (_, row) in enumerate(ledger_df.iterrows()):
                r = 5 + di
                risk = row["风险等级"]
                if risk == "高":
                    row_fill = red_fill
                elif risk == "一般":
                    row_fill = yellow_fill
                else:
                    row_fill = None

                values = [
                    row["序号"], row["规则编号"], row["核心问题"], row["受影响站点"],
                    row["影响站点数"], row["风险等级"], row["整改建议"], row["责任部门"], row["整改期限"]
                ]
                for ci, val in enumerate(values, 1):
                    cell = ws2.cell(row=r, column=ci, value=val)
                    cell.border = thin_border
                    cell.alignment = Alignment(wrap_text=True, vertical="center")
                    if row_fill:
                        cell.fill = row_fill
                    if ci == 6 and risk == "高":
                        cell.font = Font(bold=True, color="DC2626")
                ws2.row_dimensions[r].height = 28

            # 冻结 + 筛选
            ws2.freeze_panes = "A5"
            ws2.auto_filter.ref = f"A4:{get_column_letter(len(ledger_cols))}{ws2.max_row}"

        # 列宽
        col_widths_s2 = {1: 8, 2: 14, 3: 42, 4: 28, 5: 14, 6: 12, 7: 46, 8: 16, 9: 14}
        for ci, col_w in col_widths_s2.items():
            ws2.column_dimensions[get_column_letter(ci)].width = col_w

    return output.getvalue(), filename


# ==================== 侧边栏内容 ====================
with st.sidebar:
    # 顶部品牌区
    st.markdown("""
    <div style="display:flex;align-items:center;gap:12px;padding:8px 0 14px 0;border-bottom:1px solid #e5e7eb;margin-bottom:6px;">
        <div style="background:linear-gradient(135deg,#4f46e5,#7c3aed);width:40px;height:40px;border-radius:10px;display:flex;align-items:center;justify-content:center;font-size:20px;box-shadow:0 2px 8px rgba(79,70,229,0.3);">🏗️</div>
        <div>
        <div style="font-size:1.1rem;font-weight:700;color:#111827;line-height:1.3;">数智化交付</div>
        <div style="font-size:0.7rem;color:#6b7280;">控制台 | demo0.2</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ========== 第一部分：数据管理（简化） ==========
    st.markdown("### 数据管理")

    # 上传区域 — 增强错误提示
    with st.container():
        def on_upload():
            if st.session_state.file_uploader:
                st.session_state.pending_upload = st.session_state.file_uploader

        st.file_uploader(
            "上传设计元数据表 / Upload Metadata Table (CSV/XLSX/XLS)",
            type=["csv", "xlsx", "xls"],
            key="file_uploader",
            on_change=on_upload
        )

    # v3 关键改进：文件上传后自动验证 + 自动切换数据源
    if st.session_state.pending_upload:
        f = st.session_state.pending_upload
        names = [x["name"] for x in st.session_state.uploaded_files]
        new_name = f.name

        if new_name not in names:
            df = read_file(f)
            if df is None:
                st.error(f"无法读取文件: {new_name}")
            else:
                # v3: 上传时立即验证
                is_valid, errors, warnings, missing_new = validate_uploaded_file(df, new_name)
                if errors:
                    st.error(f"**文件验证失败** — {new_name}")
                    for err in errors:
                        st.error(f"  {err}")
                    st.warning("文件已暂存但标记为无效，请修正后重新上传，或先用内置样本体验。")
                    # 保存但标记为无效
                    st.session_state.uploaded_files.append({
                        "name": new_name,
                        "df": df,
                        "valid": False,
                        "errors": errors,
                        "warnings": warnings
                    })
                    st.session_state.current_idx = len(st.session_state.uploaded_files) - 1
                else:
                    # 验证通过 — 自动切换数据源
                    st.session_state.uploaded_files.append({
                        "name": new_name,
                        "df": df,
                        "valid": True,
                        "errors": [],
                        "warnings": warnings
                    })
                    st.session_state.current_idx = len(st.session_state.uploaded_files) - 1
                    st.session_state.ai_generation_done = False
                    st.session_state.ai_data = {}
                    st.session_state.ai_dataframes = {}
                    st.session_state.result_df = None
                    st.session_state.offline_mode = False
                    st.session_state.excel_bytes = None
                    st.session_state.word_report_bytes = None
                    st.session_state.word_bop_bytes = None
                    st.session_state.review_failures_bytes = None

                    # 自动触发合规审查
                    st.session_state.review_results = run_compliance_review(df)

                    if warnings:
                        for w in warnings:
                            st.warning(w)
                    st.success(f"文件验证通过 — 已自动切换为当前数据源，合规审查已完成")
        else:
            # 文件已存在，自动切换并清空旧生成状态
            st.session_state.current_idx = names.index(new_name)
            st.session_state.ai_generation_done = False
            st.session_state.ai_data = {}
            st.session_state.ai_dataframes = {}
            st.session_state.result_df = None
            st.session_state.offline_mode = False
            st.session_state.excel_bytes = None
            st.session_state.word_report_bytes = None
            st.session_state.word_bop_bytes = None
            st.session_state.review_failures_bytes = None
            # 自动触发合规审查
            st.session_state.review_results = run_compliance_review(
                st.session_state.uploaded_files[st.session_state.current_idx]["df"]
            )

        st.session_state.pending_upload = None
        save_uploaded_files()
        st.session_state.expand_raw_data = True
        st.rerun()

    # v3: 数据源切换 — 一键切换，瞬时生效
    file_list = [x["name"] for x in st.session_state.uploaded_files]
    if st.session_state.current_idx >= len(file_list):
        st.session_state.current_idx = 0

    current_item = st.session_state.uploaded_files[st.session_state.current_idx]
    current_df = current_item["df"]

    # v3: 数据源显示 — 当前文件高亮卡片
    is_valid = current_item.get("valid", True)
    file_errors = current_item.get("errors", [])
    file_warnings = current_item.get("warnings", [])

    status_color = "#059669" if is_valid else "#dc2626"
    highlight_color = "#4f46e5" if is_valid else "#dc2626"
    status_text = "有效" if is_valid else "有错误"

    # 只在数据有效时才显示当前数据源卡片
    if is_valid:
        st.markdown(f"""
        <div style="margin:12px 0 4px 0;">
            <span style="font-size:0.65rem;font-weight:700;color:{highlight_color};text-transform:uppercase;letter-spacing:0.5px;">◆ 当前数据源</span>
        </div>
        <div style="background:#fff;border:1.5px solid {highlight_color};border-left:4px solid {highlight_color};border-radius:8px;padding:10px 12px;box-shadow:0 1px 6px {highlight_color}22;">
            <div style="display:flex;align-items:center;justify-content:space-between;">
                <div style="font-size:0.82rem;font-weight:600;color:#111827;flex:1;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;" title="{current_item['name']}">{current_item["name"]}</div>
                <span style="font-size:0.66rem;font-weight:700;color:#fff;background:{status_color};padding:2px 10px;border-radius:10px;flex-shrink:0;margin-left:8px;">{status_text}</span>
            </div>
            <div style="font-size:0.68rem;color:#6b7280;margin-top:5px;">{len(current_df)} 站点 · {len(current_df.columns)} 字段</div>
        </div>
        """, unsafe_allow_html=True)



    # 数据源切换选择器 + 快捷操作
    if len(file_list) > 1:
        st.markdown(f"""
        <div style="background:#f0f6ff;border:1px solid #bfdbfe;border-radius:8px;padding:6px 10px;margin-bottom:6px;">
        """, unsafe_allow_html=True)
        sel = st.selectbox(
            "切换数据源",
            file_list,
            index=st.session_state.current_idx,
            label_visibility="collapsed",
            key="sb_file_sel"
        )
        st.markdown("</div>", unsafe_allow_html=True)
        if sel in file_list and file_list.index(sel) != st.session_state.current_idx:
            st.session_state.current_idx = file_list.index(sel)
            st.session_state.ai_generation_done = False
            st.session_state.ai_data = {}
            st.session_state.ai_dataframes = {}
            st.session_state.result_df = None
            st.session_state.offline_mode = False
            st.session_state.excel_bytes = None
            st.session_state.word_report_bytes = None
            st.session_state.word_bop_bytes = None
            st.session_state.review_failures_bytes = None
            st.session_state.ai_step_times = {}
            st.session_state.ai_step_index = 0
            # 切换数据源后自动触发合规审查
            st.session_state.review_results = run_compliance_review(
                st.session_state.uploaded_files[st.session_state.current_idx]["df"]
            )
            st.rerun()

    # v3: 多选批量删除
    non_sample_files = [f["name"] for f in st.session_state.uploaded_files if "样例数据" not in f.get("name", "")]
    if non_sample_files:
        selected_to_delete = st.multiselect(
            "选择要删除的文件",
            non_sample_files,
            default=[],
            key="sb_multi_del",
            placeholder="点击选择要批量删除的文件..."
        )
        if st.button("批量删除选中文件", use_container_width=True, key="sb_batch_del"):
            if not selected_to_delete:
                st.info("请先勾选要删除的文件")
            else:
                st.session_state.uploaded_files = [
                    f for f in st.session_state.uploaded_files
                    if f["name"] not in selected_to_delete
                ]
                st.session_state.current_idx = 0
                st.session_state.ai_generation_done = False
                st.session_state.ai_data = {}
                st.session_state.ai_dataframes = {}
                st.session_state.review_results = None
                st.session_state.result_df = None
                st.session_state.offline_mode = False
                st.session_state.excel_bytes = None
                st.session_state.word_report_bytes = None
                st.session_state.word_bop_bytes = None
                st.session_state.review_failures_bytes = None
                save_uploaded_files()
                st.rerun()

    is_sample = "样例数据" in current_item.get("name", "")
    if is_sample:
        col_op2, col_op3 = st.columns(2)
        with col_op2:
            if st.button("详情", use_container_width=True, key="sb_sample_detail"):
                st.session_state.show_file_detail = True
                st.rerun()
        with col_op3:
            if st.button("重分析", use_container_width=True, key="sb_sample_reanalyze"):
                st.session_state.ai_generation_done = False
                st.session_state.ai_data = {}
                st.session_state.ai_dataframes = {}
                st.session_state.result_df = None
                st.session_state.offline_mode = False
                st.session_state.excel_bytes = None
                st.session_state.word_report_bytes = None
                st.session_state.word_bop_bytes = None
                st.session_state.review_failures_bytes = None
                st.session_state.review_failed = False
                # 重分析时重新触发合规审查
                st.session_state.review_results = run_compliance_review(
                    st.session_state.uploaded_files[st.session_state.current_idx]["df"]
                )
                st.rerun()
    else:
        col_op1, col_op2, col_op3 = st.columns(3)
        with col_op1:
            if st.button("删除", use_container_width=True, key="sb_del"):
                del st.session_state.uploaded_files[file_list.index(current_item["name"])]
                st.session_state.current_idx = 0
                st.session_state.ai_generation_done = False
                st.session_state.ai_data = {}
                st.session_state.ai_dataframes = {}
                st.session_state.review_results = None
                st.session_state.result_df = None
                st.session_state.offline_mode = False
                st.session_state.excel_bytes = None
                st.session_state.word_report_bytes = None
                st.session_state.word_bop_bytes = None
                st.session_state.review_failures_bytes = None
                save_uploaded_files()
                st.rerun()
        with col_op2:
            if st.button("详情", use_container_width=True, key="sb_detail"):
                st.session_state.show_file_detail = True
                st.rerun()
        with col_op3:
            if st.button("重分析", use_container_width=True, key="sb_reanalyze"):
                st.session_state.ai_generation_done = False
                st.session_state.ai_data = {}
                st.session_state.ai_dataframes = {}
                st.session_state.result_df = None
                st.session_state.offline_mode = False
                st.session_state.excel_bytes = None
                st.session_state.word_report_bytes = None
                st.session_state.word_bop_bytes = None
                st.session_state.review_failures_bytes = None
                st.session_state.review_failed = False
                # 重分析时重新触发合规审查
                st.session_state.review_results = run_compliance_review(
                    st.session_state.uploaded_files[st.session_state.current_idx]["df"]
                )
                st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("---")

    # ========== 第二部分：AI配置（v3 极简设计） ==========
    st.markdown("### AI 引擎")

    # v3: 默认内置模式，简洁的切换
    ai_mode = st.radio(
        "模式",
        ["标准AI引擎", "外部AI引擎（推荐）"],
        index=0 if st.session_state.use_builtin_ai else 1,
        horizontal=False,
        key="ai_engine_mode",
        help="标准：开箱即用 | 外部：高质量，需配置"
    )
    st.session_state.use_builtin_ai = (ai_mode == "标准AI引擎")

    if st.session_state.use_builtin_ai:
        # v3: 内置模式 — 简洁提示
        st.markdown("""
        <div style="background:#dcfce7;border:1px solid #a7f3d0;border-radius:8px;padding:10px 14px;margin:8px 0;font-size:0.78rem;color:#166534;">
            <b>标准AI引擎</b><br>
            无需配置API Key，开箱即用。基于系统内置规则自动生成施工BOM、资源清单、工艺指导书、纤芯分配表和风险提示。
        </div>
        """, unsafe_allow_html=True)


    else:
        # v3: 外部API — 可折叠配置 + 清晰引导
        with st.expander("API 配置", expanded=st.session_state.ai_config_expanded):
            # v3: 平台快速选择
            platform = st.selectbox(
                "平台",
                list(PLATFORM_PRESETS.keys()),
                index=list(PLATFORM_PRESETS.keys()).index(st.session_state.ai_platform) if st.session_state.ai_platform in PLATFORM_PRESETS else 0,
                key="sb_platform"
            )
            st.session_state.ai_platform = platform
            preset = PLATFORM_PRESETS[platform]
            is_custom = (platform == "自定义/本地")

            # 平台切换时自动重置地址和模型
            if "sb_platform_prev" not in st.session_state:
                st.session_state.sb_platform_prev = platform
            if st.session_state.sb_platform_prev != platform:
                st.session_state.ai_base_url = preset["base_url"]
                st.session_state.ai_model = preset["model"]
                st.session_state.sb_platform_prev = platform
                # 同时重置 Widget 状态 key，确保 text_input 同步刷新
                if "sb_base_url" in st.session_state:
                        st.session_state.sb_base_url = preset["base_url"]
                if "sb_model" in st.session_state:
                        st.session_state.sb_model = preset["model"]

            base_url = st.text_input(
                "API 地址",
                value=st.session_state.ai_base_url,
                disabled=not is_custom,
                key="sb_base_url"
            )
            st.session_state.ai_base_url = base_url

            model = st.text_input(
                "模型名称",
                value=st.session_state.ai_model,
                key="sb_model"
            )
            st.session_state.ai_model = model

            # API Key 输入框
            api_key = st.text_input(
                "API Key",
                type="password",
                value=st.session_state.ai_api_key,
                placeholder="请输入 API Key",
                key="sb_api_key"
            )
            st.session_state.ai_api_key = api_key

            # 按钮横向排列：2 按钮居中（spacer 列），3 按钮满宽等分
            if is_custom:
                _, col_btn1, col_btn2, _ = st.columns([1, 2, 2, 1])
            else:
                col_btn0, col_btn1, col_btn2 = st.columns([1, 1, 1])
                with col_btn0:
                    if st.button("获取Key", use_container_width=True, key="btn_help"):
                        st.session_state.show_api_guide = not st.session_state.show_api_guide
                        st.rerun()
            with col_btn1:
                if st.button("测试连接", use_container_width=True, key="btn_test"):
                        st.session_state.api_test_result = None
                        st.session_state.api_connection_verified = False
                        if not api_key:
                            st.session_state.api_test_result = {"status": "error", "msg": "请先填入 API Key"}
                        elif not base_url:
                            st.session_state.api_test_result = {"status": "error", "msg": "请填写 API 地址"}
                        else:
                            # === 第一层：本地格式强校验 ===
                            is_valid, err_msg = validate_api_key(api_key, platform)
                            if not is_valid:
                                st.session_state.api_test_result = {"status": "error", "msg": f"Key 格式无效: {err_msg}"}
                            else:
                                # === 第二层：真实网络请求验证（含超时+重试） ===
                                import time as _time
                                max_retries = 2
                                last_error = ""
                                for attempt in range(max_retries + 1):
                                    try:
                                        client = get_client(base_url, api_key, platform=platform)
                                        # 双层验证：先发一个极短请求确认连通性
                                        resp = client.chat.completions.create(
                                            model=model,
                                            messages=[{"role": "user", "content": "OK"}],
                                            timeout=15,
                                            max_tokens=5
                                        )
                                        # 验证响应是否包含有效内容（先校验类型再取属性）
                                        if not hasattr(resp, 'choices'):
                                            last_error = "响应格式错误：返回非标准对象，请检查 API 地址和模型配置"
                                        elif not resp.choices or len(resp.choices) == 0:
                                            last_error = "API 返回异常：无有效 choices"
                                        elif not hasattr(resp.choices[0], 'message'):
                                            last_error = "响应格式错误：choices 缺少 message 字段"
                                        else:
                                            content = resp.choices[0].message.content
                                            if content is not None:
                                                st.session_state.api_test_result = {"status": "success"}
                                                st.session_state.api_connection_verified = True
                                                break
                                            else:
                                                last_error = "API 返回空内容，Key 可能无效或额度不足"
                                    except AttributeError:
                                        last_error = "响应格式错误：返回非标准结构，请检查 API 地址和模型配置"
                                    except Exception as e:
                                        last_error = str(e)[:200]
                                        if attempt < max_retries:
                                            _time.sleep(1.0)
                                else:
                                    # 所有重试均失败
                                    st.session_state.api_test_result = {"status": "error", "msg": f"连接失败: {last_error}"}
                        st.rerun()
            with col_btn2:
                if st.button("清除", use_container_width=True, key="btn_clear"):
                        st.session_state.ai_api_key = ""
                        st.session_state.ai_base_url = preset["base_url"]
                        st.session_state.ai_model = preset["model"]
                        st.session_state.api_test_result = None
                        st.session_state.api_connection_verified = False
                        st.rerun()

            # 测试连接结果提示（列外渲染，占满侧边栏宽度，避免列内挤压竖排）
            if st.session_state.get("api_test_result"):
                result = st.session_state.api_test_result
                if result["status"] == "success":
                    st.markdown("""
                    <div style="background:#ecfdf5;border:1px solid #10b981;border-left:3px solid #10b981;border-radius:8px;padding:8px 16px;margin-bottom:12px;display:flex;align-items:center;gap:10px;font-size:0.8rem;">
                        <span style="font-weight:700;color:#10b981;flex-shrink:0;">✅</span>
                        <span style="color:#6b7280;">连接成功 — API 配置有效</span>
                    </div>
                    """, unsafe_allow_html=True)
                elif result["status"] == "error":
                    st.markdown(f"""
                    <div style="background:#fef2f2;border:1px solid #dc2626;border-left:3px solid #dc2626;border-radius:8px;padding:8px 16px;margin-bottom:12px;display:flex;align-items:center;gap:10px;font-size:0.8rem;">
                        <span style="font-weight:700;color:#dc2626;flex-shrink:0;">⚠️</span>
                        <span style="color:#6b7280;">{result['msg']}</span>
                    </div>
                    """, unsafe_allow_html=True)

            # 状态文字单独一行
            if st.session_state.show_api_guide:
                # 为每个平台提供对应的API指引
                api_guides = {
                    "硅基流动": """
                    <div class="config-guide-box">
                    <b>获取硅基流动 API Key 步骤：</b><br>
                    1. 访问 <code>https://cloud.siliconflow.cn</code><br>
                    2. 注册后在"API Keys"页面创建新密钥<br>
                    3. 复制 Key 粘贴到上方输入框<br>
                    4. 点击"测试连接"验证
                    </div>
                    """,
                    "DeepSeek": """
                    <div class="config-guide-box">
                    <b>获取 DeepSeek API Key 步骤：</b><br>
                    1. 访问 <code>https://platform.deepseek.com</code><br>
                    2. 注册后在"API Keys"页面创建新密钥<br>
                    3. 复制 Key 粘贴到上方输入框<br>
                    4. 点击"测试连接"验证
                    </div>
                    """,
                    "OpenAI": """
                    <div class="config-guide-box">
                    <b>获取 OpenAI API Key 步骤：</b><br>
                    1. 访问 <code>https://platform.openai.com</code><br>
                    2. 注册后在"API Keys"页面创建新密钥<br>
                    3. 复制 Key 粘贴到上方输入框<br>
                    4. 点击"测试连接"验证
                    </div>
                    """,
                    "通义千问(阿里云百炼)": """
                    <div class="config-guide-box">
                    <b>获取通义千问 API Key 步骤：</b><br>
                    1. 访问 <code>https://dashscope.console.aliyun.com</code><br>
                    2. 注册后在"API Keys"页面创建新密钥<br>
                    3. 复制 Key 粘贴到上方输入框<br>
                    4. 点击"测试连接"验证
                    </div>
                    """,
                    "智谱AI": """
                    <div class="config-guide-box">
                    <b>获取智谱AI API Key 步骤：</b><br>
                    1. 访问 <code>https://open.bigmodel.cn</code><br>
                    2. 注册后在"API Keys"页面创建新密钥<br>
                    3. 复制 Key 粘贴到上方输入框<br>
                    4. 点击"测试连接"验证
                    </div>
                    """,
                    "自定义/本地": """
                    <div class="config-guide-box">
                    <b>配置自定义/本地模型：</b><br>
                    1. 输入您的本地模型API地址<br>
                    2. 输入对应的模型名称<br>
                    3. 如有API Key，请填写到密钥输入框<br>
                    4. 点击"测试连接"验证
                    </div>
                    """
                }
                
                # 显示对应平台的指引
                guide = api_guides.get(st.session_state.ai_platform, api_guides["OpenAI"])
                st.markdown(guide, unsafe_allow_html=True)

    st.markdown("---")

    # ========== 第三部分：启动按钮 ==========
    
    # v4: 内置模式分步进度展示
    if st.session_state.ai_builtin_running:
        BUILTIN_STEPS = [
            {"key": "compliance", "label": "合规审查", "weight": 30},
            {"key": "bom", "label": "施工BOM", "weight": 14},
            {"key": "bor", "label": "资源清单", "weight": 14},
            {"key": "bop", "label": "工艺指导书", "weight": 14},
            {"key": "fiber", "label": "纤芯分配表", "weight": 14},
            {"key": "risk", "label": "风险提示", "weight": 14},
        ]
        TOTAL_WEIGHT = 100
        step_idx = st.session_state.ai_builtin_step - 1
        total_steps = len(BUILTIN_STEPS)
        
        # 超时检查
        BUILTIN_TIMEOUT = 600
        elapsed = time.time() - st.session_state.ai_builtin_start_time
        if elapsed > BUILTIN_TIMEOUT:
            st.warning(f"任务已执行 {int(elapsed)} 秒，超过10分钟。当前进度：{st.session_state.get('_builtin_status_text', '处理中')}")
        
        # 计算进度百分比
        base_progress = sum(BUILTIN_STEPS[i]["weight"] for i in range(step_idx)) if step_idx > 0 else 0
        current_label = BUILTIN_STEPS[step_idx]["label"] if step_idx < total_steps else "完成"
        current_weight = BUILTIN_STEPS[step_idx]["weight"] if step_idx < total_steps else 0
        
        sites = st.session_state.ai_builtin_sites
        if sites is not None and step_idx == 0:
            cr_total = len(sites)
            cr_done = st.session_state.cr_chunk_start
            sub_pct = cr_done / cr_total if cr_total > 0 else 1.0
            progress_pct = int((base_progress + current_weight * sub_pct) / TOTAL_WEIGHT * 100)
            status_text = f"正在执行合规审查...已完成 {cr_done}/{cr_total} 行"
        else:
            progress_pct = int((base_progress + current_weight) / TOTAL_WEIGHT * 100) if step_idx < total_steps else 100
            status_text = f"正在执行{current_label}..."
        
        st.session_state._builtin_status_text = status_text
        elapsed_str = f"{int(elapsed)}s" if elapsed < 60 else f"{int(elapsed//60)}分{int(elapsed%60)}秒"
        
        st.components.v1.html(f"""
        <style>
        @keyframes pulse-builtin {{
            0%, 100% {{ opacity: 1; }}
            50% {{ opacity: 0.65; }}
        }}
        .blt-progress-fill {{
            background: linear-gradient(90deg, #4f46e5, #7c3aed);
            height: 6px;
            border-radius: 8px;
            width: {progress_pct}%;
            transition: width 0.4s ease;
            animation: pulse-builtin 1.8s ease-in-out infinite;
        }}
        </style>
        <div style="background:#fff;border:1px solid #bfdbfe;border-radius:10px;padding:12px 14px;margin:8px 0;box-shadow:0 1px 3px rgba(0,0,0,0.06);font-family:-apple-system,BlinkMacSystemFont,sans-serif;">
            <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:6px;">
                <span style="font-size:0.78rem;font-weight:600;color:#1e40af;">{status_text}</span>
                <span style="font-size:0.68rem;color:#6b7280;">{elapsed_str}</span>
            </div>
            <div style="background:#e5e7eb;border-radius:8px;height:6px;overflow:hidden;">
                <div class="blt-progress-fill"></div>
            </div>
            <div style="font-size:0.6rem;color:#9ca3af;margin-top:4px;">任务执行中，请稍候...</div>
        </div>
        """, height=72)
        
        # 状态机：按步骤执行
        step = st.session_state.ai_builtin_step
        if step == 1:
            # 合规审查（分块处理，每100行更新一次进度）
            cr_result = run_compliance_review(
                sites,
                start_row=st.session_state.cr_chunk_start,
                chunk_size=100,
                prev_results=st.session_state.cr_chunk_results,
                prev_connection_pairs=st.session_state.cr_chunk_cp
            )
            if isinstance(cr_result, tuple):
                # 分块未完成: (False, results, cp, next_start)
                st.session_state.cr_chunk_results = cr_result[1]
                st.session_state.cr_chunk_cp = cr_result[2]
                st.session_state.cr_chunk_start = cr_result[3]
                time.sleep(0.02)
                st.rerun()
            else:
                # 合规审查完成，cr_result 是 DataFrame
                review_df = cr_result
                st.session_state.review_results = review_df
                st.session_state.excel_bytes = None
                st.session_state.word_report_bytes = None
                st.session_state.word_bop_bytes = None
                st.session_state.review_failures_bytes = None
                
                failures = review_df[review_df["结果"].str.contains("失败", na=False)]
                if len(failures) > 0:
                    # 不通过 → 生成审查报告后直接结束
                    st.session_state.review_failures_bytes = build_review_failures_excel_bytes()
                    st.session_state.compliance_full_excel = build_compliance_review_full_excel_bytes(current_df)
                    st.session_state.word_report_bytes = build_word_report_bytes(current_df)
                    st.session_state.excel_bytes = None
                    st.session_state.word_bop_bytes = None
                    st.session_state.review_failed = True
                    st.session_state.ai_data = {}
                    st.session_state.ai_dataframes = {}
                    st.session_state.result_df = current_df.copy()
                    st.session_state.offline_mode = False
                    st.session_state.ai_step_times = {"total": time.time() - st.session_state.ai_builtin_start_time}
                    st.session_state.ai_builtin_running = False
                    st.session_state.ai_builtin_step = 0
                    st.session_state.ai_generation_done = True
                    st.session_state.ai_running = False
                    # 清理临时状态
                    for k in ["_builtin_status_text"]:
                        if k in st.session_state:
                            del st.session_state[k]
                    time.sleep(0.1)
                    st.rerun()
                else:
                    st.session_state.result_df = current_df.copy()
                    st.session_state.ai_builtin_step = 2
                    time.sleep(0.02)
                    st.rerun()
        
        elif step == 2:
            st.session_state._bom_df = generate_bom_data(sites)
            st.session_state.ai_builtin_step = 3
            time.sleep(0.02)
            st.rerun()
        
        elif step == 3:
            st.session_state._bor_df = generate_bor_data(sites)
            st.session_state.ai_builtin_step = 4
            time.sleep(0.02)
            st.rerun()
        
        elif step == 4:
            st.session_state._bop_content = generate_bop_content(sites)
            st.session_state.ai_builtin_step = 5
            time.sleep(0.02)
            st.rerun()
        
        elif step == 5:
            st.session_state._fiber_df = generate_fiber_data(sites)
            st.session_state.ai_builtin_step = 6
            time.sleep(0.02)
            st.rerun()
        
        elif step == 6:
            st.session_state._risk_content = generate_risk_content(sites)
            st.session_state.ai_builtin_step = 7
            time.sleep(0.02)
            st.rerun()
        
        elif step == 7:
            # 生成所有导出文件
            bom_df = st.session_state._bom_df
            bor_df = st.session_state._bor_df
            fiber_df = st.session_state._fiber_df
            st.session_state.ai_data = {
                "bom": bom_df.to_markdown(index=False),
                "bor": bor_df.to_markdown(index=False),
                "bop": st.session_state._bop_content,
                "fiber": fiber_df.to_markdown(index=False),
                "risk": st.session_state._risk_content
            }
            st.session_state.ai_dataframes = {
                "bom": bom_df, "bor": bor_df, "fiber": fiber_df
            }
            st.session_state.excel_bytes = build_excel_bytes(current_df)
            st.session_state.word_report_bytes = build_word_report_bytes(current_df)
            st.session_state.word_bop_bytes = build_word_bop_bytes()
            st.session_state.review_failures_bytes = build_review_failures_excel_bytes()
            st.session_state.compliance_full_excel = build_compliance_review_full_excel_bytes(current_df)
            st.session_state.offline_mode = True
            st.session_state.review_failed = False
            st.session_state.ai_step_times = {"total": time.time() - st.session_state.ai_builtin_start_time}
            
            # 清理临时状态
            for k in ["_bom_df", "_bor_df", "_fiber_df", "_bop_content", "_risk_content", "_builtin_status_text"]:
                if k in st.session_state:
                    del st.session_state[k]
            
            st.session_state.ai_builtin_running = False
            st.session_state.ai_builtin_step = 0
            st.session_state.ai_generation_done = True
            st.session_state.ai_running = False
            time.sleep(0.15)
            st.rerun()
    
    # 原有按钮逻辑（内置模式已走上方进度，仅外部API模式或空闲时走此分支）
    if not st.session_state.ai_running and not st.session_state.ai_builtin_running:
        # 检查合规审查结果
        review_fail_count = 0
        if st.session_state.review_results is not None:
            review_df = st.session_state.review_results
            review_fail_count = len(review_df[review_df["结果"].str.contains("失败", na=False)])

        if st.session_state.use_builtin_ai:
            btn_label = "启动分析"
        else:
            if not st.session_state.ai_api_key:
                btn_label = "启动分析（需先配置API）"
            elif not st.session_state.api_connection_verified:
                btn_label = "启动分析（需先测试连接）"
            else:
                btn_label = "启动分析"

        btn_disabled = (not st.session_state.use_builtin_ai and not st.session_state.api_connection_verified)

        # v3: 无效文件阻止分析
        if not is_valid:
            st.error("当前文件存在数据错误，无法执行分析。请先修正文件或切换到其他数据源。")
            btn_disabled = True

        # 合规审查不通过 → 阻止BOM生成
        if review_fail_count > 0:
            st.error(f"合规审查存在 {review_fail_count} 项不通过，请修正后重新上传文件，或点击「重分析」重新审查。")
            btn_disabled = True

        clicked = st.button(btn_label, type="primary", use_container_width=True, key="sb_launch", disabled=btn_disabled)
        if clicked:
            # 执行分析前再次验证
            re_valid, re_errors, re_warnings, _ = validate_uploaded_file(current_df, current_item.get("name", ""))
            if re_errors:
                st.error("文件验证失败，无法执行分析。")
            else:
                if st.session_state.use_builtin_ai:
                    # v4: 内置模式 → 启动分步进度执行
                    st.session_state.result_df = current_df.copy()
                    st.session_state.ai_builtin_running = True
                    st.session_state.ai_builtin_step = 1
                    st.session_state.ai_builtin_start_time = time.time()
                    st.session_state.ai_builtin_sites = current_df.copy()
                    st.session_state.cr_chunk_start = 0
                    st.session_state.cr_chunk_results = []
                    st.session_state.cr_chunk_cp = {}
                    st.session_state.excel_bytes = None
                    st.session_state.word_report_bytes = None
                    st.session_state.word_bop_bytes = None
                    st.session_state.review_failures_bytes = None
                    st.session_state.review_failed = False
                    st.session_state.offline_mode = False
                    st.session_state.ai_data = {}
                    st.session_state.ai_dataframes = {}
                    st.session_state.ai_step_times = {}
                    time.sleep(0.05)
                    st.rerun()
                else:
                    # 外部API模式（保持原有逻辑）
                    review_df = run_compliance_review(current_df)
                    st.session_state.review_results = review_df
                    st.session_state.excel_bytes = None
                    st.session_state.word_report_bytes = None
                    st.session_state.word_bop_bytes = None
                    st.session_state.review_failures_bytes = None

                    failures = review_df[review_df["结果"].str.contains("失败", na=False)]
                    if len(failures) > 0:
                        st.session_state.review_failures_bytes = build_review_failures_excel_bytes()
                        st.session_state.compliance_full_excel = build_compliance_review_full_excel_bytes(current_df)
                        st.session_state.excel_bytes = None
                        st.session_state.word_report_bytes = build_word_report_bytes(current_df)
                        st.session_state.word_bop_bytes = None
                        st.session_state.review_failed = True
                        st.session_state.ai_running = False
                        st.session_state.ai_generation_done = True
                        st.session_state.offline_mode = False
                        st.session_state.ai_data = {}
                        st.session_state.ai_dataframes = {}
                        st.session_state.result_df = current_df.copy()
                        st.session_state.ai_step_times = {"total": 0}
                        time.sleep(0.15)
                        st.rerun()

                    st.session_state.result_df = current_df.copy()
                    st.session_state.ai_running = True
                    st.session_state.ai_generation_done = False
                    st.session_state.ai_step_index = 0
                    st.session_state.ai_start_time = time.time()
                    st.session_state.ai_step_times = {}
                    st.session_state.ai_timeout = False
                    st.session_state.ai_data = {}
                    st.session_state.offline_mode = False
                    time.sleep(0.15)
                    st.rerun()

    # AI运行中状态
    if st.session_state.ai_running:
        TIMEOUT_SECONDS = 600
        elapsed = time.time() - st.session_state.ai_start_time
        if elapsed > TIMEOUT_SECONDS:
            st.session_state.ai_running = False
            st.session_state.ai_generation_done = True
            st.session_state.ai_timeout = True
            if "total" not in st.session_state.ai_step_times:
                st.session_state.ai_step_times["total"] = elapsed
            st.rerun()

        si = st.session_state.ai_step_index
        total_steps = len(AI_STEPS)
        current_label = AI_STEPS[si]["label"] if si < total_steps else "完成"
        progress_pct = int((si / total_steps) * 100) if total_steps > 0 else 0
        start_js = int(st.session_state.ai_start_time * 1000)

        st.components.v1.html(f"""
        <style>
        @keyframes pulse-progress {{
            0%, 100% {{ opacity: 1; }}
            50% {{ opacity: 0.65; }}
        }}
        .progress-fill-a {{
            background: linear-gradient(90deg, #4f46e5, #7c3aed);
            height: 6px;
            border-radius: 8px;
            width: {progress_pct}%;
            transition: width 0.6s ease;
            animation: pulse-progress 1.8s ease-in-out infinite;
        }}
        </style>
        <div style="background:#fff;border:1px solid #bfdbfe;border-radius:10px;padding:12px 14px;margin:0;box-shadow:0 1px 3px rgba(0,0,0,0.06);font-family:-apple-system,BlinkMacSystemFont,sans-serif;">
            <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:6px;">
                <span style="font-size:0.78rem;font-weight:600;color:#1e40af;">生成中 · {current_label}</span>
                <span id="live-timer-a" style="font-size:0.68rem;color:#6b7280;">0s</span>
            </div>
            <div style="background:#e8edf2;border-radius:8px;height:6px;overflow:hidden;">
                <div class="progress-fill-a"></div>
            </div>
            <div style="font-size:0.65rem;color:#9ca3af;margin-top:4px;">{si}/{total_steps} 步骤</div>
        </div>
        <script>
        (function(){{
            var st = {start_js};
            var el = document.getElementById('live-timer-a');
            function t(){{ el.textContent = Math.floor((Date.now()-st)/1000)+'s'; }}
            setInterval(t,1000); t();
        }})();
        </script>
        """, height=90)



        if st.button("取消生成", use_container_width=True, key="sb_cancel"):
            st.session_state.ai_running = False
            st.session_state.ai_generation_done = True
            st.rerun()

    # ========== 第四部分：一键下载 ==========
    if st.session_state.ai_generation_done and not st.session_state.ai_running:
        total = st.session_state.ai_step_times.get("total", 0)
        if st.session_state.get("ai_timeout", False):
            st.error("生成超时（已超过 10 分钟），已自动终止")
        # v5: 任务历史记录
        from datetime import datetime as dt_now
        task_review = st.session_state.review_results
        if task_review is not None:
            task_fail = len(task_review[task_review["结果"].str.contains("失败", na=False)])
            task_warn = len(task_review[task_review["结果"].str.contains("警告", na=False)])
            task_pass = len(task_review[task_review["结果"].str.contains("通过", na=False)])
        else:
            task_fail, task_warn, task_pass = 0, 0, 0
        
        # 判断状态
        is_rf = st.session_state.get("review_failed", False)
        if is_rf:
            task_status = "failed"
        elif task_warn > 0:
            task_status = "warning"
        else:
            task_status = "success"
        
        # 受影响站点数
        task_sites = set()
        if task_review is not None:
            for _, r in task_review.iterrows():
                res = str(r.get("结果", ""))
                if "失败" in res or "警告" in res:
                    task_sites.add(str(r.get("站点", "")))
        
        task_files = []
        if st.session_state.excel_bytes: task_files.append("excel_bytes")
        if st.session_state.word_report_bytes: task_files.append("word_report_bytes")
        if st.session_state.word_bop_bytes: task_files.append("word_bop_bytes")
        if st.session_state.review_failures_bytes: task_files.append("review_failures_bytes")
        
        st.session_state.task_history.append({
            "id": f"TASK-{len(st.session_state.task_history)+1:04d}",
            "status": task_status,
            "timestamp": dt_now.now().strftime("%Y-%m-%d %H:%M:%S"),
            "duration": st.session_state.ai_step_times.get("total", 0),
            "compliance_summary": {"total": len(task_review) if task_review is not None else 0, "fail": task_fail, "warn": task_warn, "pass": task_pass},
            "files": task_files,
            "sites_count": len(task_sites),
            "error_msg": "合规审查未通过" if is_rf else ""
        })

        st.markdown(f"""
        <div class="completion-banner">
            生成完成 · 总耗时 {total:.1f}s · 交付文件已就绪，请前往「工程概览」页面底部下载
        </div>
        """, unsafe_allow_html=True)

# ==================== 主页面 (v3 升级) ====================
st.markdown("""
<div style="margin-bottom:6px;">
    <div class="page-title">5G通信基建数智化交付系统</div>
    <div class="page-subtitle">RK-001~RK-010 合规审查 · 12项标准BOM · 16项站点字段 · Word/Excel 双格式交付</div>
</div>
""", unsafe_allow_html=True)

# 状态提示 — 三态：内置 / 外部已配置 / 外部未配置
if st.session_state.use_builtin_ai:
    mode_color = "#10b981"
    mode_bg = "#ecfdf5"
    mode_text = "标准AI引擎"
    mode_hint = "开箱即用，无需配置"
    mode_desc = "内置合规审查模型驱动，本地处理文档，数据不出设备"
elif st.session_state.ai_api_key and st.session_state.api_connection_verified:
    mode_color = "#4f46e5"
    mode_bg = "#eef2ff"
    mode_text = "外部AI引擎"
    mode_hint = "已连接远程大模型，全量功能可用"
    mode_desc = f"平台：{st.session_state.ai_platform} · 模型：{st.session_state.ai_model}"
elif st.session_state.ai_api_key:
    mode_color = "#f59e0b"
    mode_bg = "#fffbeb"
    mode_text = "外部AI引擎"
    mode_hint = "已填写API配置，点击测试连接启用AI"
    mode_desc = f"平台：{st.session_state.ai_platform} · 模型：{st.session_state.ai_model} · 请点击「测试连接」验证"
else:
    mode_color = "#ef4444"
    mode_bg = "#fef2f2"
    mode_text = "API 未配置"
    mode_hint = "合规审查可用，配置 API Key 即可启用全量 AI 生成"
    mode_desc = "在侧边栏「AI 引擎 → API 配置」中填入 API Key 激活全量功能"

st.markdown(f"""
<div style="background:{mode_bg};border:1px solid {mode_color}22;border-left:3px solid {mode_color};border-radius:8px 8px 0 0;padding:8px 16px;display:flex;align-items:center;gap:10px;font-size:0.8rem;">
    <span style="font-weight:700;color:{mode_color};">{mode_text}</span>
    <span style="color:#6b7280;">{mode_hint}</span>
</div>
<div style="background:{mode_bg};border:1px solid {mode_color}11;border-top:none;border-left:3px solid transparent;border-radius:0 0 8px 8px;padding:2px 16px 8px 16px;margin-bottom:16px;font-size:0.75rem;color:#9ca3af;">
    {mode_desc}
</div>
""", unsafe_allow_html=True)

current_df = st.session_state.uploaded_files[st.session_state.current_idx]["df"]

# ===== 工程概览卡片行 (3列网格) =====
st.markdown('<div class="section-title">工程概览</div>', unsafe_allow_html=True)

# 快速状态栏
has_review = st.session_state.review_results is not None
has_generated = st.session_state.ai_generation_done
if has_review:
    rdf = st.session_state.review_results
    r_fail = len(rdf[rdf["结果"].str.contains("失败", na=False)])
    r_warn = len(rdf[rdf["结果"].str.contains("警告", na=False)])
    if r_fail > 0:
        st.markdown(f"""
        <div style="background:#fee2e2;border:1px solid #fca5a5;border-radius:8px;padding:10px 16px;margin-bottom:14px;display:flex;align-items:center;gap:12px;">
            <span style="font-weight:700;color:#dc2626;font-size:0.82rem;">审查状态</span>
            <span style="background:#dc2626;color:#fff;padding:2px 12px;border-radius:10px;font-size:0.7rem;font-weight:700;">{r_fail} 项不通过</span>
            {f'<span style="background:#f59e0b;color:#fff;padding:2px 12px;border-radius:10px;font-size:0.7rem;font-weight:700;">{r_warn} 项警告</span>' if r_warn > 0 else ''}
            <span style="color:#6b7280;font-size:0.75rem;margin-left:auto;">请前往「合规审查」标签页查看详情</span>
        </div>
        """, unsafe_allow_html=True)



# 第一行：2列
ov_row1 = st.columns(2)

with ov_row1[0]:
    st.markdown(f"""
    <div class="metric-card status-completed">
        <div class="label">站点数量</div>
        <div class="value">{len(current_df)}<span class="suffix">个</span></div>
    </div>
    """, unsafe_allow_html=True)

with ov_row1[1]:
    if "路由长度(m)" in current_df.columns and not current_df["路由长度(m)"].dropna().empty:
        total_len = int(current_df["路由长度(m)"].apply(safe_float_route).sum())
        max_len = int(current_df["路由长度(m)"].apply(safe_float_route).max())
        st.markdown(f"""
        <div class="metric-card status-in-progress">
            <div class="label">总路由长度</div>
            <div class="value">{total_len}<span class="suffix">米</span></div>
            <div style="font-size:0.72rem;color:#6b7280;margin-top:4px;">最长单段: {max_len}米</div>
        </div>
        """, unsafe_allow_html=True)


    else:
        st.markdown(f"""
        <div class="metric-card status-pending">
            <div class="label">总路由长度</div>
            <div style="font-size:2.5rem;font-weight:800;color:#9ca3af;">—</div>
            <div style="font-size:0.72rem;color:#d1d5db;margin-top:4px;">推荐补充字段</div>
        </div>
        """, unsafe_allow_html=True)



# 第二行：2列
ov_row2 = st.columns(2)

with ov_row2[0]:
    if "AAU型号" in current_df.columns:
        aaus = current_df["AAU型号"].dropna().unique().tolist()
        bbus = current_df["BBU型号"].dropna().unique().tolist()
        st.markdown(f"""
        <div class="metric-card status-info">
            <div class="label">设备型号</div>
            <div class="value">{len(aaus)+len(bbus)}<span class="suffix">种</span></div>
            <div style="font-size:0.72rem;color:#6b7280;margin-top:4px;">AAU: {len(aaus)}种 | BBU: {len(bbus)}种</div>
        </div>
        """, unsafe_allow_html=True)


    else:
        st.markdown(f"""
        <div class="metric-card status-pending">
            <div class="label">设备型号</div>
            <div style="font-size:2.5rem;font-weight:800;color:#9ca3af;">—</div>
        </div>
        """, unsafe_allow_html=True)



with ov_row2[1]:
    if "站点类型" in current_df.columns:
        types = current_df["站点类型"].value_counts().to_dict()
        main_type = max(types, key=types.get) if types else ""
        main_count = types.get(main_type, 0)
        st.markdown(f"""
        <div class="metric-card status-completed">
            <div class="label">主要站点类型</div>
            <div class="value">{main_count}<span class="suffix">个</span></div>
            <div style="font-size:0.72rem;color:#6b7280;margin-top:4px;">{main_type}</div>
        </div>
        """, unsafe_allow_html=True)


    else:
        st.markdown(f"""
        <div class="metric-card status-pending">
            <div class="label">站点类型分布</div>
            <div style="font-size:2.5rem;font-weight:800;color:#9ca3af;">—</div>
        </div>
        """, unsafe_allow_html=True)



st.markdown("---")

# ===== v3: 文件详情弹窗 =====
if st.session_state.show_file_detail:
    current_item = st.session_state.uploaded_files[st.session_state.current_idx]
    with st.expander("文件详情", expanded=True):
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            st.markdown(f"**文件名**: {current_item.get('name','')}")
            st.markdown(f"**验证状态**: {'有效' if current_item.get('valid', True) else '有错误'}")
            st.markdown(f"**站点数**: {len(current_item['df'])}")
            st.markdown(f"**字段数**: {len(current_item['df'].columns)}")
        with col_d2:
            errors = current_item.get('errors', [])
            warnings = current_item.get('warnings', [])
            if errors:
                st.error("**错误:**")
                for e in errors:
                    st.error(f"- {e}")
            if warnings:
                st.warning("**警告:**")
                for w in warnings:
                    st.warning(f"- {w}")
        st.dataframe(current_item["df"], use_container_width=True, hide_index=True)
    if st.button("关闭详情", key="close_detail"):
        st.session_state.show_file_detail = False
        st.rerun()

# ===== 标签页 =====
tab1, tab2, tab3, tab4 = st.tabs(["工程概览", "AI生成施工资料", "合规审查", "任务中心/执行日志"])

with tab1:
    # v3: 改进字段校验展示
    missing_new = [c for c in NEW_FIELD_TEMPLATE if c not in current_df.columns]
    if not missing_new:
        st.success(f"16项字段齐全 · {len(current_df)}站点 · {len(current_df.columns)}字段")
    else:
        st.warning(f"缺少推荐字段: {', '.join(missing_new)}（不影响内置分析）")

    with st.expander("查看完整原始数据", expanded=st.session_state.expand_raw_data):
        st.dataframe(current_df, use_container_width=True, hide_index=True)
        st.session_state.expand_raw_data = False
    # ===== 4列总览卡片（审查完成后显示）=====
    if st.session_state.review_results is not None:
        review_df = st.session_state.review_results
        fail_count = len(review_df[review_df["结果"].str.contains("失败", na=False)])
        warn_count = len(review_df[review_df["结果"].str.contains("警告", na=False)])
        pass_count = len(review_df[review_df["结果"].str.contains("通过", na=False)])
        high_risk_count = len(review_df[review_df["风险"].str.contains("高风险", na=False)])
        total_elapsed = st.session_state.ai_step_times.get("total", 0)
        
        if fail_count > 0:
            compliance_status = "不通过"
            compliance_color = "#dc2626"
            compliance_bg = "#fef2f2"
            compliance_border = "#fca5a5"
        elif warn_count > 0:
            compliance_status = "警告"
            compliance_color = "#d97706"
            compliance_bg = "#fffbeb"
            compliance_border = "#fde68a"
        else:
            compliance_status = "通过"
            compliance_color = "#059669"
            compliance_bg = "#ecfdf5"
            compliance_border = "#86efac"
        
        affected_sites = set()
        for _, row in review_df.iterrows():
            r = str(row.get("结果", ""))
            if "失败" in r or "警告" in r:
                affected_sites.add(str(row.get("站点", "")))
        affected_count = len(affected_sites)

        st.markdown('<div class="section-title">分析概览</div>', unsafe_allow_html=True)
        ov_cols = st.columns(4)
        card_items = [
            ("合规状态", compliance_status, compliance_color, compliance_bg, compliance_border,
             "通过" if compliance_status == "通过" else ("警告" if compliance_status == "警告" else "不通过")),
            ("高风险项", str(high_risk_count), "#dc2626", "#fef2f2", "#fca5a5", "项高风险"),
            ("总耗时", f"{total_elapsed:.1f}s", "#2563eb", "#eff6ff", "#93c5fd", "秒"),
            ("受影响站点", str(affected_count), "#ea580c", "#fff7ed", "#fdba74", "个站点"),
        ]
        for col, (title, val, color, bg, border, sub) in zip(ov_cols, card_items):
            with col:
                st.markdown(f"""
                <div style="background:{bg};border:2px solid {border};border-radius:12px;padding:16px 18px;text-align:center;">
                    <div style="font-size:0.7rem;font-weight:700;color:{color};text-transform:uppercase;letter-spacing:0.5px;margin-bottom:6px;">{title}</div>
                    <div style="font-size:2rem;font-weight:800;color:#111827;line-height:1.1;">{val}</div>
                    <div style="font-size:0.72rem;color:#6b7280;margin-top:4px;">{sub}</div>
                </div>
                """, unsafe_allow_html=True)

    # ===== AI生成进度（内置模式）=====
    if st.session_state.ai_builtin_running:
        TIMEOUT_SECONDS = 600
        elapsed = time.time() - st.session_state.ai_builtin_start_time
        if elapsed > TIMEOUT_SECONDS:
            st.session_state.ai_builtin_running = False
            st.session_state.ai_builtin_step = 0
            st.session_state.ai_generation_done = True
            st.session_state.ai_timeout = True
            if "total" not in st.session_state.ai_step_times:
                st.session_state.ai_step_times["total"] = elapsed
            st.rerun()

        step = st.session_state.ai_builtin_step
        step_names = {1: "合规审查", 2: "施工BOM", 3: "资源需求清单", 4: "工艺指导书", 5: "纤芯分配表", 6: "风险提示", 7: "导出文件", 8: "完成"}
        step_name = step_names.get(step, "处理中")
        progress_pct = int((step / 8) * 100)
        
        st.subheader("AI 生成进度")
        st.progress(progress_pct / 100, text=f"步骤 {step}/8: {step_name}")
        
        desc_map = {1: "正在加载数据(10%)", 2: "AI分析(30%)", 3: "AI分析(30%)", 4: "合规审查(60%)", 5: "生成报告(90%)", 6: "生成报告(90%)", 7: "导出文件(95%)", 8: "完成(100%)"}
        desc_text = desc_map.get(step, "处理中")
        status_placeholder = st.empty()
        status_placeholder.info(f"当前阶段: {desc_text}")
        
        progress_cols = st.columns(8)
        step_labels = ["合规审查", "施工BOM", "资源清单", "工艺指导书", "纤芯分配", "风险提示", "导出文件", "完成"]
        for idx, (col, label) in enumerate(zip(progress_cols, step_labels)):
            s_idx = idx + 1
            with col:
                with st.container(border=True):
                    if s_idx < step:
                        st.success(f"{s_idx}. {label}")
                    elif s_idx == step:
                        st.warning(f"{s_idx}. {label}")
                    else:
                        st.info(f"{s_idx}. {label}")

    # ===== AI生成进度（外部API模式）=====
    elif st.session_state.ai_running:
        TIMEOUT_SECONDS = 600
        elapsed = time.time() - st.session_state.ai_start_time
        if elapsed > TIMEOUT_SECONDS:
            st.session_state.ai_running = False
            st.session_state.ai_generation_done = True
            st.session_state.ai_timeout = True
            if "total" not in st.session_state.ai_step_times:
                st.session_state.ai_step_times["total"] = elapsed
            st.rerun()

        si = st.session_state.ai_step_index
        st.subheader("AI 生成进度")
        
        start_js = int(st.session_state.ai_start_time * 1000)
        total_steps = len(AI_STEPS)
        safe_pct = int((si / total_steps) * 100) if total_steps > 0 else 0
        st.components.v1.html(f"""
        <style>
        * {{ margin:0; padding:0; box-sizing:border-box; }}
        @keyframes pulse-progress-b {{
            0%, 100% {{ opacity: 1; }}
            50% {{ opacity: 0.65; }}
        }}
        .wrap {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            padding: 4px 0 0 0;
        }}
        .bar-bg {{
            background: #e8edf2;
            border-radius: 8px;
            height: 6px;
            overflow: hidden;
        }}
        .bar-fill {{
            background: linear-gradient(90deg, #4f46e5, #7c3aed);
            height: 6px;
            border-radius: 8px;
            width: {safe_pct}%;
            transition: width 0.6s ease;
            animation: pulse-progress-b 1.8s ease-in-out infinite;
        }}
        .info {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-top: 5px;
            font-size: 0.75rem;
            color: #6b7280;
        }}
        </style>
        <div class="wrap">
            <div class="bar-bg"><div class="bar-fill"></div></div>
            <div class="info">
                <span>总进度 {si}/{total_steps} 步骤</span>
                <span id="live-timer-b" style="font-weight:500;">0s</span>
            </div>
        </div>
        <script>
        (function(){{
            var start = {start_js};
            var el = document.getElementById('live-timer-b');
            function tick(){{ el.textContent = Math.floor((Date.now() - start) / 1000) + 's'; }}
            setInterval(tick, 1000);
            tick();
        }})();
        </script>
        """, height=80)
        
        progress_cols = st.columns(len(AI_STEPS))
        for idx, (col, s) in enumerate(zip(progress_cols, AI_STEPS)):
            with col:
                with st.container(border=True):
                    if idx < si:
                        t = st.session_state.ai_step_times.get(s['key'], 0)
                        st.success(f"{s['label']}")
                        st.caption(f"{t:.1f}s")
                    elif idx == si:
                        st.warning(f"{s['label']}")
                    else:
                        st.info(f"{s['label']}")

        if si < len(AI_STEPS):
            elapsed = time.time() - st.session_state.ai_start_time
            if elapsed > 600:
                st.session_state.ai_running = False
                st.session_state.ai_generation_done = True
                st.session_state.ai_timeout = True
                if "total" not in st.session_state.ai_step_times:
                    st.session_state.ai_step_times["total"] = elapsed
                st.rerun()
            step = AI_STEPS[si]
            try:
                client = get_client(st.session_state.ai_base_url, st.session_state.ai_api_key, platform=st.session_state.ai_platform)
                t0 = time.time()
                rv = (step["key"] == "risk")
                resp = client.chat.completions.create(
                    model=st.session_state.ai_model,
                    messages=[
                        {"role": "system", "content": step["system"]},
                        {"role": "user", "content": f"站点数据（{len(st.session_state.result_df)}个）：\\n{st.session_state.result_df.to_string()}"}
                    ],
                    timeout=90 if rv else 180,
                    temperature=0.3,
                    max_tokens=2048 if rv else 4096
                )
                st.session_state.ai_data[step["key"]] = resp.choices[0].message.content if hasattr(resp, 'choices') and resp.choices else "响应格式错误：返回非标准对象，请检查 API 地址和模型配置"
                st.session_state.ai_step_times[step["key"]] = time.time() - t0
            except Exception as e:
                st.session_state.ai_data[step["key"]] = f"生成失败: {str(e)[:300]}"
            st.session_state.ai_step_index = si + 1
            time.sleep(0.3)
            st.rerun()
        else:
            elapsed = time.time() - st.session_state.ai_start_time
            if elapsed > 600:
                st.session_state.ai_timeout = True
            st.session_state.ai_step_times["total"] = elapsed
            st.session_state.ai_running = False
            st.session_state.ai_generation_done = True
            st.session_state.excel_bytes = build_excel_bytes(current_df)
            st.session_state.word_report_bytes = build_word_report_bytes(current_df)
            st.session_state.word_bop_bytes = build_word_bop_bytes()
            st.session_state.review_failures_bytes = build_review_failures_excel_bytes()
            st.session_state.compliance_full_excel = build_compliance_review_full_excel_bytes(current_df)
            st.rerun()

    # ===== 分析完成提示 =====
    if st.session_state.ai_generation_done and not st.session_state.ai_running and not st.session_state.ai_builtin_running:
        is_review_failed = st.session_state.get("review_failed", False)
        is_timeout = st.session_state.get("ai_timeout", False)

        if is_timeout:
            st.error("生成超时（已超过 10 分钟），已自动终止。请减少数据量后重试。")

        if is_review_failed:
            st.warning("合规审查未通过，已拦截物料清单生成。请前往「交付结果」标签页下载审查报告修正数据后重新上传。")
        elif st.session_state.review_results is not None:
            review_df = st.session_state.review_results
            fail_n = len(review_df[review_df["结果"].str.contains("失败", na=False)])
            warn_n = len(review_df[review_df["结果"].str.contains("警告", na=False)])
            total_n = len(review_df)
            total_elapsed = st.session_state.ai_step_times.get("total", 0)
            
            affected_sites = set()
            for _, row in review_df.iterrows():
                r = str(row.get("结果", ""))
                if "失败" in r or "警告" in r:
                    affected_sites.add(str(row.get("站点", "")))
            affected_n = len(affected_sites)
            
            if fail_n > 0:
                st.error(f"分析完成！共检测{total_n}项，发现 警告{warn_n}项， 不通过{fail_n}项，影响 {affected_n}个站点，总耗时 {total_elapsed:.1f}秒")
            elif warn_n > 0:
                st.warning(f"分析完成！共检测{total_n}项，发现 警告{warn_n}项， 不通过{fail_n}项，影响 {affected_n}个站点，总耗时 {total_elapsed:.1f}秒")
            else:
                st.success(f"分析完成！共检测{total_n}项，全部通过，影响 {affected_n}个站点，总耗时 {total_elapsed:.1f}秒")
        elif not is_review_failed:
            total_elapsed = st.session_state.ai_step_times.get("total", 0)
            st.success(f"全部生成完成 · 总耗时 {total_elapsed:.1f}s")

    # ===== 交付结果下载模块 =====
    st.markdown("---")
    st.markdown('<div class="section-title">📦 交付结果下载</div>', unsafe_allow_html=True)

    if not st.session_state.ai_generation_done:
        st.markdown("""
        <div style="background:#f3f4f6;border:1px solid #d1d5db;border-radius:10px;padding:20px;text-align:center;margin:20px 0;">
            <div style="font-size:1.1rem;font-weight:700;color:#6b7280;margin-bottom:8px;">⏳ 正在生成交付文件，下载按钮将在生成后解锁</div>
            <div style="font-size:0.8rem;color:#9ca3af;">请在侧边栏底部点击「启动分析」，系统将自动生成所有交付文件</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        is_review_failed = st.session_state.get("review_failed", False)
        is_timeout = st.session_state.get("ai_timeout", False)

        if is_timeout:
            st.error("生成超时（已超过 10 分钟），已自动终止")
        elif is_review_failed:
            st.warning("合规审查未通过，已拦截物料清单生成。请下载审查报告修正数据后重新上传。")
        else:
            st.success("✅ 所有交付文件已生成，可下载使用")

        # 确保所有 bytes 已缓存
        if st.session_state.excel_bytes is None:
            st.session_state.excel_bytes = build_excel_bytes(current_df)
        if st.session_state.word_report_bytes is None:
            st.session_state.word_report_bytes = build_word_report_bytes(current_df)
        if st.session_state.word_bop_bytes is None:
            st.session_state.word_bop_bytes = build_word_bop_bytes()
        if st.session_state.review_failures_bytes is None:
            st.session_state.review_failures_bytes = build_review_failures_excel_bytes()
        if st.session_state.compliance_full_excel is None:
            st.session_state.compliance_full_excel = build_compliance_review_full_excel_bytes(current_df)

        # 【一键交付】
        st.caption("【一键交付】")
        import io, zipfile
        from datetime import datetime

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            if not is_review_failed and st.session_state.excel_bytes:
                zf.writestr("工程数据包（BOM_纤芯表）.xlsx", st.session_state.excel_bytes)
            if st.session_state.word_report_bytes:
                zf.writestr("综合交付报告.docx", st.session_state.word_report_bytes)
            if not is_review_failed and st.session_state.word_bop_bytes:
                zf.writestr("施工工艺指导书.docx", st.session_state.word_bop_bytes)
            if st.session_state.review_failures_bytes:
                zf.writestr("问题整改台账.xlsx", st.session_state.review_failures_bytes)
            if st.session_state.compliance_full_excel:
                cdata, _ = st.session_state.compliance_full_excel if isinstance(st.session_state.compliance_full_excel, tuple) else (st.session_state.compliance_full_excel, "")
                if cdata:
                    zf.writestr("合规审查完整明细.xlsx", cdata)

        zip_name = f"工程交付结果_{datetime.now().strftime('%Y%m%d%H%M')}.zip"
        st.download_button(
            "📦 一键下载全套交付成果（ZIP包）",
            zip_buffer.getvalue(),
            zip_name,
            "application/zip",
            use_container_width=True,
            type="primary",
            key="tab1_dl_zip",
            help="打包包含所有工程交付文件，可直接用于项目存档/甲方交付"
        )

        st.markdown("---")

        # 【核心工程文件】
        st.caption("【核心工程文件】")
        core_cols = st.columns(3)
        with core_cols[0]:
            if not is_review_failed and st.session_state.excel_bytes:
                st.download_button(
                    "📊 工程数据Excel包（BOM/纤芯表）",
                    st.session_state.excel_bytes,
                    "5G基站AI交付结果.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True, key="tab1_dl_excel",
                    help="包含施工BOM、纤芯分配表、工程统计数据"
                )
            else:
                st.button("📊 工程数据Excel包（BOM/纤芯表）", disabled=True, use_container_width=True, key="tab1_dl_excel_dis")
        with core_cols[1]:
            if st.session_state.word_report_bytes:
                st.download_button(
                    "📄 综合交付报告（Word）",
                    st.session_state.word_report_bytes,
                    "5G基站工程交付报告.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, key="tab1_dl_report",
                    help="含工程概况、合规结论、施工建议的完整交付报告"
                )
            else:
                st.button("📄 综合交付报告（Word）", disabled=True, use_container_width=True, key="tab1_dl_report_dis")
        with core_cols[2]:
            if not is_review_failed and st.session_state.word_bop_bytes:
                st.download_button(
                    "🛠️ 施工工艺指导书（Word）",
                    st.session_state.word_bop_bytes,
                    "5G基站施工工艺指导书.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, key="tab1_dl_bop",
                    help="含施工步骤、规范要求的现场操作指引"
                )
            else:
                st.button("🛠️ 施工工艺指导书（Word）", disabled=True, use_container_width=True, key="tab1_dl_bop_dis")

        # 【合规整改文件】
        st.caption("【合规整改文件】")
        aux_cols = st.columns(2)
        with aux_cols[0]:
            if st.session_state.review_failures_bytes:
                st.download_button(
                    "⚠️ 问题整改台账（带整改建议）",
                    st.session_state.review_failures_bytes,
                    "合规审查不通过项.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True, key="tab1_dl_failures",
                    help="仅展示需整改的警告/不通过项，含风险等级和整改建议"
                )
            else:
                st.button("⚠️ 问题整改台账（带整改建议）", disabled=True, use_container_width=True, key="tab1_dl_failures_dis")
        with aux_cols[1]:
            if st.session_state.compliance_full_excel:
                cdata, cname = st.session_state.compliance_full_excel if isinstance(st.session_state.compliance_full_excel, tuple) else (st.session_state.compliance_full_excel, "合规审查报告.xlsx")
                st.download_button(
                    "📋 合规审查完整明细（存档用）",
                    cdata,
                    cname,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True, key="tab1_dl_compliance",
                    help="包含本次审查的所有通过/警告/不通过项，用于存档追溯"
                )
            else:
                st.button("📋 合规审查完整明细（存档用）", disabled=True, use_container_width=True, key="tab1_dl_compliance_dis")


with tab2:
    st.markdown('<div class="section-title">AI 生成施工资料</div>', unsafe_allow_html=True)
    if not st.session_state.ai_generation_done:
        st.markdown("""
        <div style="background:#f3f4f6;border:1px solid #d1d5db;border-radius:10px;padding:20px;text-align:center;margin:20px 0;">
            <div style="font-size:1.1rem;font-weight:700;color:#6b7280;margin-bottom:8px;">尚未生成施工资料</div>
            <div style="font-size:0.8rem;color:#9ca3af;">请先在左侧控制台底部点击「启动分析」</div>
        </div>
        """, unsafe_allow_html=True)


    else:
        status_cols = st.columns(5)
        keys = ["bom", "bor", "bop", "fiber", "risk"]
        labels = ["施工BOM", "资源清单", "工艺指导书", "纤芯分配表", "风险提示"]
        icons = ["BOM", "BOR", "BOP", "FC", "RK"]

        for col, key, label, icon in zip(status_cols, keys, labels, icons):
            with col:
                content = st.session_state.ai_data.get(key, "")
                has_data = content and content != "暂无数据" and not content.startswith("生成失败")
                bg = "#f0fdf4" if has_data else "#fef2f2"
                border = "#86efac" if has_data else "#fca5a5"
                text_color = "#166534" if has_data else "#991b1b"
                badge_bg = "#dcfce7" if has_data else "#fee2e2"
                badge_text = "已生成" if has_data else "未生成"
                st.markdown(f"""
                <div style="background:{bg};border:1.5px solid {border};border-radius:10px;padding:12px 14px;text-align:center;">
                    <div style="font-size:0.7rem;font-weight:700;color:#6b7280;text-transform:uppercase;letter-spacing:0.5px;margin-bottom:4px;">{icon}</div>
                    <div style="font-weight:700;color:#111827;font-size:0.85rem;margin-bottom:8px;">{label}</div>
                    <span style="background:{badge_bg};color:{text_color};font-size:0.65rem;font-weight:700;padding:2px 12px;border-radius:8px;">{badge_text}</span>
                    </div>
                """, unsafe_allow_html=True)

        st.markdown("---")

        all_sites = list(st.session_state.result_df["站点编号"]) if st.session_state.result_df is not None else []
        view_mode = st.radio("查看模式", ["全部站点汇总", "单站点详细"], horizontal=True, key="tab2_view")
        selected_site = None
        if "单站点" in view_mode and all_sites:
            selected_site = st.selectbox("选择站点", all_sites, key="tab2_site")

        subtabs = st.tabs(["施工BOM", "资源需求清单", "工艺指导书", "纤芯分配表"])
        sub_keys = ["bom", "bor", "bop", "fiber"]
        for i, t in enumerate(subtabs):
            with t:
                key = sub_keys[i]
                # BOP（工艺指导书）始终以 Markdown 渲染
                if key == "bop":
                    content = st.session_state.ai_data.get("bop", "暂无数据")
                    if not content or content == "暂无数据":
                        st.info("暂无工艺指导书数据，请先生成。")
                    else:
                        st.markdown(content)
                    continue

                # 优先使用存储的 DataFrame（内置模式），否则从文本解析（外部 API 模式）
                df_display = st.session_state.ai_dataframes.get(key)
                if df_display is None:
                    content = st.session_state.ai_data.get(key, "")
                    if not content or content == "暂无数据" or content.startswith("生成失败"):
                        st.info(f"暂无{key}数据，请先生成。")
                        continue
                    df_display = parse_markdown_table(content)
                    if df_display is None or df_display.empty:
                        # 清洗 Markdown 标点符号再展示
                        cleaned = content
                        cleaned = re.sub(r'^\|[-:\s|]+\|\s*$', '', cleaned, flags=re.MULTILINE)
                        cleaned = re.sub(r'^\||\|$', '', cleaned, flags=re.MULTILINE)
                        cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
                        # 转义 HTML 标签，防止源码泄露
                        cleaned = html.escape(cleaned)
                        with st.container(height=400):
                            st.markdown(cleaned)
                        continue

                # 纤芯分配表支持单站点筛选（列名模糊匹配）
                if key == "fiber" and selected_site and df_display is not None and not df_display.empty:
                    site_col = None
                    if "站点编号" in df_display.columns:
                        site_col = "站点编号"
                    elif "站点" in df_display.columns:
                        site_col = "站点"
                    else:
                        for c in df_display.columns:
                            if "站点" in str(c):
                                site_col = c
                                break
                    if site_col:
                        df_display = df_display[df_display[site_col].astype(str) == str(selected_site)]
                    else:
                        st.caption("（纤芯分配表不含站点编号列，无法按站点筛选）")

                # 摘要行：条数 + 关键指标
                if df_display is None or df_display.empty:
                    st.caption(f"（{key} 表格为空或解析失败）")
                    continue
                summary_parts = [f"共 **{len(df_display)}** 条"]
                if key == "fiber":
                    # 站点列名模糊匹配
                    site_col = None
                    if "站点编号" in df_display.columns:
                        site_col = "站点编号"
                    elif "站点" in df_display.columns:
                        site_col = "站点"
                    else:
                        for c in df_display.columns:
                            if "站点" in str(c):
                                site_col = c
                                break
                    if site_col:
                        sites_n = df_display[site_col].nunique()
                        summary_parts.append(f"覆盖 **{sites_n}** 个站点")
                elif key == "bor":
                    for col_name in ["物料类别", "类别", "资源类型"]:
                        if col_name in df_display.columns:
                            types = df_display[col_name].value_counts()
                            type_str = " · ".join([f"{k}: {v}" for k, v in types.head(5).items()])
                            if type_str:
                                summary_parts.append(f"类型: {type_str}")
                            break
                st.caption(" · ".join(summary_parts))

                # 纤芯分配表多条件筛选
                if key == "fiber" and df_display is not None and not df_display.empty:
                    st.markdown("---")
                    st.caption("筛选条件")
                    site_col = None
                    if "站点编号" in df_display.columns:
                        site_col = "站点编号"
                    elif "站点" in df_display.columns:
                        site_col = "站点"
                    else:
                        for c in df_display.columns:
                            if "站点" in str(c):
                                site_col = c
                                break

                    cable_col = next((c for c in df_display.columns if "光缆" in str(c)), None)
                    color_col = next((c for c in df_display.columns if "颜色" in str(c)), None)
                    biz_col = next((c for c in df_display.columns if "业务" in str(c)), None)
                    start_col = next((c for c in df_display.columns if "起始端子" in str(c)), None)
                    end_col = next((c for c in df_display.columns if "终止端子" in str(c)), None)

                    f1, f2, f3, f4, f5, f6 = st.columns([1, 1, 1, 1, 1, 0.6])
                    filter_site = []
                    if site_col:
                        with f1:
                            all_sites_fiber = sorted(df_display[site_col].dropna().astype(str).unique().tolist())
                            filter_site = st.multiselect("站点编号", all_sites_fiber, key="fsite", label_visibility="collapsed", placeholder="站点编号")
                    filter_cable = ""
                    if cable_col:
                        with f2:
                            all_cables = sorted(df_display[cable_col].dropna().astype(str).unique().tolist())
                            filter_cable = st.selectbox("光缆编号", [""] + all_cables, key="fcable", label_visibility="collapsed")
                    filter_color = ""
                    if color_col:
                        with f3:
                            fiber_colors = ["", "蓝", "橙", "绿", "棕", "灰", "白"]
                            filter_color = st.selectbox("纤芯颜色", fiber_colors, key="fcolor", label_visibility="collapsed")
                    filter_biz = ""
                    if biz_col:
                        with f4:
                            filter_biz = st.selectbox("业务类型", ["", "4G业务", "5G业务"], key="fbiz", label_visibility="collapsed")
                    filter_start = ""
                    if start_col:
                        with f5:
                            filter_start = st.text_input("起始端子", key="fstart", label_visibility="collapsed", placeholder="起始端子")
                    with f6:
                        reset_btn = st.button("重置", key="freset", use_container_width=True)

                    if reset_btn:
                        st.session_state.pop("fsite", None)
                        st.session_state.pop("fcable", None)
                        st.session_state.pop("fcolor", None)
                        st.session_state.pop("fbiz", None)
                        st.session_state.pop("fstart", None)
                        st.rerun()

                    # 应用筛选
                    df_filtered = df_display.copy()
                    if filter_site and site_col:
                        df_filtered = df_filtered[df_filtered[site_col].astype(str).isin(filter_site)]
                    if filter_cable and cable_col:
                        df_filtered = df_filtered[df_filtered[cable_col].astype(str) == filter_cable]
                    if filter_color and color_col:
                        df_filtered = df_filtered[df_filtered[color_col].astype(str) == filter_color]
                    if filter_biz and biz_col:
                        df_filtered = df_filtered[df_filtered[biz_col].astype(str) == filter_biz]
                    if filter_start and start_col:
                        df_filtered = df_filtered[df_filtered[start_col].astype(str).str.contains(filter_start, na=False)]
                    if filter_start and end_col:
                        df_filtered = df_filtered[df_filtered[end_col].astype(str).str.contains(filter_start, na=False)]

                    if len(df_filtered) < len(df_display):
                        st.caption(f"当前筛选：{len(df_filtered)}条 / 总{len(df_display)}条")
                    df_display = df_filtered

                # 纤芯分配表高亮说明
                if key == "fiber" and df_display is not None and not df_display.empty:
                    st.markdown("""
                    <div style="display:flex;gap:12px;font-size:0.7rem;color:#6b7280;margin-bottom:6px;">
                        <span style="display:flex;align-items:center;gap:4px;">
                            <span style="display:inline-block;width:12px;height:12px;border-radius:2px;background:#E8F4F8;border:1px solid #bcd4de;"></span> 5G业务
                        </span>
                        <span style="display:flex;align-items:center;gap:4px;">
                            <span style="display:inline-block;width:12px;height:12px;border-radius:2px;background:#F0F0F0;border:1px solid #d0d0d0;"></span> 4G业务
                        </span>
                        <span style="display:flex;align-items:center;gap:4px;">
                            <span style="display:inline-block;width:12px;height:12px;border-radius:2px;background:#FFF8E1;border:1px solid #e6d88a;"></span> BBU-P1端子
                        </span>
                    </div>
                    """, unsafe_allow_html=True)

                st.dataframe(df_display, use_container_width=True, hide_index=True, height=400)

with tab3:
    st.markdown('<div class="section-title">合规审查结果（RK-001~RK-010）</div>', unsafe_allow_html=True)
    if st.session_state.review_results is not None:
        review_df = st.session_state.review_results.copy()

        # ===== 整改建议映射函数 =====
        def _get_suggestion(rule):
            suggestions = {
                "RK-001": "修正站点编号格式，确保符合大写字母+数字+连字符，6-25位规范",
                "RK-002": "修正站点类型为枚举值之一，参考系统预设站点类型列表",
                "RK-003": "修正网络制式为枚举值之一，参考系统预设网络制式列表",
                "RK-004": "补充线缆类型字段，路由长度>0时线缆类型为必填项",
                "RK-005": "补充接地设备数量，室外站点必须填写且>=1",
                "RK-006": "增加光口配线对数至>=12，满足5G站点最低配线要求",
                "RK-007": "修正起止点，确保每个站点的起点与终点不重复",
                "RK-008": "直流远供线路长度超过150米，需增加中继设备或改用其他取电方式",
                "RK-009": "检查端口连接规划，消除重复连接冲突",
                "RK-010": "统一BBU与AAU设备厂家，降低兼容性风险",
            }
            return suggestions.get(rule, "请核实数据后修正")

        def _remap_risk(row):
            result = str(row.get("结果", ""))
            rule = str(row.get("规则", ""))
            if rule in ["RK-001", "RK-004", "RK-005"]:
                return "高风险"
            elif rule in ["RK-002", "RK-003"]:
                return "中风险"
            elif rule in ["RK-006", "RK-008"]:
                return "高风险" if "失败" in result else "中风险"
            elif rule in ["RK-007"]:
                return "中风险" if "失败" in result else "低风险"
            elif rule in ["RK-009"]:
                return "高风险"
            elif rule in ["RK-010"]:
                return "低风险"
            return "低风险"

        # 添加风险等级列和整改建议列
        review_df["风险等级"] = review_df.apply(_remap_risk, axis=1)
        review_df["整改建议"] = review_df["规则"].apply(_get_suggestion)

        # 排序：不通过 > 警告 > 通过；同结果内 高风险 > 中风险 > 低风险
        result_order = {"失败": 0, "警告": 1, "通过": 2}
        risk_level_order = {"高风险": 0, "中风险": 1, "低风险": 2}
        review_df["_result_sort"] = review_df["结果"].apply(lambda x: result_order.get(x, 9) if isinstance(x, str) else 9)
        review_df["_risk_sort"] = review_df["风险等级"].map(risk_level_order).fillna(9).astype(int)
        review_df = review_df.sort_values(["_result_sort", "_risk_sort"]).drop(columns=["_result_sort", "_risk_sort"])

        # ===== 合并重复项（相同规则编号+站点合并）=====
        merged_rows = []
        seen = {}
        for _, row in review_df.iterrows():
            rule = str(row.get("规则", ""))
            site = str(row.get("站点", ""))
            result_val = str(row.get("结果", ""))
            key = (rule, result_val)
            if key not in seen:
                seen[key] = {"row": row, "sites": [site]}
            else:
                seen[key]["sites"].append(site)
        
        for key, data in seen.items():
            row = data["row"].copy()
            sites = data["sites"]
            if len(sites) > 1:
                # 压缩显示：取前3个 + "等N个"
                if len(sites) <= 3:
                    row["站点"] = "、".join(sites)
                else:
                    row["站点"] = f"{'、'.join(sites[:3])}等{len(sites)}个"
            merged_rows.append(row)
        
        review_df_merged = pd.DataFrame(merged_rows)

        # 统计
        fail_count = len(review_df_merged[review_df_merged["结果"].str.contains("失败", na=False)])
        warn_count = len(review_df_merged[review_df_merged["结果"].str.contains("警告", na=False)])
        pass_count = len(review_df_merged[review_df_merged["结果"].str.contains("通过", na=False)])
        high_risk_count = len(review_df_merged[review_df_merged["风险等级"].str.contains("高风险", na=False)])
        mid_risk_count = len(review_df_merged[review_df_merged["风险等级"].str.contains("中风险", na=False)])

        # ==== 审查概要卡片 ====
        st.markdown('<div class="section-title">审查概要</div>', unsafe_allow_html=True)
        alert_cols = st.columns(4)
        card_data = [
            ("不通过", fail_count, "#dc2626", "#fee2e2", "项未通过审查"),
            ("警告", warn_count, "#d97706", "#fef3c7", "项需关注"),
            ("高风险项", high_risk_count, "#b91c1c", "#fecaca", "项高风险"),
            ("通过率", f"{pass_count}/{len(review_df_merged)}", "#059669", "#dcfce7", f"{pass_count*100//max(len(review_df_merged),1)}% 合规"),
        ]
        for col, (title, val, color, bg, sub) in zip(alert_cols, card_data):
            with col:
                st.markdown(f"""
                <div style="background:{bg};border:1.5px solid {color};border-left:5px solid {color};border-radius:10px;padding:14px 16px;">
                    <div style="font-size:0.7rem;font-weight:700;color:{color};text-transform:uppercase;letter-spacing:0.5px;">{title}</div>
                    <div style="font-size:1.8rem;font-weight:800;color:#111827;line-height:1.2;">{val}</div>
                    <div style="font-size:0.72rem;color:#6b7280;">{sub}</div>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("---")

        # ==== 默认筛选：仅显示需整改项 ====
        show_all = st.checkbox("显示全部（含通过项）", value=False, key="show_all_review")
        
        if show_all:
            display_df = review_df_merged.copy()
        else:
            display_df = review_df_merged[
                review_df_merged["结果"].str.contains("失败|警告", na=False)
            ].copy()

        if display_df.empty and not show_all:
            st.success("所有审查项均已通过，无需整改！")
        else:
            st.caption(f"当前显示：{len(display_df)} 项（共 {len(review_df_merged)} 项）")

            # ==== 不通过项优先展示 ====
            fail_df = display_df[display_df["结果"].str.contains("失败", na=False)]
            if len(fail_df) > 0:
                st.markdown(f"""
                <div style="background:#fef2f2;border:2px solid #dc2626;border-radius:12px;padding:16px 20px;margin:12px 0;">
                    <div style="display:flex;align-items:center;gap:8px;margin-bottom:12px;">
                        <span style="font-size:1.0rem;font-weight:700;color:#dc2626;">不通过项</span>
                        <span style="background:#dc2626;color:#fff;font-size:0.7rem;font-weight:700;padding:2px 10px;border-radius:10px;">{len(fail_df)} 项</span>
                    </div>
                """, unsafe_allow_html=True)
                for _, row in fail_df.iterrows():
                    risk_val = str(row.get("风险等级", ""))
                    risk_color = {"高风险": "#dc2626", "中风险": "#d97706", "低风险": "#6b7280"}.get(risk_val, "#6b7280")
                    suggestion = str(row.get("整改建议", ""))
                    st.markdown(f"""
                    <div style="background:#fff;border:1px solid #fecaca;border-radius:8px;padding:12px 16px;margin-bottom:8px;display:flex;align-items:flex-start;gap:12px;">
                        <span style="background:{risk_color};color:#fff;font-size:0.65rem;font-weight:700;padding:3px 10px;border-radius:6px;flex-shrink:0;min-width:52px;text-align:center;">{risk_val}</span>
                        <div style="flex:1;">
                            <div style="font-weight:700;color:#111827;font-size:0.85rem;">{row['规则']} — {row['站点']}</div>
                            <div style="font-size:0.78rem;color:#6b7280;margin-top:4px;">{row['提示']}</div>
                            <div style="font-size:0.75rem;color:#2563eb;margin-top:3px;">整改建议: {suggestion}</div>
                        </div>
                        <span style="background:#dc2626;color:#fff;font-size:0.65rem;font-weight:700;padding:3px 10px;border-radius:6px;flex-shrink:0;">{row['结果']}</span>
                    </div>
                    """, unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)

            # ==== 警告项 ====
            warn_df = display_df[display_df["结果"].str.contains("警告", na=False)]
            if len(warn_df) > 0:
                with st.expander(f"警告项（{len(warn_df)} 项）", expanded=(len(fail_df) == 0)):
                    for _, row in warn_df.iterrows():
                        risk_val = str(row.get("风险等级", ""))
                        risk_color = {"高风险": "#dc2626", "中风险": "#d97706", "低风险": "#6b7280"}.get(risk_val, "#6b7280")
                        suggestion = str(row.get("整改建议", ""))
                        st.markdown(f"""
                        <div style="background:#fffbeb;border:1px solid #fde68a;border-radius:8px;padding:10px 14px;margin-bottom:6px;display:flex;align-items:flex-start;gap:10px;">
                            <span style="background:{risk_color};color:#fff;font-size:0.65rem;font-weight:700;padding:3px 10px;border-radius:6px;flex-shrink:0;min-width:52px;text-align:center;">{risk_val}</span>
                            <div style="flex:1;">
                                <div style="font-weight:700;color:#111827;font-size:0.85rem;">{row['规则']} — {row['站点']}</div>
                                <div style="font-size:0.78rem;color:#6b7280;margin-top:2px;">{row['提示']}</div>
                                <div style="font-size:0.75rem;color:#2563eb;margin-top:3px;">整改建议: {suggestion}</div>
                            </div>
                            <span style="background:#d97706;color:#fff;font-size:0.65rem;font-weight:700;padding:3px 10px;border-radius:6px;flex-shrink:0;">{row['结果']}</span>
                        </div>
                        """, unsafe_allow_html=True)

            # ==== 完整审查表 ====
            with st.expander(f"完整审查明细表（{len(review_df_merged)} 项）", expanded=False):
                def color_result(val):
                    if pd.isna(val):
                        return ""
                    if "失败" in str(val):
                        return "background-color:#fecaca;color:#991b1b;font-weight:700"
                    elif "警告" in str(val):
                        return "background-color:#fde68a;color:#92400e;font-weight:700"
                    return "background-color:#dcfce7;color:#166534"

                def color_risk(val):
                    colors = {
                        "高风险": "background-color:#fecaca;color:#991b1b;font-weight:700",
                        "中风险": "background-color:#fde68a;color:#92400e;font-weight:700",
                        "低风险": "background-color:#dbeafe;color:#1e40af",
                    }
                    return colors.get(str(val), "")

                # 选择要展示的列
                display_cols = [c for c in ["规则", "站点", "结果", "风险等级", "提示", "整改建议"] if c in review_df_merged.columns]
                styled = review_df_merged[display_cols].style \
                    .map(color_result, subset=["结果"]) \
                    .map(color_risk, subset=["风险等级"])
                st.dataframe(styled, use_container_width=True, hide_index=True, height=400)

        # ==== 风险分析（原风险提示标签页内容，合并为折叠区）====
        with st.expander("风险分析（YD/T 5264-2021）", expanded=False):
            if st.session_state.ai_generation_done:
                risk_content = st.session_state.ai_data.get("risk", "暂无")
                risk_cols = st.columns(3)
                with risk_cols[0]:
                    st.markdown("""
                    <div style="background:#fef2f2;border:2px solid #fca5a5;border-radius:12px;padding:16px 20px;text-align:center;">
                        <div style="font-size:1.8rem;margin-bottom:4px;">高风险</div>
                        <div style="font-weight:700;color:#dc2626;font-size:0.85rem;">取电安全 · 高空作业 · 接地</div>
                        <div style="font-size:0.72rem;color:#6b7280;margin-top:6px;">YD/T 5264-2021 强制项</div>
                    </div>
                    """, unsafe_allow_html=True)
                with risk_cols[1]:
                    st.markdown("""
                    <div style="background:#fffbeb;border:2px solid #fde68a;border-radius:12px;padding:16px 20px;text-align:center;">
                        <div style="font-size:1.8rem;color:#d97706;margin-bottom:4px;">中风险</div>
                        <div style="font-weight:700;color:#d97706;font-size:0.85rem;">线缆牵引 · 光缆熔接</div>
                        <div style="font-size:0.72rem;color:#6b7280;margin-top:6px;">工艺控制关键项</div>
                    </div>
                    """, unsafe_allow_html=True)
                with risk_cols[2]:
                    st.markdown("""
                    <div style="background:#f0fdf4;border:2px solid #bbf7d0;border-radius:12px;padding:16px 20px;text-align:center;">
                        <div style="font-size:1.8rem;color:#059669;margin-bottom:4px;">低风险</div>
                        <div style="font-weight:700;color:#059669;font-size:0.85rem;">常规检查 · 资料归档</div>
                        <div style="font-size:0.72rem;color:#6b7280;margin-top:6px;">施工日志记录项</div>
                    </div>
                    """, unsafe_allow_html=True)

                st.markdown("---")
                if risk_content and risk_content != "暂无数据" and not risk_content.startswith("生成失败"):
                    st.markdown("### 详细风险分析")
                    st.markdown(risk_content)
                else:
                    st.info("暂无详细风险分析数据")
            else:
                st.info("启动生成后，系统将依据 YD/T 5264-2021 自动分析施工风险并提供合规建议。")
    else:
        st.info("请先启动生成，系统将自动执行 RK-001~RK-010 合规审查。")

with tab4:
    st.markdown('<div class="section-title">任务中心 / 执行日志</div>', unsafe_allow_html=True)
    
    task_history = st.session_state.get("task_history", [])
    
    if not task_history:
        st.markdown("""
        <div style="background:#f3f4f6;border:1px solid #d1d5db;border-radius:10px;padding:20px;text-align:center;margin:20px 0;">
            <div style="font-size:1.1rem;font-weight:700;color:#6b7280;margin-bottom:8px;">暂无任务记录</div>
            <div style="font-size:0.8rem;color:#9ca3af;">启动分析后，任务记录将自动显示在此处</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        # 统计卡片
        total_tasks = len(task_history)
        success_count = sum(1 for t in task_history if t.get("status") == "success")
        warning_count = sum(1 for t in task_history if t.get("status") == "warning")
        failed_count = sum(1 for t in task_history if t.get("status") == "failed")
        
        stat_cols = st.columns(4)
        stat_data = [
            ("总任务数", total_tasks, "#4f46e5", "#eef2ff"),
            ("成功", success_count, "#059669", "#ecfdf5"),
            ("警告", warning_count, "#d97706", "#fffbeb"),
            ("失败", failed_count, "#dc2626", "#fef2f2"),
        ]
        for col, (title, val, color, bg) in zip(stat_cols, stat_data):
            with col:
                st.markdown(f"""
                <div style="background:{bg};border:2px solid {color}33;border-radius:10px;padding:12px 16px;text-align:center;">
                    <div style="font-size:0.7rem;font-weight:700;color:{color};">{title}</div>
                    <div style="font-size:1.6rem;font-weight:800;color:#111827;">{val}</div>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("---")
        
        # 状态筛选
        status_filter = st.selectbox("按状态筛选", ["全部", "成功", "警告", "失败"], key="task_filter")
        
        filtered = task_history
        if status_filter != "全部":
            status_map = {"成功": "success", "警告": "warning", "失败": "failed"}
            filtered = [t for t in task_history if t.get("status") == status_map.get(status_filter, "")]
        
        # 构建表格数据
        table_data = []
        for t in reversed(filtered):  # 最新在前
            status = t.get("status", "success")
            if status == "success":
                status_html = '<span style="background:#dcfce7;color:#166534;padding:2px 10px;border-radius:8px;font-weight:700;font-size:0.75rem;">成功</span>'
            elif status == "warning":
                status_html = '<span style="background:#fef3c7;color:#92400e;padding:2px 10px;border-radius:8px;font-weight:700;font-size:0.75rem;">警告</span>'
            else:
                status_html = '<span style="background:#fee2e2;color:#991b1b;padding:2px 10px;border-radius:8px;font-weight:700;font-size:0.75rem;">失败</span>'
            
            cs = t.get("compliance_summary", {})
            compliance_str = f"总{cs.get('total',0)}/失败{cs.get('fail',0)}/警告{cs.get('warn',0)}"
            
            task_id = t.get("id", "")
            
            # 操作按钮
            actions = f'<button onclick="window._task_action=\'detail_{task_id}\'" style="font-size:0.7rem;padding:2px 8px;margin:1px;">查看明细</button>'
            
            table_data.append({
                "任务ID": task_id,
                "时间": t.get("timestamp", ""),
                "状态": status_html,
                "耗时": f"{t.get('duration', 0):.1f}s",
                "合规摘要": compliance_str,
                "操作": actions,
                "_raw": t
            })
        
        # 使用 dataframe 显示（HTML列用 st.markdown 渲染不便，用 pandas）
        if table_data:
            display_df = pd.DataFrame([{
                "任务ID": d["任务ID"],
                "时间": d["时间"],
                "状态": "成功" if d["_raw"].get("status")=="success" else ("警告" if d["_raw"].get("status")=="warning" else "失败"),
                "耗时": d["耗时"],
                "合规摘要": d["合规摘要"],
                "站点数": d["_raw"].get("sites_count", 0),
            } for d in table_data])
            
            st.dataframe(display_df, use_container_width=True, hide_index=True, height=300)
            
            # 操作区
            st.markdown("---")
            op_cols = st.columns(3)
            with op_cols[0]:
                selected_task_id = st.selectbox(
                    "选择任务查看详情",
                    [d["任务ID"] for d in table_data],
                    key="task_detail_select"
                )
            
            # 查找选中任务
            selected_task = None
            for d in table_data:
                if d["任务ID"] == selected_task_id:
                    selected_task = d["_raw"]
                    break
            
            if selected_task:
                with op_cols[1]:
                    if st.button("重新分析", use_container_width=True, key="task_reanalyze"):
                        st.session_state.result_df = current_df.copy()
                        st.session_state.ai_builtin_running = True
                        st.session_state.ai_builtin_step = 1
                        st.session_state.ai_builtin_start_time = time.time()
                        st.session_state.ai_builtin_sites = current_df.copy()
                        st.session_state.cr_chunk_start = 0
                        st.session_state.cr_chunk_results = []
                        st.session_state.cr_chunk_cp = {}
                        st.session_state.excel_bytes = None
                        st.session_state.word_report_bytes = None
                        st.session_state.word_bop_bytes = None
                        st.session_state.review_failures_bytes = None
                        st.session_state.review_failed = False
                        st.session_state.offline_mode = False
                        st.session_state.ai_data = {}
                        st.session_state.ai_dataframes = {}
                        st.session_state.ai_step_times = {}
                        st.session_state.ai_generation_done = False
                        time.sleep(0.05)
                        st.rerun()
                
                with op_cols[2]:
                    if st.button("下载该任务文件", use_container_width=True, key="task_download"):
                        import io as io2
                        zip_buf = io2.BytesIO()
                        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf2:
                            if st.session_state.excel_bytes:
                                zf2.writestr("5G基站AI交付结果.xlsx", st.session_state.excel_bytes)
                            if st.session_state.word_report_bytes:
                                zf2.writestr("5G基站工程交付报告.docx", st.session_state.word_report_bytes)
                            if st.session_state.word_bop_bytes:
                                zf2.writestr("5G基站施工工艺指导书.docx", st.session_state.word_bop_bytes)
                            if st.session_state.review_failures_bytes:
                                zf2.writestr("合规审查不通过项.xlsx", st.session_state.review_failures_bytes)
                        st.download_button(
                            f"下载 {selected_task_id}.zip",
                            zip_buf.getvalue(),
                            f"{selected_task_id}.zip",
                            "application/zip",
                            use_container_width=True,
                            key=f"dl_{selected_task_id}"
                        )
                
                # 任务详情展开
                with st.expander(f"任务详情: {selected_task_id}", expanded=True):
                    st.markdown(f"**任务ID**: {selected_task.get('id', '')}")
                    st.markdown(f"**执行时间**: {selected_task.get('timestamp', '')}")
                    st.markdown(f"**状态**: {selected_task.get('status', '')}")
                    st.markdown(f"**耗时**: {selected_task.get('duration', 0):.1f}s")
                    st.markdown(f"**受影响站点**: {selected_task.get('sites_count', 0)}个")
                    
                    cs = selected_task.get("compliance_summary", {})
                    if cs:
                        st.markdown(f"**合规摘要**: 总{cs.get('total',0)}项 / 不通过{cs.get('fail',0)} / 警告{cs.get('warn',0)} / 通过{cs.get('pass',0)}")
                    
                    if selected_task.get("error_msg"):
                        st.error(f"错误信息: {selected_task['error_msg']}")
                    
                    files_list = selected_task.get("files", [])
                    if files_list:
                        st.markdown("**交付文件**: " + "、".join(files_list))
        else:
            st.info("当前筛选条件下无任务记录。")

