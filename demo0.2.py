
import os
import json
import tempfile
import pickle
import io
os.environ["PYTHONIOENCODING"] = "utf-8"
os.environ["PYTHONUTF8"] = "1"

import streamlit as st
import pandas as pd
import time
import re
import math
from io import BytesIO
from openai import OpenAI
from openpyxl.styles import Border, Side, Alignment, PatternFill, Font
from openpyxl.utils import get_column_letter
from docx.oxml.ns import qn

st.set_page_config(
    page_title="5G通信基建数智化交付系统",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==================== 现代化 CSS 主题 (v3 升级) ====================
def inject_custom_css():
    st.markdown("""
    <style>
    :root {
        --primary: #2563eb;
        --primary-light: #dbeafe;
        --primary-dark: #1e40af;
        --success: #059669;
        --success-light: #dcfce7;
        --warning: #d97706;
        --warning-light: #fef3c7;
        --danger: #dc2626;
        --danger-light: #fee2e2;
        --gray-50: #f9fafb;
        --gray-100: #f3f4f6;
        --gray-200: #e5e7eb;
        --gray-300: #d1d5db;
        --gray-400: #9ca3af;
        --gray-500: #6b7280;
        --gray-600: #4b5563;
        --gray-700: #374151;
        --gray-800: #1f2937;
        --gray-900: #111827;
        --radius-sm: 6px;
        --radius: 10px;
        --radius-lg: 12px;
        --shadow-sm: 0 1px 2px rgba(0,0,0,0.05);
        --shadow: 0 1px 3px rgba(0,0,0,0.1);
        --shadow-md: 0 4px 6px rgba(0,0,0,0.07);
        --shadow-lg: 0 10px 15px rgba(0,0,0,0.1);
    }

    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #ffffff 0%, #f8fafc 100%);
        border-right: 1px solid var(--gray-200);
    }

    [data-testid="stSidebar"] .stMarkdown h2 {
        font-size: 1.05rem !important;
        font-weight: 600 !important;
        color: var(--gray-800) !important;
        margin-top: 18px !important;
    }

    .stButton > button {
        border-radius: 8px !important;
        font-weight: 500 !important;
        padding: 8px 16px !important;
        font-family: -apple-system, BlinkMacSystemFont, sans-serif;
    }

    .stButton > button[kind="primary"] {
        background: linear-gradient(135deg, var(--primary), #3b82f6) !important;
        color: #fff !important;
        border: none !important;
        box-shadow: 0 2px 4px rgba(37,99,235,0.3) !important;
    }
    .stButton > button[kind="primary"]:hover {
        transform: translateY(-1px);
        box-shadow: 0 4px 8px rgba(37,99,235,0.4) !important;
    }

    .metric-card {
        background: linear-gradient(135deg, #fff 0%, #f8fafc 100%);
        border: 1px solid var(--gray-200);
        border-radius: var(--radius-lg);
        padding: 20px;
        box-shadow: var(--shadow-sm);
        text-align: center;
        display: flex;
        flex-direction: column;
        justify-content: center;
        align-items: center;
        min-height: 130px;
        height: 100%;
        box-sizing: border-box;
    }
    .metric-card:hover {
        box-shadow: var(--shadow-md);
        transform: translateY(-2px);
    }
    .metric-card .label {
        font-size: 0.78rem;
        font-weight: 600;
        color: var(--gray-500);
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }
    .metric-card .value {
        font-size: 1.8rem;
        font-weight: 700;
        color: var(--gray-900);
        margin: 4px 0;
    }
    .metric-card .suffix {
        font-size: 0.85rem;
        color: var(--gray-500);
    }

    .status-badge {
        display: inline-flex;
        align-items: center;
        padding: 3px 10px;
        border-radius: 20px;
        font-size: 0.72rem;
        font-weight: 600;
    }
    .status-badge.success {
        background: var(--success-light);
        color: var(--success);
    }
    .status-badge.warning {
        background: var(--warning-light);
        color: var(--warning);
    }
    .status-badge.danger {
        background: var(--danger-light);
        color: var(--danger);
    }
    .status-badge.info {
        background: var(--primary-light);
        color: var(--primary);
    }

    .main-header {
        padding: 24px;
        background: linear-gradient(135deg, #f8fafc 0%, #fff 100%);
        border-radius: var(--radius);
        border-bottom: 1px solid var(--gray-200);
        margin-bottom: 20px;
    }
    .main-header h1 {
        font-size: 1.7rem !important;
        font-weight: 700 !important;
        background: linear-gradient(135deg, var(--primary), #3b82f6);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    .main-header .subtitle {
        color: var(--gray-600);
        font-size: 0.87rem;
    }

    .stTabs [data-baseweb="tab-list"] {
        gap: 4px;
        background: var(--gray-50);
        border-radius: var(--radius);
        padding: 4px;
    }
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px !important;
        font-weight: 500 !important;
    }
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        background: #fff !important;
        color: var(--primary) !important;
    }

    [data-testid="stDataFrame"] {
        border-radius: var(--radius) !important;
        border: 1px solid var(--gray-200) !important;
    }

    .stProgress > div > div {
        background: linear-gradient(90deg, var(--primary), #60a5fa) !important;
        border-radius: 4px !important;
        height: 6px !important;
    }

    .completion-banner {
        background: linear-gradient(135deg, var(--success-light) 0%, #d1fae5 100%);
        border: 1px solid #a7f3d0;
        border-radius: var(--radius-lg);
        padding: 14px 20px;
        margin: 12px 0;
        font-weight: 600;
        color: var(--success);
        box-shadow: var(--shadow-sm);
    }

    .download-section {
        background: linear-gradient(135deg, var(--primary-light) 0%, #eff6ff 100%);
        border: 1px solid #bfdbfe;
        border-radius: var(--radius-lg);
        padding: 18px;
        margin: 12px 0;
    }

    .config-guide-box {
        background: linear-gradient(135deg, #fff7ed 0%, #fff1f2 100%);
        border: 1px solid #fed7aa;
        border-radius: var(--radius);
        padding: 14px 16px;
        margin: 10px 0;
        font-size: 0.82rem;
        line-height: 1.6;
    }

    .config-guide-box code {
        background: #fff;
        padding: 2px 6px;
        border-radius: 4px;
        font-size: 0.78rem;
        color: var(--primary);
    }

    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header[data-testid="stHeader"] {background: transparent;}
    [data-testid="stHeaderActionElements"] {display: none;}
    </style>
    """, unsafe_allow_html=True)

inject_custom_css()

# ==================== 新字段模板（16项） ====================
NEW_FIELD_TEMPLATE = [
    "站点编号", "站点名称", "站点类型", "网络制式", "覆盖场景",
    "线缆类型", "路由长度(m)", "设备总台数", "接地设备数量", "室外接头数量",
    "光口配线对数", "取电方式", "起点", "终点", "BBU型号", "AAU型号"
]

REQUIRED_FIELDS = ["站点编号", "站点名称", "站点类型", "网络制式", "覆盖场景", "取电方式"]

ENUMS = {
    "站点类型": ["室外宏站", "室内分布", "杆塔站", "塔房站", "其他"],
    "网络制式": ["4G", "5G", "4G+5G", "其他"],
    "覆盖场景": ["城区", "园区", "校园", "楼宇", "隧道", "其他"],
    "线缆类型": ["光缆", "电源线", "馈线", "接地线", "网线", "其他"],
    "取电方式": ["市电直供", "直流远供", "太阳能+储能", "混合供电", "其他"],
}

BOM_FIELDS = ["编号", "项目编码", "专业类别", "设备/材料名称", "规格型号", "单位", "数量", "项目特征", "工程量计算规则", "工作内容", "安装位置", "备注"]

DEMO_DATA = {
    "站点编号": ["SZ-BS-001", "SZ-BS-002", "GZ-BS-001", "BJ-BS-001", "SH-BS-001"],
    "站点名称": ["深圳大学城1号站", "深圳科技园2号站", "广州天河3号站", "北京海淀4号站", "上海浦东5号站"],
    "站点类型": ["室外宏站", "杆塔站", "室内分布", "塔房站", "室外宏站"],
    "网络制式": ["5G", "4G+5G", "4G", "5G", "4G+5G"],
    "覆盖场景": ["城区", "园区", "楼宇", "城区", "校园"],
    "线缆类型": ["光缆", "电源线", "馈线", "光缆", "光缆"],
    "路由长度(m)": [200, 350, 180, 420, 280],
    "设备总台数": [3, 4, 2, 3, 3],
    "接地设备数量": [2, 2, 1, 2, 2],
    "室外接头数量": [2, 3, 1, 3, 2],
    "光口配线对数": [24, 12, 8, 24, 12],
    "取电方式": ["市电直供", "直流远供", "市电直供", "混合供电", "市电直供"],
    "起点": ["BBU-P1", "AAU-P1", "ODF-P1", "BBU-P2", "AAU-P2"],
    "终点": ["AAU-P1", "ODF-P2", "BBU-P1", "AAU-P1", "ODF-P1"],
    "BBU型号": ["BBU5900", "BBU5900", "BBU3910", "BBU5900", "BBU5900"],
    "AAU型号": ["AAU5613", "AAU5636", "AAU5339", "AAU5639", "AAU5613"],
}
DEMO_DF = pd.DataFrame(DEMO_DATA)

PLATFORM_PRESETS = {
    "硅基流动": {"base_url": "https://api.siliconflow.cn/v1", "model": "deepseek-ai/DeepSeek-V3"},
    "DeepSeek官方": {"base_url": "https://api.deepseek.com", "model": "deepseek-chat"},
    "OpenAI": {"base_url": "https://api.openai.com/v1", "model": "gpt-4o"},
    "通义千问(阿里云百炼)": {"base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1", "model": "qwen-plus"},
    "智谱AI": {"base_url": "https://open.bigmodel.cn/api/paas/v4", "model": "glm-4"},
    "百度千帆": {"base_url": "https://qianfan.baidubce.com/v2", "model": "ernie-4.0-8k"},
    "火山方舟": {"base_url": "https://ark.cn-beijing.volces.com/api/v3", "model": "doubao-pro-32k"},
    "腾讯混元": {"base_url": "https://api.hunyuan.cloud.tencent.com/v1", "model": "hunyuan-pro"},
    "Groq": {"base_url": "https://api.groq.com/openai/v1", "model": "llama-3.3-70b-versatile"},
    "自定义/本地": {"base_url": "http://localhost:11434/v1", "model": "qwen2.5:7b"},
}

AI_STEPS = [
    {"key": "bom", "label": "施工BOM",
     "system": "你是5G基站施工BOM专家。按12项BOM字段生成物料清单：编号、项目编码、专业类别、设备/材料名称、规格型号、单位、数量、项目特征、工程量计算规则、工作内容、安装位置、备注。必须以标准Markdown表格格式输出，第一行为列名，第二行为分隔线，之后每行一条数据。每行必须以 | 开头和结尾。"},
    {"key": "bor", "label": "资源需求清单",
     "system": "你是5G基站施工管理专家。汇总全部站点生成资源需求清单。必须以标准Markdown表格格式输出，列名：资源类别、资源名称、规格/型号、数量、单位、备注。第一行为列名，第二行为分隔线，之后每行一条数据。每行必须以 | 开头和结尾。"},
    {"key": "bop", "label": "工艺指导书",
     "system": "你是5G基站施工工艺专家。按6章结构生成工艺指导书：一、站点勘测与规划 二、设备安装要求 三、线缆敷设要求 四、取电与接地要求 五、质量检查标准 六、施工注意事项。直接输出。"},
    {"key": "fiber", "label": "纤芯分配表",
     "system": "你是5G光缆工程专家。生成纤芯分配表。必须以标准Markdown表格格式输出，列名：站点编号、光缆编号、纤芯序号、纤芯颜色、起始端子、终止端子、业务类型。第一行为列名，第二行为分隔线，之后每行一条数据。每行必须以 | 开头和结尾。"},
    {"key": "risk", "label": "风险提示",
     "system": "你是5G通信基建审查专家。依据YD/T 5264-2021分析全部站点数据，输出施工风险提示、注意事项、合规建议。"},
]

# ==================== Session State 初始化 ====================
def init_session_state():
    defaults = {
        "uploaded_files": [{"name": "样例数据", "df": DEMO_DF, "valid": True, "errors": []}],
        "current_idx": 0,
        "result_df": None,
        "pending_upload": None,
        "ai_data": {},
        "ai_running": False,
        "ai_generation_done": False,
        "ai_step_index": 0,
        "ai_start_time": 0,
        "ai_step_times": {},
        "offline_mode": False,
        "review_results": None,
        "use_builtin_ai": True,          # v3: 默认内置模式
        "ai_platform": "硅基流动",
        "ai_base_url": "https://api.siliconflow.cn/v1",
        "ai_model": "deepseek-ai/DeepSeek-V3",
        "ai_api_key": "",
        "ai_config_expanded": False,     # v3: 默认折叠，简化界面
        "show_api_guide": False,          # v3: API获取指引
        "show_file_detail": False,
        "expand_raw_data": False,       # 上传文件后自动展开原始数据
        "excel_bytes": None,
        "word_report_bytes": None,
        "word_bop_bytes": None,
        "validation_errors": {},          # v3: 文件验证错误记录
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_session_state()

# ==================== v3 新增：文件持久化 ====================
def save_uploaded_files():
    """保存上传文件到临时目录"""
    try:
        data_to_save = []
        for item in st.session_state.uploaded_files:
            # 将 DataFrame 转换为 CSV 字符串保存
            if "df" in item:
                csv_str = item["df"].to_csv(index=False)
                data_item = {
                    "name": item["name"],
                    "csv_data": csv_str,
                    "valid": item.get("valid", True),
                    "errors": item.get("errors", []),
                    "warnings": item.get("warnings", [])
                }
                data_to_save.append(data_item)
        
        # 保存到临时文件
        temp_dir = tempfile.gettempdir()
        save_path = os.path.join(temp_dir, "5g_delivery_system_uploaded_files.json")
        with open(save_path, "w", encoding="utf-8") as f:
            json.dump(data_to_save, f, ensure_ascii=False)
    except Exception as e:
        st.error(f"保存文件失败: {e}")

def load_uploaded_files():
    """从临时目录加载上传文件"""
    try:
        temp_dir = tempfile.gettempdir()
        save_path = os.path.join(temp_dir, "5g_delivery_system_uploaded_files.json")
        if os.path.exists(save_path):
            with open(save_path, "r", encoding="utf-8") as f:
                loaded_data = json.load(f)
            
            loaded_files = []
            for item in loaded_data:
                # 从 CSV 字符串恢复 DataFrame
                df = pd.read_csv(io.StringIO(item["csv_data"]))
                loaded_item = {
                    "name": item["name"],
                    "df": df,
                    "valid": item.get("valid", True),
                    "errors": item.get("errors", []),
                    "warnings": item.get("warnings", [])
                }
                loaded_files.append(loaded_item)
            
            # 如果当前 session 没有文件，则加载
            if not st.session_state.uploaded_files or len(st.session_state.uploaded_files) == 1:
                st.session_state.uploaded_files = loaded_files
            elif loaded_files:
                # 合并已加载文件，避免重复
                existing_names = [f["name"] for f in st.session_state.uploaded_files]
                for f in loaded_files:
                    if f["name"] not in existing_names:
                        st.session_state.uploaded_files.append(f)
    except Exception as e:
        # 加载失败不影响正常使用
        pass

# 页面初始化时加载
load_uploaded_files()

# ==================== v3 新增：文件上传验证 ====================
def validate_uploaded_file(df, filename):
    """验证上传文件的字段完整性和数据有效性"""
    errors = []
    warnings = []

    # 检查必备字段是否缺失
    missing_fields = [c for c in REQUIRED_FIELDS if c not in df.columns]
    if missing_fields:
        errors.append(f"缺少必备字段: {', '.join(missing_fields)}")

    # 检查全部16项模板字段
    missing_new = [c for c in NEW_FIELD_TEMPLATE if c not in df.columns]
    if missing_new:
        warnings.append(f"缺少推荐字段(不影响使用): {', '.join(missing_new)}")

    # 检查空行
    if "站点编号" in df.columns:
        empty_ids = df["站点编号"].isna().sum() + (df["站点编号"].astype(str).str.strip() == "").sum()
        if empty_ids > 0:
            errors.append(f"存在{empty_ids}行站点编号为空，将被跳过")

    # 检查路由长度合法性
    if "路由长度(m)" in df.columns:
        try:
            route_len = pd.to_numeric(df["路由长度(m)"], errors='coerce')
            invalid_routes = route_len.isna().sum()
            zero_routes = (route_len <= 0).sum()
            if invalid_routes > 0:
                errors.append(f"存在{invalid_routes}行路由长度非数值，无法分析")
            if zero_routes > 0 and invalid_routes == 0:
                warnings.append(f"存在{zero_routes}行路由长度<=0，将触发合规审查告警")
        except:
            errors.append("路由长度列格式异常，无法解析")

    # 检查枚举值有效性
    for field, valid_values in ENUMS.items():
        if field in df.columns:
            actual_values = df[field].dropna().astype(str).unique()
            invalid = [v for v in actual_values if v not in valid_values and v != "nan"]
            if invalid:
                warnings.append(f"字段'{field}'包含非标准值: {', '.join(invalid[:5])}")

    is_valid = len(errors) == 0
    return is_valid, errors, warnings, missing_new


def read_file(f):
    """读取上传文件，支持 CSV 和 Excel"""
    try:
        if f.name.endswith(".csv"):
            return pd.read_csv(f)
        else:
            return pd.read_excel(f)
    except Exception as e:
        return None


def get_client(url, key):
    keys = [k.strip() for k in key.split(",") if k.strip()]
    return OpenAI(base_url=url, api_key=keys[0] if keys else key)


def parse_markdown_table(text):
    """从 Markdown 文本中提取表格为 DataFrame。多策略解析。"""
    if not text:
        return None

    lines = text.strip().split('\n')

    # ---- 策略 A：标准 Markdown 表格 ----
    table_lines = []
    in_table = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if in_table and table_lines:
                break
            continue
        is_pipe_line = stripped.startswith('|') or stripped.endswith('|') or '|' in stripped
        if is_pipe_line:
            is_sep = bool(re.match(r'^[\s\-:|]+$', stripped))
            if is_sep:
                in_table = True
                continue
            table_lines.append(stripped)
            in_table = True
        else:
            if in_table and table_lines:
                break
            in_table = False

    if len(table_lines) >= 2:
        df = _lines_to_df(table_lines)
        if df is not None and not df.empty:
            return df

    # ---- 策略 B：按 | 分割所有行 ----
    pipe_lines = [l.strip() for l in lines if l.strip() and '|' in l.strip()]
    pipe_lines = [l for l in pipe_lines if not re.match(r'^[\s\-:|]+$', l)]
    if len(pipe_lines) >= 2:
        df = _lines_to_df(pipe_lines)
        if df is not None and not df.empty:
            return df

    # ---- 策略 C：pd.read_csv 自动检测 ----
    try:
        df = pd.read_csv(io.StringIO(text), sep=None, engine='python', on_bad_lines='skip')
        df = df.dropna(axis=1, how='all').dropna(axis=0, how='all')
        if not df.empty and len(df.columns) >= 2:
            # 清理 Unnamed 列名
            df.columns = [c if not str(c).startswith('Unnamed') else '' for c in df.columns]
            return df
    except Exception:
        pass

    return None


def _lines_to_df(table_lines):
    """将 | 分隔的行列表转为 DataFrame，自动检测表头。"""
    # 归一化：去掉首尾 |
    normalized = []
    for tl in table_lines:
        t = tl.strip()
        if t.startswith('|'):
            t = t[1:]
        if t.endswith('|'):
            t = t[:-1]
        normalized.append(t)

    # 解析单元格
    parsed_rows, max_cols = [], 0
    for line in normalized:
        cells = [cell.strip() for cell in line.split('|')]
        max_cols = max(max_cols, len(cells))
        parsed_rows.append(cells)

    # 补齐列数
    for row in parsed_rows:
        row += [''] * (max_cols - len(row))

    if len(parsed_rows) < 2:
        return None

    # 检测第一行是否为表头（表头通常全文本，数据行常含数字）
    first_row = parsed_rows[0]
    has_numeric_in_first = any(re.search(r'\d', c) for c in first_row if c)
    second_row = parsed_rows[1]
    has_numeric_in_second = any(re.search(r'\d', c) for c in second_row if c)

    if has_numeric_in_first and not has_numeric_in_second:
        # 第一行更像是数据，第二行更像表头 → 交换
        header = second_row
        data = [first_row] + parsed_rows[2:]
    else:
        header = first_row
        data = parsed_rows[1:]

    # 确保 header 无空名且唯一
    unique_header = []
    for j, h in enumerate(header):
        h = h.strip()
        if not h:
            h = f'列{j}'
        # 确保唯一性
        base = h
        counter = 1
        while h in unique_header:
            h = f'{base}_{counter}'
            counter += 1
        unique_header.append(h)
    
    df = pd.DataFrame(data, columns=unique_header)

    # 去全空列
    empty_cols = [c for c in df.columns if df[c].astype(str).str.strip().eq('').all()]
    if empty_cols:
        df = df.drop(columns=empty_cols)

    return df


def run_compliance_review(df):
    results = []
    connection_pairs = {}
    for i, row in df.iterrows():
        site_id = str(row.get("站点编号", ""))
        site_type = str(row.get("站点类型", ""))
        network = str(row.get("网络制式", ""))
        route_len = float(row.get("路由长度(m)", 0) or 0)
        cable_type = str(row.get("线缆类型", ""))
        grounding = int(float(row.get("接地设备数量", 0) or 0))
        fiber_pairs = int(float(row.get("光口配线对数", 0) or 0))
        power_mode = str(row.get("取电方式", ""))
        start = str(row.get("起点", ""))
        end = str(row.get("终点", ""))
        bbu = str(row.get("BBU型号", ""))
        aau = str(row.get("AAU型号", ""))
        outdoor_types = ["室外宏站", "杆塔站", "塔房站"]

        if re.match(r'^[A-Z0-9-]{6,25}$', site_id):
            results.append({"站点": site_id, "规则": "RK-001", "结果": "通过", "风险": "建议", "提示": "站点编号格式正确"})
        else:
            results.append({"站点": site_id, "规则": "RK-001", "结果": "失败", "风险": "高风险", "提示": f"站点编号'{site_id}'格式不符，需满足大写字母+数字+连字符，6-25位"})

        if site_type in ENUMS["站点类型"]:
            results.append({"站点": site_id, "规则": "RK-002", "结果": "通过", "风险": "建议", "提示": "站点类型有效"})
        else:
            results.append({"站点": site_id, "规则": "RK-002", "结果": "失败", "风险": "高风险", "提示": f"站点类型'{site_type}'不在枚举范围内"})

        if network in ENUMS["网络制式"]:
            results.append({"站点": site_id, "规则": "RK-003", "结果": "通过", "风险": "建议", "提示": "网络制式有效"})
        else:
            results.append({"站点": site_id, "规则": "RK-003", "结果": "失败", "风险": "高风险", "提示": f"网络制式'{network}'不在枚举范围内"})

        if route_len > 0 and not cable_type:
            results.append({"站点": site_id, "规则": "RK-004", "结果": "失败", "风险": "高风险", "提示": "路由长度>0时必须填写线缆类型"})
        else:
            results.append({"站点": site_id, "规则": "RK-004", "结果": "通过", "风险": "建议", "提示": "线缆类型已填写"})

        if site_type in outdoor_types and grounding < 1:
            results.append({"站点": site_id, "规则": "RK-005", "结果": "失败", "风险": "高风险", "提示": f"室外站点必须填写接地设备数量且>=1"})
        else:
            results.append({"站点": site_id, "规则": "RK-005", "结果": "通过", "风险": "建议", "提示": "接地设备数量满足要求"})

        if network in ("5G", "4G+5G") and fiber_pairs < 12:
            results.append({"站点": site_id, "规则": "RK-006", "结果": "失败", "风险": "高风险", "提示": f"5G站点光口配线对数必须>=12，当前为{fiber_pairs}"})
        else:
            results.append({"站点": site_id, "规则": "RK-006", "结果": "通过", "风险": "建议", "提示": "光口配线对数满足要求"})

        if start and end and start == end:
            results.append({"站点": site_id, "规则": "RK-007", "结果": "失败", "风险": "中风险", "提示": "起点与终点不能相同"})
        else:
            results.append({"站点": site_id, "规则": "RK-007", "结果": "通过", "风险": "建议", "提示": "起止点有效"})

        if "直流远供" in power_mode and route_len > 150:
            results.append({"站点": site_id, "规则": "RK-008", "结果": "警告", "风险": "高风险", "提示": f"直流远供线路长度{route_len}米>150米，高风险"})
        else:
            results.append({"站点": site_id, "规则": "RK-008", "结果": "通过", "风险": "建议", "提示": "直流远供长度合规"})

        # RK-009: 端口重复连接冲突 — 检测是否存在重复的(起点, 终点)连接对
        line_start = str(row.get("起点", "")).strip()
        line_end = str(row.get("终点", "")).strip()
        pair_key = f"{line_start} → {line_end}"
        if line_start and line_end and line_start != "nan" and line_end != "nan":
            if pair_key not in connection_pairs:
                connection_pairs[pair_key] = []
            connection_pairs[pair_key].append(site_id)

        # RK-010: BBU与AAU厂家一致性
        if bbu and aau and bbu != "nan" and aau != "nan":
            bbu_vendor = bbu[:2] if len(bbu) >= 2 else ""
            aau_vendor = aau[:2] if len(aau) >= 2 else ""
            if bbu_vendor and aau_vendor and bbu_vendor != aau_vendor:
                results.append({"站点": site_id, "规则": "RK-010", "结果": "警告", "风险": "建议", "提示": f"BBU({bbu})与AAU({aau})厂家不一致，建议统一"})
            else:
                results.append({"站点": site_id, "规则": "RK-010", "结果": "通过", "风险": "建议", "提示": "设备厂家一致"})

    # 汇总重复连接
    all_sites = set(df["站点编号"].astype(str))
    conflict_sites = set()
    for pair, sites in connection_pairs.items():
        if len(sites) > 1:
            for s in sites:
                conflict_sites.add(s)
                results.append({"站点": s, "规则": "RK-009", "结果": "警告", "风险": "高风险",
                                "提示": f"端口连接重复冲突: {pair} 被 {len(sites)} 个站点共用（{', '.join(sites)}）"})
    for s in all_sites - conflict_sites:
        results.append({"站点": s, "规则": "RK-009", "结果": "通过", "风险": "建议", "提示": "端口连接无重复冲突"})

    return pd.DataFrame(results)


# ==================== 内置生成函数（保持原逻辑） ====================
def generate_bom_data(sites_df):
    total_sites = len(sites_df)
    L = sum(float(row.get("路由长度(m)", 0) or 0) for _, row in sites_df.iterrows())
    J_val = sum(int(float(row.get("室外接头数量", 0) or 0)) for _, row in sites_df.iterrows())
    K = sum(int(float(row.get("光口配线对数", 0) or 0)) for _, row in sites_df.iterrows())
    N_d = sum(int(float(row.get("接地设备数量", 0) or 0)) for _, row in sites_df.iterrows())

    bom_items = [
        {"编号": 1, "项目编码": "TX-GC-001", "专业类别": "通信线路工程", "设备/材料名称": "光缆", "规格型号": "GYTA-按芯数", "单位": "米", "数量": round(L * 1.06, 1), "项目特征": "单模铠装", "工程量计算规则": "路由长度x1.06", "工作内容": "敷设接续", "安装位置": "站点路由", "备注": "预留6%"},
        {"编号": 2, "项目编码": "TX-GC-002", "专业类别": "通信线路工程", "设备/材料名称": "防水套件", "规格型号": "热缩式防水胶带+胶泥套装", "单位": "套", "数量": J_val, "项目特征": "户外防水", "工程量计算规则": "室外接头数量", "工作内容": "接头防水处理", "安装位置": "室外接头处", "备注": ""},
        {"编号": 3, "项目编码": "TX-GC-003", "专业类别": "无线通信设备安装工程", "设备/材料名称": "光纤跳线", "规格型号": "LC-LC单模双芯", "单位": "条", "数量": math.ceil(K * 1.15), "项目特征": "双芯跳线", "工程量计算规则": "光口配线对数x1.15", "工作内容": "端口连接", "安装位置": "ODF-AAU/BBU", "备注": "含15%冗余"},
        {"编号": 4, "项目编码": "TX-GC-004", "专业类别": "通信线路工程", "设备/材料名称": "接地线", "规格型号": "BVR 16mm2", "单位": "条", "数量": N_d, "项目特征": "黄绿双色", "工程量计算规则": "接地设备数量", "工作内容": "接地连接", "安装位置": "设备接地排", "备注": ""},
    ]
    return pd.DataFrame(bom_items)


def generate_bor_data(sites_df):
    total_sites = len(sites_df)
    bor_data = [
        {"资源类别": "工具", "资源名称": "光纤熔接机", "规格/型号": "单芯熔接", "数量": 1, "单位": "台", "备注": "光缆接续"},
        {"资源类别": "工具", "资源名称": "OTDR", "规格/型号": "光时域反射仪", "数量": 1, "单位": "台", "备注": "光缆测试"},
        {"资源类别": "人员", "资源名称": "光纤技工", "规格/型号": "持证", "数量": 2, "单位": "人", "备注": "光缆接续"},
        {"资源类别": "人员", "资源名称": "电工", "规格/型号": "低压电工证", "数量": 1, "单位": "人", "备注": "取电接入"},
        {"资源类别": "安全", "资源名称": "安全帽", "规格/型号": "-", "数量": total_sites * 2, "单位": "个", "备注": "全员"},
    ]
    return pd.DataFrame(bor_data)


def generate_fiber_data(sites_df):
    fiber_data = []
    color_map = {1: "蓝", 2: "橙", 3: "绿", 4: "棕", 5: "灰", 6: "白", 7: "红", 8: "黑", 9: "黄", 10: "紫", 11: "粉", 12: "青"}

    for _, row in sites_df.iterrows():
        sid = row["站点编号"]
        cores = int(float(row.get("光口配线对数", 12) or 12)) // 2
        start = str(row.get("起点", "起点") or "起点")
        end = str(row.get("终点", "终点") or "终点")
        for c in range(1, min(max(cores, 4), 12) + 1):
            fiber_data.append({
                "工程编号": f"SZ-{sid}",
                "光缆编号": f"FC-{sid}-{c:02d}",
                "站点编号": sid,
                "纤芯序号": c,
                "纤芯颜色": color_map.get(c, '备用'),
                "起始端子": f"{start}-P{c}",
                "终止端子": f"{end}-P{c}",
                "BBU端口": f"BBU-P1-P{c}",
                "AAU端口": f"AAU-P1-P{c}",
                "业务类型": '数据' if c <= cores//2 else '备用'
            })
    return pd.DataFrame(fiber_data)


def generate_bop_content(sites_df):
    total_sites = len(sites_df)
    site_types = sites_df["站点类型"].value_counts().to_dict() if "站点类型" in sites_df.columns else {}
    type_summary = "、".join([f"{k}{v}个" for k, v in site_types.items()])
    power_modes = sites_df["取电方式"].value_counts().to_dict() if "取电方式" in sites_df.columns else {}
    power_summary = "、".join([f"{k}{v}个" for k, v in power_modes.items()])
    total_len = sum(float(row.get("路由长度(m)", 0) or 0) for _, row in sites_df.iterrows())

    content = f"""# 5G基站施工工艺指导书

## 一、站点勘测与规划

**勘测目标**：确认{total_sites}个站点的实际位置、环境条件、取电方式及线缆路由走向。

**站点分布**：{type_summary}，分布在{total_sites}个点位，总路由长度约{total_len:.0f}米。

**勘测内容**：
1. 确认站点经纬度坐标，核对设计图纸与实际场地的一致性
2. 检查现场供电条件，确认取电方式的可行性（{power_summary}）
3. 勘察线缆路由，测量实际距离，标记高风险路段
4. 记录现场障碍物、施工限制条件

**风险点记录**：对长距离线缆站点标记为高风险段，制定专项施工方案。

## 二、设备安装要求

**安装位置要求**：
- AAU安装于抱杆顶端，确保天线高度符合设计要求
- BBU安装于机柜内，保持通风良好，便于维护
- 设备安装应水平、牢固，水平误差<=3mm

**固定要求**：
- 宏站/杆塔站：AAU采用抱杆安装，紧固螺栓力矩>=40N-m
- 室内分布：设备壁挂或机架安装，确保承重安全
- 室外设备必须做好防水密封，防水等级>=IP65

**接地要求**：
- 所有设备必须可靠接地
- 接地电阻<=10欧（室外宏站<=4欧）
- 接地线采用黄绿双色BVR 16mm2线缆
- 防雷接地与工作接地分设，间距>=5m

## 三、线缆敷设要求

**光纤敷设规范**：
- 选用G.652D单模光缆，按芯数需求配置
- 弯曲半径>=20倍光缆外径（施工中>=30倍）
- 敷设张力不超过光缆短期允许张力的80%
- 光缆接续采用熔接方式，单芯熔接损耗<=0.05dB
- 光缆余长盘留整齐，标签清晰

**电源线敷设规范**：
- 电源线与其他线缆保持安全距离（>=150mm）
- 直流远供线路长度不超过150米（超过需增加中继）
- 电源接头做好防水、绝缘处理
- 线径根据传输距离和功率核算，确保末端压降<=5%

## 四、取电与接地要求

**取电方式**（本工程涉及：{power_summary}）：
- 市电直供：确认供电容量满足设备需求，安装自动切换装置
- 直流远供：检查远供设备输出电压、电流，配置隔离变压器和漏电保护
- 太阳能+储能：确认储能容量满足备用时间要求，定期检查电池状态

**安全规范**：
- 电工必须持低压电工证上岗
- 操作前必须确认断电，使用验电器验证
- 做好绝缘防护，使用绝缘工具
- 配电箱内空开、浪涌保护器、电表安装规范

**接地检查**：
- 施工前测量原有机房/铁塔接地电阻
- 接地体埋深>=0.6m（冻土层以下）
- 接地引入线截面积>=16mm2
- 所有接地点用防锈螺栓紧固

## 五、质量检查标准

| 检查项 | 验收指标 | 异常处理 |
|--------|----------|----------|
| 设备安装水平度 | <=1.5mm/m | 重新调整并复测 |
| 光缆熔接损耗 | <=0.05dB/点 | 重新熔接 |
| 接地电阻 | <=10欧（室外<=4欧） | 检查接地连接，增加接地体 |
| 电源电压 | 额定值+-10% | 检查供电设备，调整分接头 |
| 信号覆盖RSRP | >=-105dBm（边缘） | 调整天线方向角、下倾角 |
| 驻波比 | <=1.5 | 检查馈线接头、天线匹配 |
| 光纤接口清洁度 | 符合IEC61300-3-35 | 重新清洁并测试 |

## 六、施工注意事项

**安全风险**：
1. 高空作业必须持证上岗，佩戴安全带和安全帽
2. 带电作业必须两人以上，一人操作一人监护
3. 施工现场设置安全警示标志，禁止无关人员进入
4. 雷雨天气严禁室外作业

**现场协调**：
1. 与物业、业主提前沟通，取得施工许可
2. 施工时间避开高峰期，减少对周边影响
3. 做好现场清洁，施工结束后恢复原状
4. 与其他施工单位协调交叉作业，避免冲突

**施工记录**：
1. 填写施工日志，记录每日工作内容和人员
2. 关键工序拍照留存（设备安装、光缆接续、接地连接）
3. 整理竣工资料，包括测试报告、竣工图纸、隐蔽工程记录
4. 及时归档，确保可追溯
"""
    return content


def generate_risk_content(sites_df):
    total_sites = len(sites_df)
    outdoor_sites = 0
    long_route_sites = []
    for _, row in sites_df.iterrows():
        stype = str(row.get("站点类型", ""))
        if stype in ["室外宏站", "杆塔站", "塔房站"]:
            outdoor_sites += 1
        route_len = float(row.get("路由长度(m)", 0) or 0)
        if route_len > 300:
            long_route_sites.append((str(row.get("站点编号", "")), route_len))

    long_route_text = ""
    if long_route_sites:
        long_route_text = "\n**长距离站点**：\n"
        for sid, rlen in long_route_sites:
            long_route_text += f"- {sid}：{rlen:.0f}米，需专项施工方案\n"

    return f"""**依据YD/T 5264-2021《5G数字蜂窝移动通信网工程施工监理规范》分析**

本工程共{total_sites}个站点，其中室外站点{outdoor_sites}个。{long_route_text}

**施工风险提示：**

1. **线缆敷设风险（中）** — 长距离站点需注意牵引力控制，敷设张力不超过光缆短期允许张力的80%，避免光缆损伤。管沟或架空敷设时注意弯曲半径>=30倍光缆外径。

2. **取电安全风险（高）** — 市电直供站点需确认断电后再操作，防止触电事故。直流远供线路需配置隔离变压器和漏电保护，电压波动范围+-10%。

3. **高空作业风险（高）** — 室外宏站/杆塔站共{outdoor_sites}个，需登高作业。必须持证上岗，佩戴安全带、安全帽，雷雨天气严禁室外作业。

4. **光缆熔接质量风险（中）** — 熔接损耗超标需重新接续。单芯熔接损耗<=0.05dB，OTDR测试曲线无异常反射。光纤接口清洁度符合IEC61300-3-35标准。

5. **接地系统风险（高）** — 室外站点接地电阻<=4欧，室内站点<=10欧。接地电阻不达标影响设备安全，需定期检查。接地体埋深>=0.6m，引入线截面>=16mm2。

**合规建议：**
- 严格按照YD/T 5264-2021《5G数字蜂窝移动通信网工程施工监理规范》执行
- 施工人员必须持证上岗（电工证、登高证），做好安全培训和交底
- 关键工序必须拍照留存（设备安装、光缆接续、接地连接），确保可追溯
- 竣工后进行全面测试（RSRP>=-105dBm、SINR>=15dB、业务成功率100%）
- 填写施工日志，每日记录工作内容、人员和异常情况
"""


def style_excel_worksheet(ws, df, title=None):
    """为Excel工作表添加真实单元格样式，可选标题行"""
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    header_fill = PatternFill(start_color="2563EB", end_color="2563EB", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    data_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
    alt_fill = PatternFill(start_color="F4F6FA", end_color="F4F6FA", fill_type="solid")

    # 如果有标题，在第1行插入合并标题行，数据从第3行开始
    data_start_row = 1
    if title:
        ws.insert_rows(1)
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(df.columns))
        title_cell = ws.cell(row=1, column=1, value=title)
        title_cell.font = Font(bold=True, size=14, color="1F2937")
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        title_cell.fill = PatternFill(start_color="F0F4FF", end_color="F0F4FF", fill_type="solid")
        ws.row_dimensions[1].height = 30
        data_start_row = 2
        # 给标题行加边框
        for c in range(1, len(df.columns) + 1):
            ws.cell(row=1, column=c).border = thin_border

    for row_idx, row in enumerate(ws.iter_rows(min_row=data_start_row, max_row=ws.max_row, max_col=ws.max_column), data_start_row):
        for cell in row:
            cell.border = thin_border
            cell.alignment = Alignment(wrap_text=True, vertical='center', horizontal='center')
            if row_idx == data_start_row:
                cell.fill = header_fill
                cell.font = header_font
            else:
                cell.fill = alt_fill if (row_idx - data_start_row) % 2 == 1 else data_fill

    for col_idx, column in enumerate(df.columns, 1):
        col_letter = get_column_letter(col_idx)
        max_length = 0
        for cell in ws[col_letter]:
            try:
                if cell.value:
                    cell_str = str(cell.value)
                    length = sum(2 if ord(char) > 127 else 1 for char in cell_str)
                    max_length = max(max_length, length)
            except:
                pass
        adjusted_width = min(max(max_length + 2, 8), 50)
        ws.column_dimensions[col_letter].width = adjusted_width

    freeze_row = data_start_row + 1
    ws.freeze_panes = f"A{freeze_row}"


def create_word_document(content, title="5G基站施工工艺指导书"):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    try:
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        if style._element.rPr is not None and style._element.rPr.rFonts is not None:
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    except:
        pass

    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.name = '黑体'
        run.font.size = Pt(18)
        run.font.bold = True
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        except:
            pass

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith('# ') and not line.startswith('## '):
            heading = doc.add_heading(line[2:], level=1)
            for run in heading.runs:
                run.font.name = '黑体'
                run.font.size = Pt(16)
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                except:
                    pass
            i += 1
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            for run in heading.runs:
                run.font.name = '黑体'
                run.font.size = Pt(14)
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                except:
                    pass
            i += 1
        elif line.startswith('|') and line.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            data_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl[1:-1].split('|')]
                if all(c.replace('-','').replace(':','').replace(' ','') == '' for c in cells):
                    continue
                data_rows.append(cells)
            if data_rows:
                max_cols = max(len(r) for r in data_rows)
                table = doc.add_table(rows=len(data_rows), cols=max_cols, style='Table Grid')
                for r_idx, row in enumerate(data_rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < max_cols:
                            cell = table.cell(r_idx, c_idx)
                            cell.text = cell_text
                            if r_idx == 0:
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.bold = True
                                        run.font.name = '宋体'
                                        try:
                                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                        except:
                                            pass
                doc.add_paragraph()
        else:
            text = line.replace('**', '')
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = '宋体'
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                except:
                    pass
            i += 1
    return doc


def build_excel_bytes(current_df):
    """预构建 Excel 交付结果（一键下载优化）"""
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as w:
        sheet_titles = {"bom": "施工BOM清单", "bor": "资源需求清单", "fiber": "纤芯分配表", "review": "合规审查结果"}

        if st.session_state.review_results is not None:
            review_df = st.session_state.review_results.copy()
            review_df.to_excel(w, index=False, sheet_name="合规审查结果")
            ws = w.sheets["合规审查结果"]
            style_excel_worksheet(ws, review_df, title="合规审查结果")
            # 失败 → 红色背景；警告 → 橙色背景
            red_fill = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")
            orange_fill = PatternFill(start_color="FFE5CC", end_color="FFE5CC", fill_type="solid")
            for row in ws.iter_rows(min_row=3, max_row=ws.max_row):
                for cell in row:
                    if cell.value and "失败" in str(cell.value):
                        cell.fill = red_fill
                    elif cell.value and "警告" in str(cell.value):
                        cell.fill = orange_fill

        sheet_names = {"bom": "施工BOM", "bor": "资源需求清单", "fiber": "纤芯分配表"}
        if st.session_state.offline_mode and hasattr(st.session_state, 'ai_dataframes'):
            for key in ["bom", "bor", "fiber"]:
                if key in st.session_state.ai_dataframes:
                    df = st.session_state.ai_dataframes[key]
                    sheet_name = sheet_names.get(key, key)[:31]
                    df.to_excel(w, index=False, sheet_name=sheet_name)
                    ws = w.sheets[sheet_name]
                    style_excel_worksheet(ws, df, title=sheet_titles.get(key))
        else:
            for s in AI_STEPS:
                key = s["key"]
                if key in ("bop", "risk"):
                    continue
                content = st.session_state.ai_data.get(key, "")
                sheet_name = sheet_names.get(key, key)[:31]
                df_parsed = parse_markdown_table(content)
                if df_parsed is not None and not df_parsed.empty:
                    df_parsed.columns = [str(c).replace("*","").replace("_","").strip() for c in df_parsed.columns]
                    df_parsed.to_excel(w, index=False, sheet_name=sheet_name)
                    ws = w.sheets[sheet_name]
                    style_excel_worksheet(ws, df_parsed, title=sheet_titles.get(key))
                else:
                    pd.DataFrame({"内容": [content]}).to_excel(w, index=False, sheet_name=sheet_name)
    return output.getvalue()


def build_word_report_bytes(current_df):
    """预构建综合交付报告 Word"""
    report_content = f"""# 5G基站工程交付报告

## 一、工程概况

| 项目 | 数值 |
|------|------|
| 站点总数 | {len(current_df)}个 |
"""
    if "站点类型" in current_df.columns:
        types_count = current_df["站点类型"].value_counts().to_dict()
        for t, c in types_count.items():
            report_content += f"| {t} | {c}个 |\n"
    if "网络制式" in current_df.columns:
        network_count = current_df["网络制式"].value_counts().to_dict()
        for n, c in network_count.items():
            report_content += f"| {n} | {c}个 |\n"
    if "路由长度(m)" in current_df.columns:
        total_len = float(current_df["路由长度(m)"].sum())
        report_content += f"| 总路由长度 | {total_len:.1f}米 |\n"
    report_content += "\n"

    report_content += f"""

## 二、施工工艺指导书

{st.session_state.ai_data.get('bop', '暂无数据')}

## 三、风险提示

{st.session_state.ai_data.get('risk', '暂无数据')}

## 四、合规审查结果

"""
    if st.session_state.review_results is not None:
        review_df = st.session_state.review_results
        fail_count = len(review_df[review_df["结果"].str.contains("失败")])
        warn_count = len(review_df[review_df["结果"].str.contains("警告")])
        pass_count = len(review_df[review_df["结果"].str.contains("通过")])

        report_content += f"""
本次审查共{len(review_df)}项，其中：
- 通过：{pass_count}项
- 失败：{fail_count}项
- 警告：{warn_count}项
"""
        if fail_count > 0:
            report_content += "\n### 不通过项详情\n\n"
            for _, row in review_df[review_df["结果"].str.contains("失败")].iterrows():
                report_content += f"- **{row['站点']}** - {row['规则']}: {row['提示']}\n"
        if warn_count > 0:
            report_content += "\n### 警告项详情\n\n"
            for _, row in review_df[review_df["结果"].str.contains("警告")].iterrows():
                report_content += f"- **{row['站点']}** - {row['规则']}: {row['提示']}\n"

    doc = create_word_document(report_content, "5G基站工程交付报告")
    word_output = BytesIO()
    doc.save(word_output)
    return word_output.getvalue()


def build_word_bop_bytes():
    """预构建工艺指导书 Word"""
    bop_content = st.session_state.ai_data.get("bop", "暂无数据")
    doc = create_word_document(bop_content, "5G基站施工工艺指导书")
    word_output = BytesIO()
    doc.save(word_output)
    return word_output.getvalue()


# ==================== v3 全新侧边栏设计 ====================
with st.sidebar:
    # 顶部品牌区
    st.markdown("""
    <div style="display:flex;align-items:center;gap:12px;padding:8px 0 14px 0;border-bottom:1px solid #e5e7eb;margin-bottom:6px;">
        <div style="background:linear-gradient(135deg,#2563eb,#3b82f6);width:40px;height:40px;border-radius:10px;display:flex;align-items:center;justify-content:center;font-size:20px;box-shadow:0 2px 8px rgba(37,99,235,0.3);">🏗️</div>
        <div>
            <div style="font-size:1.1rem;font-weight:700;color:#111827;line-height:1.3;">数智化交付</div>
            <div style="font-size:0.7rem;color:#6b7280;">控制台 | demo0.2</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ========== 第一部分：数据管理（简化） ==========
    st.markdown("### 数据管理")

    # 上传区域 — 增强错误提示
    with st.container():
        def on_upload():
            if st.session_state.file_uploader:
                st.session_state.pending_upload = st.session_state.file_uploader

        st.file_uploader(
            "上传设计元数据表 / Upload Metadata Table (CSV/XLSX/XLS)",
            type=["csv", "xlsx", "xls"],
            key="file_uploader",
            on_change=on_upload
        )

    # v3 关键改进：文件上传后自动验证 + 自动切换数据源
    if st.session_state.pending_upload:
        f = st.session_state.pending_upload
        names = [x["name"] for x in st.session_state.uploaded_files]
        new_name = f.name

        if new_name not in names:
            df = read_file(f)
            if df is None:
                st.error(f"无法读取文件: {new_name}")
            else:
                # v3: 上传时立即验证
                is_valid, errors, warnings, missing_new = validate_uploaded_file(df, new_name)
                if errors:
                    st.error(f"**文件验证失败** — {new_name}")
                    for err in errors:
                        st.error(f"  {err}")
                    st.warning("文件已暂存但标记为无效，请修正后重新上传，或先用内置样本体验。")
                    # 保存但标记为无效
                    st.session_state.uploaded_files.append({
                        "name": new_name,
                        "df": df,
                        "valid": False,
                        "errors": errors,
                        "warnings": warnings
                    })
                    st.session_state.current_idx = len(st.session_state.uploaded_files) - 1
                else:
                    # 验证通过 — 自动切换数据源
                    st.session_state.uploaded_files.append({
                        "name": new_name,
                        "df": df,
                        "valid": True,
                        "errors": [],
                        "warnings": warnings
                    })
                    st.session_state.current_idx = len(st.session_state.uploaded_files) - 1
                    st.session_state.ai_generation_done = False
                    st.session_state.ai_data = {}
                    st.session_state.ai_dataframes = {}
                    st.session_state.review_results = None
                    st.session_state.result_df = None
                    st.session_state.excel_bytes = None
                    st.session_state.word_report_bytes = None
                    st.session_state.word_bop_bytes = None

                    if warnings:
                        for w in warnings:
                            st.warning(w)
                    st.success(f"文件验证通过 — 已自动切换为当前数据源")
        else:
            # 文件已存在，自动切换并清空旧生成状态
            st.session_state.current_idx = names.index(new_name)
            st.session_state.ai_generation_done = False
            st.session_state.ai_data = {}
            st.session_state.ai_dataframes = {}
            st.session_state.review_results = None
            st.session_state.result_df = None
            st.session_state.excel_bytes = None
            st.session_state.word_report_bytes = None
            st.session_state.word_bop_bytes = None

        st.session_state.pending_upload = None
        save_uploaded_files()
        st.session_state.expand_raw_data = True
        st.rerun()

    # v3: 数据源切换 — 一键切换，瞬时生效
    file_list = [x["name"] for x in st.session_state.uploaded_files]
    if st.session_state.current_idx >= len(file_list):
        st.session_state.current_idx = 0

    current_item = st.session_state.uploaded_files[st.session_state.current_idx]
    current_df = current_item["df"]

    # v3: 数据源显示 — 当前文件高亮卡片
    is_valid = current_item.get("valid", True)
    file_errors = current_item.get("errors", [])
    file_warnings = current_item.get("warnings", [])

    status_color = "#059669" if is_valid else "#dc2626"
    highlight_color = "#2563eb" if is_valid else "#dc2626"
    status_text = "有效" if is_valid else "有错误"

    st.markdown(f"""
    <div style="margin:12px 0 4px 0;">
        <span style="font-size:0.65rem;font-weight:700;color:{highlight_color};text-transform:uppercase;letter-spacing:0.5px;">◆ 当前数据源</span>
    </div>
    <div style="background:#fff;border:1.5px solid {highlight_color};border-left:4px solid {highlight_color};border-radius:8px;padding:10px 12px;box-shadow:0 1px 6px {highlight_color}22;">
        <div style="display:flex;align-items:center;justify-content:space-between;">
            <div style="font-size:0.82rem;font-weight:600;color:#111827;flex:1;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;" title="{current_item['name']}">{current_item["name"]}</div>
            <span style="font-size:0.66rem;font-weight:700;color:#fff;background:{status_color};padding:2px 10px;border-radius:10px;flex-shrink:0;margin-left:8px;">{status_text}</span>
        </div>
        <div style="font-size:0.68rem;color:#6b7280;margin-top:5px;">{len(current_df)} 站点 · {len(current_df.columns)} 字段</div>
    </div>
    """, unsafe_allow_html=True)

    # 数据源切换选择器 + 快捷操作
    if len(file_list) > 1:
        st.markdown(f"""
        <div style="background:#f0f6ff;border:1px solid #bfdbfe;border-radius:8px;padding:6px 10px;margin-bottom:6px;">
        """, unsafe_allow_html=True)
        sel = st.selectbox(
            "切换数据源",
            file_list,
            index=st.session_state.current_idx,
            label_visibility="collapsed",
            key="sb_file_sel"
        )
        st.markdown("</div>", unsafe_allow_html=True)
        if sel in file_list and file_list.index(sel) != st.session_state.current_idx:
            st.session_state.current_idx = file_list.index(sel)
            st.session_state.ai_generation_done = False
            st.session_state.ai_data = {}
            st.session_state.ai_dataframes = {}
            st.session_state.review_results = None
            st.session_state.result_df = None
            st.session_state.excel_bytes = None
            st.session_state.word_report_bytes = None
            st.session_state.word_bop_bytes = None
            st.rerun()

    col_op1, col_op2, col_op3 = st.columns(3)
    with col_op1:
        if "样例数据" not in current_item.get("name", "") and st.button("删除", use_container_width=True, key="sb_del"):
            del st.session_state.uploaded_files[file_list.index(current_item["name"])]
            st.session_state.current_idx = 0
            save_uploaded_files()
            st.rerun()
    with col_op2:
        if st.button("详情", use_container_width=True, key="sb_detail"):
            st.session_state.show_file_detail = True
            st.rerun()
    with col_op3:
        if st.button("重分析", use_container_width=True, key="sb_reanalyze"):
            st.session_state.ai_generation_done = False
            st.session_state.ai_data = {}
            st.session_state.review_results = None
            st.session_state.excel_bytes = None
            st.session_state.word_report_bytes = None
            st.session_state.word_bop_bytes = None
            st.rerun()

    st.markdown("---")

    # ========== 第二部分：AI配置（v3 极简设计） ==========
    st.markdown("### AI 引擎")

    # v3: 默认内置模式，简洁的切换
    ai_mode = st.radio(
        "模式",
        ["内置简易AI（推荐）", "外部大模型API"],
        index=0 if st.session_state.use_builtin_ai else 1,
        horizontal=False,
        key="ai_engine_mode",
        help="内置：开箱即用 | 外部：高质量，需配置"
    )
    st.session_state.use_builtin_ai = (ai_mode == "内置简易AI（推荐）")

    if st.session_state.use_builtin_ai:
        # v3: 内置模式 — 简洁提示
        st.markdown("""
        <div style="background:#dcfce7;border:1px solid #a7f3d0;border-radius:8px;padding:10px 14px;margin:8px 0;font-size:0.78rem;color:#166534;">
            <b>内置模式</b><br>
            无需配置API Key，开箱即用。基于系统内置规则自动生成施工BOM、资源清单、工艺指导书、纤芯分配表和风险提示。
        </div>
        """, unsafe_allow_html=True)
    else:
        # v3: 外部API — 可折叠配置 + 清晰引导
        with st.expander("API 配置", expanded=st.session_state.ai_config_expanded):
            # v3: 平台快速选择
            platform = st.selectbox(
                "平台",
                list(PLATFORM_PRESETS.keys()),
                index=list(PLATFORM_PRESETS.keys()).index(st.session_state.ai_platform) if st.session_state.ai_platform in PLATFORM_PRESETS else 0,
                key="sb_platform"
            )
            st.session_state.ai_platform = platform
            preset = PLATFORM_PRESETS[platform]
            is_custom = (platform == "自定义/本地")

            base_url = st.text_input(
                "API 地址",
                value=st.session_state.ai_base_url if st.session_state.ai_base_url else preset["base_url"],
                disabled=not is_custom,
                key="sb_base_url"
            )
            st.session_state.ai_base_url = base_url

            model = st.text_input(
                "模型名称",
                value=st.session_state.ai_model if st.session_state.ai_model else preset["model"],
                key="sb_model"
            )
            st.session_state.ai_model = model

            api_key = st.text_input(
                "API Key",
                type="password",
                value=st.session_state.ai_api_key,
                placeholder="sk-xxxxxxxxxxxx",
                key="sb_api_key"
            )
            st.session_state.ai_api_key = api_key

            # v3: API Key 获取指引
            if st.button("如何获取 API Key?", use_container_width=True, key="sb_guide"):
                st.session_state.show_api_guide = not st.session_state.show_api_guide
                st.rerun()

            if st.session_state.show_api_guide:
                st.markdown(f"""
                <div class="config-guide-box">
                <b>获取 {platform} API Key 步骤：</b><br>
                1. 访问 <code>{preset['base_url'].replace('/v1','').replace('/compatible-mode/v1','')}</code><br>
                2. 注册/登录后在"API Keys"或"密钥管理"中创建<br>
                3. 复制 Key 粘贴到上方输入框<br>
                4. 点击下方"测试连接"验证<br>
                <br><b>推荐平台：</b><br>
                - <b>硅基流动</b>：注册即送额度，DeepSeek-V3 免费<br>
                - <b>DeepSeek官方</b>：性价比高，10元可用很久<br>
                - <b>通义千问</b>：阿里云百炼，百万token免费额度
                </div>
                """, unsafe_allow_html=True)

            # v3: 测试连接按钮
            col_test1, col_test2 = st.columns([3, 2])
            with col_test1:
                if st.button("测试连接", use_container_width=True, key="sb_test"):
                    if not api_key:
                        st.error("请先填入 API Key")
                    elif not base_url:
                        st.error("请填写 API 地址")
                    else:
                        try:
                            client = get_client(base_url, api_key)
                            resp = client.chat.completions.create(
                                model=model,
                                messages=[{"role": "user", "content": "OK"}],
                                timeout=10,
                                max_tokens=10
                            )
                            st.success("连接成功 — API 配置有效")
                        except Exception as e:
                            st.error(f"连接失败: {str(e)[:200]}")
            with col_test2:
                if st.button("清空配置", use_container_width=True, key="sb_clear"):
                    st.session_state.ai_api_key = ""
                    st.session_state.ai_base_url = preset["base_url"]
                    st.session_state.ai_model = preset["model"]
                    st.rerun()

    st.markdown("---")

    # ========== 第三部分：启动按钮 ==========
    if not st.session_state.ai_running:
        if st.session_state.use_builtin_ai:
            btn_label = "启动分析"
        else:
            if not st.session_state.ai_api_key:
                btn_label = "启动分析（需先配置API）"
            else:
                btn_label = "启动分析"

        btn_disabled = (not st.session_state.use_builtin_ai and not st.session_state.ai_api_key)

        # v3: 无效文件阻止分析
        if not is_valid:
            st.error("当前文件存在数据错误，无法执行分析。请先修正文件或切换到其他数据源。")
            btn_disabled = True

        clicked = st.button(btn_label, type="primary", use_container_width=True, key="sb_launch", disabled=btn_disabled)
        if clicked:
            # 执行分析前再次验证
            re_valid, re_errors, re_warnings, _ = validate_uploaded_file(current_df, current_item.get("name", ""))
            if re_errors:
                st.error("文件验证失败，无法执行分析。")
            else:
                review_df = run_compliance_review(current_df)
                st.session_state.review_results = review_df
                st.session_state.excel_bytes = None
                st.session_state.word_report_bytes = None
                st.session_state.word_bop_bytes = None

                if st.session_state.use_builtin_ai:
                    # 内置模式：直接生成
                    st.session_state.result_df = current_df.copy()
                    st.session_state.ai_running = False
                    st.session_state.ai_generation_done = True
                    st.session_state.offline_mode = True
                    st.session_state.ai_step_times = {"total": 0}

                    sites = current_df.copy()
                    bom_df = generate_bom_data(sites)
                    bor_df = generate_bor_data(sites)
                    fiber_df = generate_fiber_data(sites)
                    bop_content = generate_bop_content(sites)
                    risk_content = generate_risk_content(sites)

                    st.session_state.ai_data = {
                        "bom": bom_df.to_markdown(index=False),
                        "bor": bor_df.to_markdown(index=False),
                        "bop": bop_content,
                        "fiber": fiber_df.to_markdown(index=False),
                        "risk": risk_content
                    }
                    st.session_state.ai_dataframes = {
                        "bom": bom_df, "bor": bor_df, "fiber": fiber_df
                    }

                    st.session_state.excel_bytes = build_excel_bytes(current_df)
                    st.session_state.word_report_bytes = build_word_report_bytes(current_df)
                    st.session_state.word_bop_bytes = build_word_bop_bytes()
                    time.sleep(0.15)
                    st.rerun()
                else:
                    # 外部API模式
                    st.session_state.result_df = current_df.copy()
                    st.session_state.ai_running = True
                    st.session_state.ai_generation_done = False
                    st.session_state.ai_step_index = 0
                    st.session_state.ai_start_time = time.time()
                    st.session_state.ai_step_times = {}
                    st.session_state.ai_data = {}
                    st.session_state.offline_mode = False
                    time.sleep(0.15)
                    st.rerun()

    # AI运行中状态
    if st.session_state.ai_running:
        si = st.session_state.ai_step_index
        elapsed = time.time() - st.session_state.ai_start_time
        total_steps = len(AI_STEPS)

        st.markdown(f"""
        <div style="background:#fef3c7;border:1px solid #fde68a;border-radius:8px;padding:10px 14px;margin:8px 0;">
            <span style="font-size:0.8rem;font-weight:600;color:#92400e;">生成中 {si}/{total_steps} · {int(elapsed)}s</span>
        </div>
        """, unsafe_allow_html=True)

        for idx, step in enumerate(AI_STEPS):
            if idx < si:
                t = st.session_state.ai_step_times.get(step['key'], 0)
                st.success(f"  {step['label']} ({t:.1f}s)")
            elif idx == si:
                st.info(f"  {step['label']} ...")
            else:
                st.caption(f"  {step['label']}")

        if st.button("取消生成", use_container_width=True, key="sb_cancel"):
            st.session_state.ai_running = False
            st.session_state.ai_generation_done = True
            st.rerun()

    # ========== 第四部分：一键下载 ==========
    if st.session_state.ai_generation_done and not st.session_state.ai_running:
        total = st.session_state.ai_step_times.get("total", 0)
        st.markdown(f"""
        <div class="completion-banner">
            生成完成 · 总耗时 {total:.1f}s
        </div>
        """, unsafe_allow_html=True)

        st.markdown("### 下载交付结果")

        if st.session_state.excel_bytes is None:
            st.session_state.excel_bytes = build_excel_bytes(current_df)
        if st.session_state.word_report_bytes is None:
            st.session_state.word_report_bytes = build_word_report_bytes(current_df)
        if st.session_state.word_bop_bytes is None:
            st.session_state.word_bop_bytes = build_word_bop_bytes()

        st.download_button(
            "Excel 交付结果",
            st.session_state.excel_bytes,
            "5G基站AI交付结果.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True, key="sb_dl_excel"
        )

        st.download_button(
            "综合交付报告 (Word)",
            st.session_state.word_report_bytes,
            "5G基站工程交付报告.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, key="sb_dl_report"
        )

        st.download_button(
            "工艺指导书 (Word)",
            st.session_state.word_bop_bytes,
            "5G基站施工工艺指导书.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, key="sb_dl_bop"
        )


# ==================== 主页面 (v3 升级) ====================
st.markdown("""
<div class="main-header">
    <h1>5G通信基建数智化交付系统</h1>
    <p class="subtitle">RK-001~RK-010 合规审查引擎 · 12项标准BOM · 16项站点字段 · Word/Excel 双格式交付 · demo0.2</p>
</div>
""", unsafe_allow_html=True)

# 状态提示 — 紧凑条形式
mode_color = "#059669" if st.session_state.use_builtin_ai else "#6b7280"
mode_bg = "#dcfce7" if st.session_state.use_builtin_ai else "#f3f4f6"
mode_text = "内置简易AI" if st.session_state.use_builtin_ai else "离线模式"
mode_hint = "开箱即用，无需配置" if st.session_state.use_builtin_ai else "合规审查可用，配置API可启用全量生成"
st.markdown(f"""
<div style="background:{mode_bg};border:1px solid {mode_color}22;border-left:3px solid {mode_color};border-radius:8px;padding:8px 16px;margin-bottom:16px;display:flex;align-items:center;gap:10px;font-size:0.8rem;">
    <span style="font-weight:700;color:{mode_color};">{mode_text}</span>
    <span style="color:#6b7280;">{mode_hint}</span>
</div>
""", unsafe_allow_html=True)

current_df = st.session_state.uploaded_files[st.session_state.current_idx]["df"]

# ===== 工程概览卡片行 (仪表盘风格) =====
st.subheader("工程概览")

# 快速状态栏
has_review = st.session_state.review_results is not None
has_generated = st.session_state.ai_generation_done
if has_review:
    rdf = st.session_state.review_results
    r_fail = len(rdf[rdf["结果"].str.contains("失败", na=False)])
    r_warn = len(rdf[rdf["结果"].str.contains("警告", na=False)])
    if r_fail > 0:
        st.markdown(f"""
        <div style="background:#fef2f2;border:1.5px solid #fca5a5;border-radius:8px;padding:10px 16px;margin-bottom:14px;display:flex;align-items:center;gap:12px;">
            <span style="font-weight:700;color:#dc2626;font-size:0.82rem;">审查状态</span>
            <span style="background:#dc2626;color:#fff;padding:2px 12px;border-radius:10px;font-size:0.7rem;font-weight:700;">{r_fail} 项不通过</span>
            {f'<span style="background:#f59e0b;color:#fff;padding:2px 12px;border-radius:10px;font-size:0.7rem;font-weight:700;">{r_warn} 项警告</span>' if r_warn > 0 else ''}
            <span style="color:#6b7280;font-size:0.75rem;margin-left:auto;">请前往「合规审查」标签页查看详情</span>
        </div>
        """, unsafe_allow_html=True)

ov_cols = st.columns(5)

with ov_cols[0]:
    st.markdown(f"""
    <div class="metric-card">
        <div class="label">站点数量</div>
        <div class="value">{len(current_df)} <span class="suffix">个</span></div>
    </div>
    """, unsafe_allow_html=True)

with ov_cols[1]:
    if "站点类型" in current_df.columns:
        types = current_df["站点类型"].value_counts().to_dict()
        items_html = "".join([f'<div style="font-size:0.78rem;color:#374151;padding:2px 0;">{k}: <b>{v}</b>个</div>' for k,v in types.items()])
        st.markdown(f"""
        <div class="metric-card">
            <div class="label">站点类型分布</div>
            <div style="padding-top:2px;">{items_html}</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown(f"""
        <div class="metric-card">
            <div class="label">站点类型分布</div>
            <div style="font-size:0.9rem;color:#9ca3af;">— 暂无数据</div>
        </div>
        """, unsafe_allow_html=True)

with ov_cols[2]:
    if "路由长度(m)" in current_df.columns and not current_df["路由长度(m)"].dropna().empty:
        total_len = int(float(current_df["路由长度(m)"].sum()))
        max_len = int(float(current_df["路由长度(m)"].max()))
        st.markdown(f"""
        <div class="metric-card">
            <div class="label">总路由长度</div>
            <div class="value">{total_len} <span class="suffix">米</span></div>
            <div style="font-size:0.72rem;color:#6b7280;margin-top:2px;">最长单段: {max_len}米</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown(f"""
        <div class="metric-card">
            <div class="label">总路由长度</div>
            <div style="font-size:0.9rem;color:#9ca3af;padding:8px 0;display:flex;align-items:center;justify-content:center;gap:6px;">
                <span style="font-size:1.2rem;">—</span> 暂无数据
            </div>
            <div style="font-size:0.72rem;color:#d1d5db;margin-top:2px;">推荐补充字段</div>
        </div>
        """, unsafe_allow_html=True)

with ov_cols[3]:
    if "取电方式" in current_df.columns:
        modes = current_df["取电方式"].value_counts().to_dict()
        items_html = "".join([f'<div style="font-size:0.78rem;color:#374151;padding:2px 0;">{k}: <b>{v}</b></div>' for k,v in modes.items()])
        st.markdown(f"""
        <div class="metric-card">
            <div class="label">取电方式</div>
            <div style="padding-top:2px;">{items_html}</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown(f"""
        <div class="metric-card">
            <div class="label">取电方式</div>
            <div style="font-size:0.9rem;color:#9ca3af;">— 暂无数据</div>
        </div>
        """, unsafe_allow_html=True)

with ov_cols[4]:
    if "AAU型号" in current_df.columns:
        aaus = current_df["AAU型号"].dropna().unique().tolist()
        bbus = current_df["BBU型号"].dropna().unique().tolist()
        st.markdown(f"""
        <div class="metric-card">
            <div class="label">设备型号</div>
            <div style="font-size:0.8rem;color:#374151;padding-top:4px;">
                AAU: {len(aaus)}种<br>
                BBU: {len(bbus)}种
            </div>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown(f"""
        <div class="metric-card">
            <div class="label">设备型号</div>
            <div style="font-size:0.9rem;color:#9ca3af;">— 暂无数据</div>
        </div>
        """, unsafe_allow_html=True)

st.markdown("---")

# ===== v3: 文件详情弹窗 =====
if st.session_state.show_file_detail:
    current_item = st.session_state.uploaded_files[st.session_state.current_idx]
    with st.expander("文件详情", expanded=True):
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            st.markdown(f"**文件名**: {current_item.get('name','')}")
            st.markdown(f"**验证状态**: {'有效' if current_item.get('valid', True) else '有错误'}")
            st.markdown(f"**站点数**: {len(current_item['df'])}")
            st.markdown(f"**字段数**: {len(current_item['df'].columns)}")
        with col_d2:
            errors = current_item.get('errors', [])
            warnings = current_item.get('warnings', [])
            if errors:
                st.error("**错误:**")
                for e in errors:
                    st.error(f"- {e}")
            if warnings:
                st.warning("**警告:**")
                for w in warnings:
                    st.warning(f"- {w}")
        st.dataframe(current_item["df"], use_container_width=True, hide_index=True)
    if st.button("关闭详情", key="close_detail"):
        st.session_state.show_file_detail = False
        st.rerun()

# ===== 标签页 =====
tab1, tab2, tab3, tab4 = st.tabs(["工程概览", "AI生成施工资料", "合规审查", "风险提示"])

with tab1:
    # v3: 改进字段校验展示
    missing_new = [c for c in NEW_FIELD_TEMPLATE if c not in current_df.columns]
    if not missing_new:
        st.success(f"16项字段齐全 · {len(current_df)}站点 · {len(current_df.columns)}字段")
    else:
        st.warning(f"缺少推荐字段: {', '.join(missing_new)}（不影响内置分析）")

    with st.expander("查看完整原始数据", expanded=st.session_state.expand_raw_data):
        st.dataframe(current_df, use_container_width=True, hide_index=True)
        st.session_state.expand_raw_data = False

    # AI生成进度
    if st.session_state.ai_running:
        si = st.session_state.ai_step_index
        st.subheader("AI 生成进度")
        progress_cols = st.columns(len(AI_STEPS))
        for idx, (col, s) in enumerate(zip(progress_cols, AI_STEPS)):
            with col:
                with st.container(border=True):
                    if idx < si:
                        t = st.session_state.ai_step_times.get(s['key'], 0)
                        st.success(f"  {s['label']}")
                        st.caption(f"{t:.1f}s")
                    elif idx == si:
                        st.warning(f"  {s['label']}")
                    else:
                        st.info(f"  {s['label']}")

        elapsed = time.time() - st.session_state.ai_start_time
        st.progress(si / len(AI_STEPS), text=f"总进度 {si}/{len(AI_STEPS)} · {int(elapsed)}s")

        if si < len(AI_STEPS):
            step = AI_STEPS[si]
            try:
                client = get_client(st.session_state.ai_base_url, st.session_state.ai_api_key)
                t0 = time.time()
                rv = (step["key"] == "risk")
                resp = client.chat.completions.create(
                    model=st.session_state.ai_model,
                    messages=[
                        {"role": "system", "content": step["system"]},
                        {"role": "user", "content": f"站点数据（{len(st.session_state.result_df)}个）：\n{st.session_state.result_df.to_string()}"}
                    ],
                    timeout=90 if rv else 180,
                    temperature=0.3,
                    max_tokens=2048 if rv else 4096
                )
                st.session_state.ai_data[step["key"]] = resp.choices[0].message.content
                st.session_state.ai_step_times[step["key"]] = time.time() - t0
            except Exception as e:
                st.session_state.ai_data[step["key"]] = f"生成失败: {str(e)[:300]}"
            st.session_state.ai_step_index = si + 1
            time.sleep(0.3)
            st.rerun()
        else:
            st.session_state.ai_step_times["total"] = time.time() - st.session_state.ai_start_time
            st.session_state.ai_running = False
            st.session_state.ai_generation_done = True
            st.session_state.excel_bytes = build_excel_bytes(current_df)
            st.session_state.word_report_bytes = build_word_report_bytes(current_df)
            st.session_state.word_bop_bytes = build_word_bop_bytes()
            st.rerun()

    if st.session_state.ai_generation_done and not st.session_state.ai_running:
        st.success(f"全部生成完成 · 总耗时 {st.session_state.ai_step_times.get('total',0):.1f}s")

        st.markdown("---")
        with st.container(border=True):
            st.subheader("一键下载交付结果")
            dl_col1, dl_col2, dl_col3 = st.columns(3)
            with dl_col1:
                if st.session_state.excel_bytes is None:
                    st.session_state.excel_bytes = build_excel_bytes(current_df)
                st.download_button(
                    "Excel 交付包",
                    st.session_state.excel_bytes,
                    "5G基站AI交付结果.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True, key="main_dl_excel"
                )
            with dl_col2:
                if st.session_state.word_report_bytes is None:
                    st.session_state.word_report_bytes = build_word_report_bytes(current_df)
                st.download_button(
                    "Word 交付报告",
                    st.session_state.word_report_bytes,
                    "5G基站工程交付报告.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, key="main_dl_report"
                )
            with dl_col3:
                if st.session_state.word_bop_bytes is None:
                    st.session_state.word_bop_bytes = build_word_bop_bytes()
                st.download_button(
                    "Word 工艺指导书",
                    st.session_state.word_bop_bytes,
                    "5G基站施工工艺指导书.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, key="main_dl_bop"
                )

with tab2:
    st.subheader("AI 生成施工资料")
    if not st.session_state.ai_generation_done:
        st.markdown("""
        <div style="background:#f3f4f6;border:1px solid #d1d5db;border-radius:10px;padding:20px;text-align:center;margin:20px 0;">
            <div style="font-size:1.1rem;font-weight:700;color:#6b7280;margin-bottom:8px;">尚未生成施工资料</div>
            <div style="font-size:0.8rem;color:#9ca3af;">请先在左侧控制台底部点击「启动分析」</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        status_cols = st.columns(5)
        keys = ["bom", "bor", "bop", "fiber", "risk"]
        labels = ["施工BOM", "资源清单", "工艺指导书", "纤芯分配表", "风险提示"]
        icons = ["BOM", "BOR", "BOP", "FC", "RK"]

        for col, key, label, icon in zip(status_cols, keys, labels, icons):
            with col:
                content = st.session_state.ai_data.get(key, "")
                has_data = content and content != "暂无数据" and not content.startswith("生成失败")
                bg = "#f0fdf4" if has_data else "#fef2f2"
                border = "#86efac" if has_data else "#fca5a5"
                text_color = "#166534" if has_data else "#991b1b"
                badge_bg = "#dcfce7" if has_data else "#fee2e2"
                badge_text = "已生成" if has_data else "未生成"
                st.markdown(f"""
                <div style="background:{bg};border:1.5px solid {border};border-radius:10px;padding:12px 14px;text-align:center;">
                    <div style="font-size:0.7rem;font-weight:700;color:#6b7280;text-transform:uppercase;letter-spacing:0.5px;margin-bottom:4px;">{icon}</div>
                    <div style="font-weight:700;color:#111827;font-size:0.85rem;margin-bottom:8px;">{label}</div>
                    <span style="background:{badge_bg};color:{text_color};font-size:0.65rem;font-weight:700;padding:2px 12px;border-radius:8px;">{badge_text}</span>
                    </div>
                """, unsafe_allow_html=True)

        st.markdown("---")

        all_sites = list(st.session_state.result_df["站点编号"]) if st.session_state.result_df is not None else []
        view_mode = st.radio("查看模式", ["全部站点汇总", "单站点详细"], horizontal=True, key="tab2_view")
        selected_site = None
        if "单站点" in view_mode and all_sites:
            selected_site = st.selectbox("选择站点", all_sites, key="tab2_site")

        subtabs = st.tabs(["施工BOM", "资源需求清单", "工艺指导书", "纤芯分配表"])
        sub_keys = ["bom", "bor", "bop", "fiber"]
        for i, t in enumerate(subtabs):
            with t:
                key = sub_keys[i]
                # BOP（工艺指导书）始终以 Markdown 渲染
                if key == "bop":
                    content = st.session_state.ai_data.get("bop", "暂无数据")
                    if not content or content == "暂无数据":
                        st.info("暂无工艺指导书数据，请先生成。")
                    else:
                        st.markdown(content)
                    continue

                # 优先使用存储的 DataFrame（内置模式），否则从文本解析（外部 API 模式）
                df_display = st.session_state.ai_dataframes.get(key)
                if df_display is None:
                    content = st.session_state.ai_data.get(key, "")
                    if not content or content == "暂无数据" or content.startswith("生成失败"):
                        st.info(f"暂无{key}数据，请先生成。")
                        continue
                    df_display = parse_markdown_table(content)
                    if df_display is None or df_display.empty:
                        with st.container(height=400):
                            st.text(content)
                        continue

                # 纤芯分配表支持单站点筛选
                if key == "fiber" and selected_site:
                    if "站点编号" in df_display.columns:
                        df_display = df_display[df_display["站点编号"].astype(str) == str(selected_site)]
                    else:
                        st.caption("（纤芯分配表不含站点编号列，无法按站点筛选）")

                # 摘要行：条数 + 关键指标
                summary_parts = [f"共 **{len(df_display)}** 条"]
                if key == "fiber":
                    if "站点编号" in df_display.columns:
                        sites_n = df_display["站点编号"].nunique()
                        summary_parts.append(f"覆盖 **{sites_n}** 个站点")
                elif key == "bor":
                    for col_name in ["物料类别", "类别", "资源类型"]:
                        if col_name in df_display.columns:
                            types = df_display[col_name].value_counts()
                            type_str = " · ".join([f"{k}: {v}" for k, v in types.head(5).items()])
                            if type_str:
                                summary_parts.append(f"类型: {type_str}")
                            break
                st.caption(" · ".join(summary_parts))

                st.dataframe(df_display, use_container_width=True, hide_index=True)

with tab3:
    st.subheader("合规审查结果（RK-001~RK-010）")
    if st.session_state.review_results is not None:
        review_df = st.session_state.review_results.copy()

        # 风险等级排序权重
        risk_order = {"高风险": 0, "中风险": 1, "低风险": 2, "建议": 3}
        result_order = {"失败": 0, "警告": 1, "通过": 2}

        # 添加排序列
        review_df["_risk_sort"] = review_df["风险"].map(risk_order).fillna(9).astype(int)
        review_df["_result_sort"] = review_df["结果"].apply(lambda x: result_order.get(x, 9) if isinstance(x, str) else 9)
        review_df = review_df.sort_values(["_result_sort", "_risk_sort"]).drop(columns=["_risk_sort", "_result_sort"])

        # 统计
        fail_count = len(review_df[review_df["结果"].str.contains("失败", na=False)])
        warn_count = len(review_df[review_df["结果"].str.contains("警告", na=False)])
        pass_count = len(review_df[review_df["结果"].str.contains("通过", na=False)])
        high_risk_count = len(review_df[review_df["风险"].str.contains("高风险", na=False)])
        mid_risk_count = len(review_df[review_df["风险"].str.contains("中风险", na=False)])

        # ==== 第一层：关键告警摘要卡片 ====
        st.markdown("### 审查概要")

        # 告警汇总卡片行
        alert_cols = st.columns(4)
        card_data = [
            ("不通过", fail_count, "#dc2626", "#fee2e2", "项未通过审查"),
            ("警告", warn_count, "#d97706", "#fef3c7", "项需关注"),
            ("高风险项", high_risk_count, "#b91c1c", "#fecaca", "项高风险"),
            ("通过率", f"{pass_count}/{len(review_df)}", "#059669", "#dcfce7", f"{pass_count*100//max(len(review_df),1)}% 合规"),
        ]
        for col, (title, val, color, bg, sub) in zip(alert_cols, card_data):
            with col:
                st.markdown(f"""
                <div style="background:{bg};border:1.5px solid {color};border-left:5px solid {color};border-radius:10px;padding:14px 16px;">
                    <div style="font-size:0.7rem;font-weight:700;color:{color};text-transform:uppercase;letter-spacing:0.5px;">{title}</div>
                    <div style="font-size:1.8rem;font-weight:800;color:#111827;line-height:1.2;">{val}</div>
                    <div style="font-size:0.72rem;color:#6b7280;">{sub}</div>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("---")

        # ==== 第二层：不通过项优先展示（红色告警区）====
        if fail_count > 0:
            fail_df = review_df[review_df["结果"].str.contains("失败", na=False)]
            st.markdown(f"""
            <div style="background:#fef2f2;border:2px solid #dc2626;border-radius:12px;padding:16px 20px;margin:12px 0;">
                <div style="display:flex;align-items:center;gap:8px;margin-bottom:12px;">
                    <span style="font-size:1.0rem;font-weight:700;color:#dc2626;">不通过项</span>
                    <span style="background:#dc2626;color:#fff;font-size:0.7rem;font-weight:700;padding:2px 10px;border-radius:10px;">{fail_count} 项</span>
                </div>
            """, unsafe_allow_html=True)
            for _, row in fail_df.iterrows():
                risk_color = {"高风险": "#dc2626", "中风险": "#d97706", "低风险": "#6b7280", "建议": "#9ca3af"}.get(row.get("风险", ""), "#6b7280")
                st.markdown(f"""
                <div style="background:#fff;border:1px solid #fecaca;border-radius:8px;padding:12px 16px;margin-bottom:8px;display:flex;align-items:flex-start;gap:12px;">
                    <span style="background:{risk_color};color:#fff;font-size:0.65rem;font-weight:700;padding:3px 10px;border-radius:6px;flex-shrink:0;min-width:52px;text-align:center;">{row.get('风险','')}</span>
                    <div style="flex:1;">
                        <div style="font-weight:700;color:#111827;font-size:0.85rem;">{row['规则']} — {row['站点']}</div>
                        <div style="font-size:0.78rem;color:#6b7280;margin-top:4px;">{row['提示']}</div>
                    </div>
                    <span style="background:#dc2626;color:#fff;font-size:0.65rem;font-weight:700;padding:3px 10px;border-radius:6px;flex-shrink:0;">{row['结果']}</span>
                </div>
                """, unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

        # ==== 第三层：警告项 ====
        if warn_count > 0:
            warn_df = review_df[review_df["结果"].str.contains("警告", na=False)]
            with st.expander(f"警告项（{warn_count} 项）", expanded=(fail_count == 0)):
                for _, row in warn_df.iterrows():
                    risk_color = {"高风险": "#dc2626", "中风险": "#d97706", "低风险": "#6b7280", "建议": "#9ca3af"}.get(row.get("风险", ""), "#6b7280")
                    st.markdown(f"""
                    <div style="background:#fffbeb;border:1px solid #fde68a;border-radius:8px;padding:10px 14px;margin-bottom:6px;display:flex;align-items:flex-start;gap:10px;">
                        <span style="background:{risk_color};color:#fff;font-size:0.65rem;font-weight:700;padding:3px 10px;border-radius:6px;flex-shrink:0;min-width:52px;text-align:center;">{row.get('风险','')}</span>
                        <div style="flex:1;">
                            <div style="font-weight:700;color:#111827;font-size:0.85rem;">{row['规则']} — {row['站点']}</div>
                            <div style="font-size:0.78rem;color:#6b7280;margin-top:2px;">{row['提示']}</div>
                        </div>
                        <span style="background:#d97706;color:#fff;font-size:0.65rem;font-weight:700;padding:3px 10px;border-radius:6px;flex-shrink:0;">{row['结果']}</span>
                    </div>
                    """, unsafe_allow_html=True)

        # ==== 第四层：完整审查表（可折叠）====
        with st.expander(f"完整审查明细表（{len(review_df)} 项）", expanded=False):
            # 颜色映射函数
            def color_result(val):
                if "失败" in str(val):
                    return "background-color:#fecaca;color:#991b1b;font-weight:700"
                elif "警告" in str(val):
                    return "background-color:#fde68a;color:#92400e;font-weight:700"
                return "background-color:#dcfce7;color:#166534"

            def color_risk(val):
                colors = {
                    "高风险": "background-color:#fecaca;color:#991b1b;font-weight:700",
                    "中风险": "background-color:#fde68a;color:#92400e;font-weight:700",
                    "低风险": "background-color:#dbeafe;color:#1e40af",
                    "建议": "background-color:#f3f4f6;color:#6b7280",
                }
                return colors.get(str(val), "")

            styled = review_df.style \
                .map(color_result, subset=["结果"]) \
                .map(color_risk, subset=["风险"])
            st.dataframe(styled, use_container_width=True, hide_index=True, height=400)
    else:
        st.info("请先启动生成，系统将自动执行 RK-001~RK-010 合规审查。")

with tab4:
    st.subheader("风险提示与注意事项")
    if st.session_state.ai_generation_done:
        content = st.session_state.ai_data.get("risk", "暂无")

        # ==== 风险速览卡片 ====
        risk_cols = st.columns(3)
        with risk_cols[0]:
            st.markdown(f"""
            <div style="background:#fef2f2;border:2px solid #fca5a5;border-radius:12px;padding:16px 20px;text-align:center;">
                <div style="font-size:1.8rem;margin-bottom:4px;">高风险</div>
                <div style="font-weight:700;color:#dc2626;font-size:0.85rem;">取电安全 · 高空作业 · 接地</div>
                <div style="font-size:0.72rem;color:#6b7280;margin-top:6px;">YD/T 5264-2021 强制项</div>
            </div>
            """, unsafe_allow_html=True)
        with risk_cols[1]:
            st.markdown(f"""
            <div style="background:#fffbeb;border:2px solid #fde68a;border-radius:12px;padding:16px 20px;text-align:center;">
                <div style="font-size:1.8rem;color:#d97706;margin-bottom:4px;">中风险</div>
                <div style="font-weight:700;color:#d97706;font-size:0.85rem;">线缆牵引 · 光缆熔接</div>
                <div style="font-size:0.72rem;color:#6b7280;margin-top:6px;">工艺控制关键项</div>
            </div>
            """, unsafe_allow_html=True)
        with risk_cols[2]:
            st.markdown(f"""
            <div style="background:#f0fdf4;border:2px solid #bbf7d0;border-radius:12px;padding:16px 20px;text-align:center;">
                <div style="font-size:1.8rem;color:#059669;margin-bottom:4px;">低风险</div>
                <div style="font-weight:700;color:#059669;font-size:0.85rem;">常规检查 · 资料归档</div>
                <div style="font-size:0.72rem;color:#6b7280;margin-top:6px;">施工日志记录项</div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("---")

        # ==== AI 详细分析内容 ====
        if content and content != "暂无数据" and not content.startswith("生成失败"):
            st.markdown("### 详细风险分析")
            st.markdown(content)
        else:
            st.info("暂无详细风险分析数据")
    else:
        st.markdown("""
        <div style="background:#f3f4f6;border:1px solid #d1d5db;border-radius:10px;padding:20px;text-align:center;">
            <div style="font-size:1.1rem;font-weight:700;color:#6b7280;margin-bottom:8px;">尚未生成风险分析</div>
            <div style="font-size:0.8rem;color:#9ca3af;">启动生成后，系统将依据 YD/T 5264-2021 自动分析施工风险并提供合规建议</div>
        </div>
        """, unsafe_allow_html=True)
